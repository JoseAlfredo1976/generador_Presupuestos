#!/usr/bin/env python3
"""
Interfaz web del Generador de Presupuestos + Analisis IA - Grupo Europa
Uso: python app.py  -> abre http://localhost:5000
"""
import json
import logging
import os
import secrets
import sys
import tempfile
import threading
import time
import unicodedata
import uuid
from datetime import datetime
from pathlib import Path

from flask import (Flask, abort, jsonify, redirect, render_template,
                   request, send_file, url_for)
from dotenv import load_dotenv

load_dotenv(Path(__file__).parent / ".env")

from core.ai_analyst import (extract_partidas_from_subcontrata, extract_solicitud_data,
                             extract_subcontrata_datos, extract_full_presupuesto)
from core.docx_generator import DocxGenerator
from core.excel_generator import ExcelGenerator
from core.html_generator import HtmlGenerator
from core.pdf_converter import PdfConverter
from utils.tarifa_loader import TarifaLoader
from utils.sheets_registro import siguiente_numero, registrar as sheets_registrar
from utils.output_saver import guardar_en_red, estado_disco as disco_estado
from utils import compartidos

BASE_DIR = Path(__file__).parent
TEMPLATES_DIR = BASE_DIR / "templates"
TARIFAS_DIR = BASE_DIR / "tarifas"
# Salidas y uploads fuera de Google Drive para evitar WinError 2 con rutas largas/acentos
SALIDAS_DIR = Path(tempfile.gettempdir()) / "acometidas_salidas"
UPLOADS_DIR = Path(tempfile.gettempdir()) / "acometidas_uploads_ia"
# VIDEOS_DIR: a diferencia de UPLOADS_DIR (frames temporales, se pueden
# perder), los videos originales hay que conservarlos para poder
# reproducirlos despues desde el enlace publico del cliente. Se guardan bajo
# BASE_DIR/config para que, SI se monta un volumen persistente de Railway en
# /app/config (recomendado, igual que en PresuPro), sobrevivan a los
# redeploys; si no se monta ningun volumen aqui, se perderan igual que ya le
# pasa hoy a SALIDAS_DIR/UPLOADS_DIR en cada reinicio.
VIDEOS_DIR = BASE_DIR / "config" / "videos_ia"
# Igual que VIDEOS_DIR: el JSON que necesita /ver/<token> para renderizar el
# informe debe sobrevivir a los redeploys tanto como el propio video y el
# token (este ultimo ya vive en config/enlaces.db), o el enlace del cliente
# se rompe (404) en el primer reinicio del contenedor aunque el video siga ahi.
INFORMES_COMPARTIDOS_DIR = BASE_DIR / "config" / "informes_compartidos"
TARIFAS_FILE = TARIFAS_DIR / "TARIFAS.xlsx"

SALIDAS_DIR.mkdir(exist_ok=True)
UPLOADS_DIR.mkdir(exist_ok=True)
VIDEOS_DIR.mkdir(parents=True, exist_ok=True)
compartidos.init_db()

TEMPLATE_FILES = {
    "obra_1":              "MODELO_MAESTRO_1OPCION.docx",
    "obra_2":              "MODELO_MAESTRO_2OPCIONES.docx",
    "desatasco":           "MODELO_DESATASCO.docx",
    "cctv_bajante":        "MODELO_CCTV_BAJANTE.docx",
    "inspeccion_zum":      "MODELO_INSPECCION_ZUM.docx",
    "limpieza_aerea":      "MODELO_LIMPIEZA_AEREA.docx",
    "fresador":            "MODELO_FRESADOR.docx",
    "robot_limpieza":      "MODELO_ROBOT_LIMPIEZA.docx",
    "fuga_agua":           "MODELO_FUGA_AGUA.docx",
    "vaciado_fosa":        "MODELO_VACIADO_FOSA.docx",
    # Nuevos modelos
    "informe_desatasco":   "MODELO_INFORME_DESATASCO.docx",
    "bajantes_amianto":    "MODELO_BAJANTES_AMIANTO.docx",
    "certificado_obra":    "MODELO_CERTIFICADO_OBRA.docx",
    "plan_seguridad":      "MODELO_PLAN_SEGURIDAD.docx",
    "contrato_saneamiento": "MODELO_CONTRATO_SANEAMIENTO.docx",
    # Nuevos oficios
    "fontaneria":           "MODELO_FONTANERIA.docx",
    "albanileria":          "MODELO_ALBANILERIA.docx",
    # Contrato subcontratacion
    "contrato_subcontrata": "MODELO_CONTRATO_SUBCONTRATA.docx",
}

ALLOWED_IA_EXTENSIONS = {
    ".jpg", ".jpeg", ".png", ".gif", ".webp",
    ".pdf", ".mp4", ".avi", ".mov", ".mkv", ".webm", ".m4v",
    ".doc", ".docx", ".txt", ".md",
}

app = Flask(__name__, template_folder="web_templates")
# Railway (o cualquier PaaS con proxy delante) termina el HTTPS antes de que
# la peticion llegue a gunicorn: sin esto, los enlaces publicos generados con
# url_for(..., _external=True) (el enlace de video para el cliente) saldrian
# como http:// en vez de https://.
from werkzeug.middleware.proxy_fix import ProxyFix
app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)
_secret_key = os.environ.get("SECRET_KEY")
if not _secret_key:
    logging.warning(
        "SECRET_KEY no esta configurada: se usa una clave aleatoria temporal "
        "(las sesiones/enlaces firmados se invalidarian al reiniciar el "
        "servidor). Configura SECRET_KEY en las variables de entorno."
    )
    _secret_key = secrets.token_hex(32)
app.secret_key = _secret_key
app.config["MAX_CONTENT_LENGTH"] = 2 * 1024 * 1024 * 1024  # 2 GB

# El proxy de Railway cierra la conexion si el servidor no manda ningun byte
# de vuelta durante 5 minutos (ver docs.railway.com/networking/public-networking/
# specs-and-limits), y el analisis de varios videos (transcodificar + IA) puede
# tardar mas que eso facilmente. Por eso /api/analizar ya no procesa en la
# propia peticion: guarda los archivos, lanza el trabajo pesado en un hilo en
# segundo plano y responde al instante con un _job_id; el frontend pregunta el
# resultado con /api/analizar_status/<job_id> (polling), evitando el timeout.
_ANALYSIS_JOBS: dict[str, dict] = {}
_ANALYSIS_JOBS_LOCK = threading.Lock()


def _limpiar_jobs_antiguos(max_edad_seg: int = 3 * 3600) -> None:
    corte = time.time() - max_edad_seg
    with _ANALYSIS_JOBS_LOCK:
        for jid in [j for j, v in _ANALYSIS_JOBS.items() if v.get("_ts", 0) < corte]:
            del _ANALYSIS_JOBS[jid]


@app.errorhandler(413)
def _too_large(_e):
    limite_mb = app.config["MAX_CONTENT_LENGTH"] // (1024 * 1024)
    return jsonify({
        "error": f"Archivos demasiado grandes. Limite total: {limite_mb} MB. "
                 "Sube menos videos a la vez o comprimelos."
    }), 413

_tarifas: TarifaLoader | None = None


def get_tarifas() -> TarifaLoader:
    global _tarifas
    if _tarifas is None:
        _tarifas = TarifaLoader(TARIFAS_FILE)
    return _tarifas


def _safe_filename(s: str, maxlen: int = 60) -> str:
    """Remove accents, keep only safe ASCII chars for Windows paths."""
    # Normalize to NFD and drop combining marks (removes accents)
    nfd = unicodedata.normalize("NFD", s)
    ascii_str = "".join(c for c in nfd if unicodedata.category(c) != "Mn")
    # Replace characters not allowed in Windows filenames
    safe = ""
    for c in ascii_str:
        if c.isalnum() or c in (" ", "-", "_", ".", "(", ")"):
            safe += c
        else:
            safe += "-"
    return safe[:maxlen].strip(". -")


def fmt_euro(amount: float) -> str:
    s = f"{amount:,.2f} €"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")


def _transcodificar_video_web(origen: Path, destino: Path) -> bool:
    """Convierte un video CCTV a H.264/MP4 web-compatible y comprimido.

    Las camaras de inspeccion CCTV suelen grabar en codecs antiguos (mpeg4/DivX,
    resolucion CIF) que ningun navegador sabe decodificar en un <video>: el
    archivo carga (headers/tamano correctos) pero no se ve ninguna imagen.
    H.264 es soportado por todos los navegadores y ademas comprime bastante
    mas que esos codecs antiguos.
    """
    import subprocess
    from core.ai_analyst import _find_ffmpeg
    try:
        ffmpeg = _find_ffmpeg()
    except Exception:
        return False
    try:
        subprocess.run(
            [ffmpeg, "-y", "-i", str(origen),
             "-c:v", "libx264", "-profile:v", "main", "-preset", "veryfast", "-crf", "26",
             "-vf", "scale='min(1280,iw)':-2",
             "-c:a", "aac", "-b:a", "96k",
             "-movflags", "+faststart",
             str(destino)],
            capture_output=True, timeout=1800, check=True,
        )
    except Exception:
        return False
    return destino.exists() and destino.stat().st_size > 0


# Intervalo de extraccion de fotogramas en core/ai_analyst.py (_extract_video_frames
# usa "fps=1/5"): 1 fotograma cada 5 segundos, empezando en el segundo 0 del video.
_SEGUNDOS_POR_FOTOGRAMA = 5.0


def _marcar_timestamps_video(wc_report: dict, session_id: str, videos_persistidos: dict[str, str]) -> list[dict]:
    """Anade '_video' (nombre de archivo WEB, ya transcodificado) y '_t' (segundos
    aprox.) a cada observacion de wc_report que referencia un fotograma de un
    video persistido, usando '_evidencia_src' (archivo ORIGINAL de cada
    fotograma, en orden) para calcular la posicion local del fotograma dentro
    de su video.

    videos_persistidos: {nombre_original: nombre_web}.
    Devuelve la lista de videos reproducibles: [{"nombre", "url"}, ...].
    """
    evidencia_src = wc_report.get("_evidencia_src") or []
    if not evidencia_src or not videos_persistidos:
        return []

    contador: dict[str, int] = {}
    local_idx_por_global: list[tuple[str, int] | None] = []
    for src in evidencia_src:
        if src in videos_persistidos:
            idx = contador.get(src, 0)
            local_idx_por_global.append((src, idx))
            contador[src] = idx + 1
        else:
            local_idx_por_global.append(None)  # imagen/PDF, no es un video

    def _anotar(obs: dict) -> None:
        for n in obs.get("imagenes") or []:
            i = int(n) - 1  # "imagenes" es 1-indexado
            if 0 <= i < len(local_idx_por_global) and local_idx_por_global[i]:
                src_original, local_idx = local_idx_por_global[i]
                obs["_video"] = videos_persistidos[src_original]
                obs["_t"] = round(local_idx * _SEGUNDOS_POR_FOTOGRAMA, 1)
                return

    for seccion in wc_report.get("secciones") or []:
        for obs in seccion.get("observaciones_tabla") or []:
            _anotar(obs)

    return [
        {"nombre": web_name, "url": url_for("servir_video_ia", session_id=session_id, filename=web_name)}
        for web_name in videos_persistidos.values()
    ]


def fmt_euro_plain(amount: float) -> str:
    s = f"{amount:,.2f}"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")


def numero_a_letras(n: float) -> str:
    """Convierte un numero entero (parte entera de euros) a palabras en castellano."""
    n = int(round(n))
    if n == 0:
        return "CERO"
    if n < 0:
        return "MENOS " + numero_a_letras(-n)

    unidades = ["", "UN", "DOS", "TRES", "CUATRO", "CINCO", "SEIS", "SIETE",
                "OCHO", "NUEVE", "DIEZ", "ONCE", "DOCE", "TRECE", "CATORCE",
                "QUINCE", "DIECISEIS", "DIECISIETE", "DIECIOCHO", "DIECINUEVE",
                "VEINTE", "VEINTIUN", "VEINTIDOS", "VEINTITRES", "VEINTICUATRO",
                "VEINTICINCO", "VEINTISEIS", "VEINTISIETE", "VEINTIOCHO", "VEINTINUEVE"]
    decenas = ["", "DIEZ", "VEINTE", "TREINTA", "CUARENTA", "CINCUENTA",
               "SESENTA", "SETENTA", "OCHENTA", "NOVENTA"]
    centenas = ["", "CIENTO", "DOSCIENTOS", "TRESCIENTOS", "CUATROCIENTOS",
                "QUINIENTOS", "SEISCIENTOS", "SETECIENTOS", "OCHOCIENTOS", "NOVECIENTOS"]

    def _menos_mil(n: int) -> str:
        if n == 0:
            return ""
        if n == 100:
            return "CIEN"
        if n < 30:
            return unidades[n]
        if n < 100:
            d, u = divmod(n, 10)
            return decenas[d] + (" Y " + unidades[u] if u else "")
        c, resto = divmod(n, 100)
        return centenas[c] + (" " + _menos_mil(resto) if resto else "")

    if n < 1000:
        return _menos_mil(n)
    if n < 2000:
        resto = n - 1000
        return "MIL" + (" " + _menos_mil(resto) if resto else "")
    if n < 1_000_000:
        miles, resto = divmod(n, 1000)
        return _menos_mil(miles) + " MIL" + (" " + _menos_mil(resto) if resto else "")
    if n < 2_000_000:
        resto = n - 1_000_000
        return "UN MILLON" + (" " + numero_a_letras(resto) if resto else "")
    millones, resto = divmod(n, 1_000_000)
    return _menos_mil(millones) + " MILLONES" + (" " + numero_a_letras(resto) if resto else "")


def _parse_fecha(fecha_raw: str):
    try:
        dt = datetime.strptime(fecha_raw, "%Y-%m-%d")
        meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio",
                 "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
        fecha_larga = f"{dt.day} de {meses[dt.month - 1]} de {dt.year}"
        fecha_display = dt.strftime("%d/%m/%Y")
    except ValueError:
        fecha_larga = fecha_raw
        fecha_display = fecha_raw
    return fecha_larga, fecha_display


def _fval(f, name: str, default: float = 0.0) -> float:
    try:
        return float(f.get(name, str(default)).replace(",", "."))
    except ValueError:
        return default


def _parse_items(f, prefix: str = "item") -> list[dict]:
    items = []
    idx = 0
    while True:
        desc = f.get(f"{prefix}_desc_{idx}", "").strip()
        if not desc:
            break
        qty = _fval(f, f"{prefix}_qty_{idx}", 1.0)
        price = _fval(f, f"{prefix}_price_{idx}", 0.0)
        unit = f.get(f"{prefix}_unit_{idx}", "PA").strip() or "PA"
        items.append({
            "descripcion": desc,
            "unidad": unit,
            "cantidad": qty,
            "precio_unitario": price,
            "importe": round(qty * price, 2),
        })
        idx += 1
    return items


def _svc_line(f, prefix: str) -> tuple[dict, float]:
    uds = _fval(f, f"{prefix.lower()}_uds", 1.0)
    precio = _fval(f, f"{prefix.lower()}_precio", 0.0)
    importe = round(uds * precio, 2)
    data = {
        f"[[{prefix}_UDS]]":     _fmt_qty_str(uds),
        f"[[{prefix}_PRECIO]]":  fmt_euro_plain(precio),
        f"[[{prefix}_IMPORTE]]": fmt_euro_plain(importe),
    }
    return data, importe


def _fmt_qty_str(qty: float) -> str:
    if qty == int(qty):
        return str(int(qty))
    return f"{qty:.2f}".replace(".", ",")


def _build_common_data(f) -> tuple[dict, str, str, str, str]:
    fecha_raw = f.get("fecha_contrato", datetime.now().strftime("%Y-%m-%d"))
    fecha_larga, fecha_display = _parse_fecha(fecha_raw)
    num_contrato = f.get("num_contrato", "").strip()
    obra = f.get("obra", "").strip()
    servicio = f.get("servicio", "").strip()
    admin_val = f.get("administracion", "").strip()
    current_year = str(datetime.now().year)
    data = {
        "[[CONTRATO_NUM]]":          num_contrato,
        "[[NUMERO_PRESUPUESTO]]":    num_contrato,
        "[[INFORME_NUM]]":           num_contrato,
        "[[FECHA_CONTRATO]]":        fecha_display,
        "[[FECHA_INFORME]]":         fecha_display,
        "[[FECHA_LARGA]]":           fecha_larga,
        "[[OBRA_COMUNIDAD]]":        obra,
        "[[NOMBRE_OBRA]]":           obra,
        "[[DIRECCION_OBRA]]":        obra,
        "[[SERVICIO_COMUNIDAD]]":    obra,   # siempre la direccion de ejecucion (el titulo propio va en el banner superior)
        "[[TIPO_SERVICIO]]":         servicio,
        "[[CLIENTE_NOMBRE]]":        f.get("cliente_nombre", "").strip(),
        "[[CLIENTE_DIRECCION]]":     f.get("cliente_dir", "").strip(),
        "[[CLIENTE_TELEFONO]]":      f.get("cliente_tel", "").strip(),
        "[[CLIENTE_EMAIL]]":         f.get("cliente_email", "").strip(),
        "[[CLIENTE_CORREO ELECTRONICO]]": f.get("cliente_email", "").strip(),
        "[[CLIENTE_CORREOELECTRONICO]]":  f.get("cliente_email", "").strip(),
        "[[PROVINCIA]]":             f.get("provincia", "Madrid").strip(),
        "[[ADMINISTRACION]]":        admin_val,
        "[[ADMINISTRACION_LABEL]]":  "ADMINISTRACION:" if admin_val else "",
        "[[ADMINISTRACION_PROVINCIA]]": f.get("provincia", "Madrid").strip() if admin_val else "",
        "[[ADMINISTRACION_TELEFONO]]":           f.get("admin_tel", "").strip(),
        "[[ADMINISTRACION_CORREO ELECTRONICO]]": f.get("admin_email", "").strip(),
        "[[ADMINISTRACION_CORREOELECTRONICO]]":  f.get("admin_email", "").strip(),
        # Correccion de años hardcodeados en plantillas antiguas
        "de 2024": f"de {current_year}",
        "de 2025": f"de {current_year}",
    }
    return data, num_contrato, fecha_larga, obra, servicio


def _inject_banner_title(docx_path, titulo: str) -> None:
    """Inserta un banner con el titulo (en MAYUSCULAS) al inicio del encabezado de pagina.

    Solo se usa cuando el flujo aporta 'titulo_servicio' (Presupuesto Integral Multioficio).
    No modifica la plantilla: se aplica sobre el DOCX ya generado, asi que no afecta al
    resto de tipos de presupuesto.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(str(docx_path))
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    if header.paragraphs:
        p = header.paragraphs[0].insert_paragraph_before()
    else:
        p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(titulo.upper())
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    doc.save(str(docx_path))


def _build_obra_data(f, total_sin_iva: float, fecha_larga: str) -> dict:
    return {
        "[[INFORME_TECNICO]]":         f.get("informe", "").strip(),
        "[[SOLUCION_ADOPTAR]]":        f.get("solucion", "").strip(),
        "[[MEMORIA_TECNICA]]":         f.get("memoria", "").strip(),
        "[[TOTAL_PRESUPUESTO]]":       fmt_euro_plain(total_sin_iva),
        "[[RESUMEN_VALORACION]]":      fmt_euro_plain(total_sin_iva),
        "[[FORMA_PAGO]]":              f.get("forma_pago", "50% inicio, 50% finalizacion").strip(),
        "[[PLAZO_EJECUCION]]":         f.get("plazo", "30").strip(),
        "[[FECHA_INICIO_OBRA_LARGA]]": fecha_larga,
        "[[FECHA_FIN_OBRA]]":          f.get("fecha_fin_obra", "").strip(),
    }


def _build_obra_2_extra(f, total_a: float, total_b: float) -> dict:
    return {
        "[[TIPO_OPCION]]":            f.get("tipo_opcion", "").strip(),
        "[[MEMORIA_TRADICIONAL]]":    f.get("memoria_tradicional", "").strip(),
        "[[MEMORIA_MULTILINER]]":     f.get("memoria_multiliner", "").strip(),
        "[[IMPORTE_TOTAL_OPCION_A]]": fmt_euro_plain(total_a),
        "[[IMPORTE_TOTAL_OPCION_B]]": fmt_euro_plain(total_b),
        "[[FIRMANTE_NOMBRE]]":        f.get("firmante_nombre", "").strip(),
        "[[FIRMANTE_CARGO]]":         f.get("firmante_cargo", "").strip(),
        "[[FIRMANTE_DNI]]":           f.get("firmante_dni", "").strip(),
        "[[FIRMANTE_FECHA]]":         f.get("firmante_fecha", "").strip(),
    }


def _build_service_data(f, tipo: str, base_data: dict) -> tuple[dict, float]:
    data = dict(base_data)
    total = 0.0

    if tipo == "desatasco":
        for prefix in ["CAMION", "CAMION_DESP", "TAPA", "INODORO", "CATA"]:
            d, imp = _svc_line(f, prefix)
            data.update(d)
            total += imp

    elif tipo == "cctv_bajante":
        for prefix in ["CCTV", "CCTV_DESP", "TAPA", "INODORO", "CATA"]:
            d, imp = _svc_line(f, prefix)
            data.update(d)
            total += imp

    elif tipo == "inspeccion_zum":
        for prefix in ["CCTV", "CCTV_DESP"]:
            d, imp = _svc_line(f, prefix)
            data.update(d)
            total += imp

    elif tipo == "limpieza_aerea":
        for prefix in ["CAMION", "CAMION_DESP", "OCUPACION", "MEDIOS", "TAPA"]:
            d, imp = _svc_line(f, prefix)
            data.update(d)
            total += imp
        data["[[HORAS_ESTIMADAS]]"] = f.get("horas_estimadas", "").strip()
        data["[[TIEMPO_ESTIMADO]]"] = f.get("tiempo_estimado", "").strip()

    elif tipo == "fresador":
        for prefix in ["FRESADOR", "FRESADOR_DESP", "CAMION", "CAMION_DESP", "MEDIOS"]:
            d, imp = _svc_line(f, prefix)
            data.update(d)
            total += imp
        data["[[HORAS_ESTIMADAS]]"] = f.get("horas_estimadas", "").strip()

    elif tipo == "robot_limpieza":
        for prefix in ["CCTV", "CCTV_DESP", "CAMION", "CAMION_DESP",
                        "OCUPACION", "MEDIOS", "LOCALIZACION"]:
            d, imp = _svc_line(f, prefix)
            data.update(d)
            total += imp
        data["[[HORAS_ESTIMADAS]]"] = f.get("horas_estimadas", "").strip()
        data["[[TIEMPO_ESTIMADO]]"] = f.get("tiempo_estimado", "").strip()

    elif tipo == "fuga_agua":
        data["[[LOCALIZACION_SERVICIO]]"] = f.get("localizacion_servicio", "").strip()
        p1 = _fval(f, "precio_localizacion", 0.0)
        p2 = _fval(f, "precio_hora_adicional", 0.0)
        p3 = _fval(f, "precio_bombona", 0.0)
        data["[[PRECIO_LOCALIZACION_FUGA]]"] = fmt_euro_plain(p1)
        data["[[PRECIO_HORA_ADICIONAL]]"]    = fmt_euro_plain(p2)
        data["[[PRECIO_BOMBONA_GAS]]"]       = fmt_euro_plain(p3)
        total = p1

    elif tipo == "vaciado_fosa":
        data["[[VALIDEZ_OFERTA]]"] = f.get("validez_oferta", "30 dias").strip()
        for field in ["UD_DESPLAZAMIENTO", "UD_SUCCION", "UD_LIMPIEZA", "UD_DESCARGA", "UD_RESIDUO"]:
            data[f"[[{field}]]"] = f.get(field.lower(), "").strip()

    # --- Nuevos modelos ---

    elif tipo == "informe_desatasco":
        data["[[INFORME_TECNICO]]"] = f.get("informe_tecnico", "").strip()

    elif tipo == "bajantes_amianto":
        ba_lines = [
            "MOVILIZACION", "PLAN_TRABAJO_MCA", "MUESTREO_AMIANTO",
            "SUSTITUCION_BAJANTE", "SUSTITUCION_INJERTO",
            "APERTURA_CIERRE", "GESTION_RESIDUOS",
        ]
        for name in ba_lines:
            key = name.lower()
            ud = _fval(f, f"ba_{key}_ud", 0.0)
            precio = _fval(f, f"ba_{key}_precio", 0.0)
            importe = round(ud * precio, 2)
            data[f"[[UD_{name}]]"]     = _fmt_qty_str(ud)
            data[f"[[PRECIO_{name}]]"] = fmt_euro_plain(precio)
            data[f"[[IMPORTE_{name}]]"] = fmt_euro_plain(importe)
            total += importe
        data["[[PRECIO_MUESTREO_CONTROL]]"]      = fmt_euro_plain(_fval(f, "ba_muestreo_control_precio", 0.0))
        data["[[PRECIO_MUESTREO_FINALIZACION]]"] = fmt_euro_plain(_fval(f, "ba_muestreo_fin_precio", 0.0))
        data["[[VALIDEZ_OFERTA]]"]  = f.get("validez_oferta", "30 dias").strip()
        data["[[IMPORTE_TOTAL]]"]   = fmt_euro_plain(total)
        data["[[FORMA_PAGO]]"]      = f.get("forma_pago_ba", "50% al inicio de los trabajos, 50% a la finalizacion").strip()

    elif tipo == "certificado_obra":
        # CONTRATO_NUM, FECHA_LARGA, OBRA_COMUNIDAD already in base_data
        for field_key, placeholder in [
            ("fecha_inicio_obra", "[[FECHA_INICIO_LARGA]]"),
            ("fecha_fin_obra",    "[[FECHA_FINAL_LARGA]]"),
        ]:
            raw = f.get(field_key, "").strip()
            if raw:
                larga, _ = _parse_fecha(raw)
                data[placeholder] = larga
            else:
                data[placeholder] = ""

    elif tipo == "plan_seguridad":
        ps_fields = [
            "empresa_contratista", "cif_empresa", "direccion_empresa", "telefono_empresa",
            "autor_plan", "dni_autor_plan", "jefe_obra", "coordinador_seguridad",
            "recurso_preventivo", "acceso_obra", "descripcion_tecnica_obra",
            "promotor_nombre", "promotor_cif", "promotor_direccion", "promotor_telefono",
            "numero_max_trabajadores", "presupuesto_ejecucion_material",
            "centro_asistencial", "direccion_centro_medico", "telefono_centro_medico",
            "mes_plan", "plano_centro_medico", "tabla_presupuesto",
        ]
        for field in ps_fields:
            data[f"[[{field.upper()}]]"] = f.get(field, "").strip()
        data["[[PLAZO_EJECUCION]]"] = f.get("plazo", "30").strip()
        # Derivar AÑO_PLAN y FECHA_PLAN de fecha_contrato
        try:
            dt_plan = datetime.strptime(f.get("fecha_contrato", ""), "%Y-%m-%d")
            meses_ps = ["enero","febrero","marzo","abril","mayo","junio",
                        "julio","agosto","septiembre","octubre","noviembre","diciembre"]
            data["[[AÑO_PLAN]]"]  = str(dt_plan.year)
            data["[[FECHA_PLAN]]"] = f"{dt_plan.day} de {meses_ps[dt_plan.month-1]} de {dt_plan.year}"
        except ValueError:
            data["[[AÑO_PLAN]]"]  = str(datetime.now().year)
            data["[[FECHA_PLAN]]"] = ""

    elif tipo == "contrato_saneamiento":
        data["[[IMPORTE_ANUAL]]"] = fmt_euro_plain(_fval(f, "importe_anual", 0.0))
        total = _fval(f, "importe_anual", 0.0)

    data["[[IMPORTE_TOTAL_ESTIMADO]]"] = fmt_euro_plain(total)
    return data, total


# ── Routes ─────────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    today = datetime.now().strftime("%Y-%m-%d")
    missing = [name for name in TEMPLATE_FILES.values()
               if not (TEMPLATES_DIR / name).exists()]
    return render_template("index.html", today=today, missing_templates=missing)


@app.route("/analizar")
def analizar():
    api_key_set = bool(os.environ.get("ANTHROPIC_API_KEY", "").strip())
    return render_template("analizar.html", api_key_set=api_key_set)


@app.route("/croquis")
def croquis():
    api_key_set = bool(os.environ.get("ANTHROPIC_API_KEY", "").strip())
    return render_template("croquis.html", api_key_set=api_key_set)


ALLOWED_PLANO = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".tiff", ".tif", ".pdf"}


def _plano_svg_desde_archivo(src: Path, api_key: str = "", contexto: str = "") -> tuple[str, dict]:
    """Analiza un boceto/croquis (imagen o PDF) y devuelve (svg, estructura).

    Extraido de la vista /api/generar_plano_ia para poder reutilizarlo desde el
    analizador CCTV, que genera el plano tecnico cuando se aporta un croquis.
    Lanza ValueError con un mensaje legible si el formato o la clave fallan.
    """
    import base64
    import re
    import anthropic as _ant

    key = (api_key or "").strip() or os.environ.get("ANTHROPIC_API_KEY", "")
    if not key:
        raise ValueError("API Key no configurada.")

    suffix = src.suffix.lower()
    if suffix not in ALLOWED_PLANO:
        raise ValueError("Formato no soportado. Usa JPG, PNG, WEBP, BMP o PDF.")

    img_b64 = base64.b64encode(src.read_bytes()).decode()
    contexto = (contexto or "").strip()
    ctx_extra = f"\nIndicaciones adicionales: {contexto}" if contexto else ""

    client = _ant.Anthropic(api_key=key)

    # Construir bloque de contenido segun tipo de archivo
    if suffix == ".pdf":
        content_block = {
            "type": "document",
            "source": {"type": "base64", "media_type": "application/pdf", "data": img_b64},
        }
    else:
        mime_map = {".jpg":"image/jpeg",".jpeg":"image/jpeg",".png":"image/png",
                    ".gif":"image/gif",".webp":"image/webp",".bmp":"image/bmp",
                    ".tiff":"image/tiff",".tif":"image/tiff"}
        content_block = {
            "type": "image",
            "source": {"type": "base64",
                       "media_type": mime_map.get(suffix, "image/jpeg"),
                       "data": img_b64},
        }

    # ── PASO 1: extraer estructura como JSON con posiciones en % ──────────
    PROMPT_PASO1 = """Analiza este croquis/plano con mucha atencion. Extrae TODOS los elementos visibles.
Para cada elemento indica su posicion como porcentaje del ancho (x) y alto (y) de la imagen, de 0 a 100.
Sé muy preciso con las posiciones relativas: si un pozo esta a la izquierda de otro, el x del primero debe ser menor.

Devuelve UNICAMENTE este JSON (sin markdown, sin texto adicional):
{
  "tuberias": [{"x1":25,"y1":40,"x2":70,"y2":40,"dn":"DN150","material":"PVC","longitud":"12m","flecha":true}],
  "pozos": [{"x":25,"y":40,"id":"P1","cota_tapa":null,"cota_solera":null}],
  "bajantes": [{"x":50,"y":20,"id":"B1"}],
  "arquetas": [{"x":80,"y":60,"id":"A1"}],
  "etiquetas": [{"x":50,"y":10,"texto":"texto visible en el boceto"}],
  "cotas": [{"x1":20,"y1":80,"x2":60,"y2":80,"valor":"15.00m"}],
  "descripcion": "descripcion breve de la instalacion representada"
}
Omite claves cuyo array este vacio. Si un campo es desconocido usa null."""

    def _api_create(client, **kwargs):
        """Llama a messages.create con reintentos en caso de 529 (overloaded)."""
        delays = [5, 15, 30]
        for attempt, delay in enumerate(delays, 1):
            try:
                return client.messages.create(**kwargs)
            except Exception as exc:
                code = getattr(exc, "status_code", None)
                if code == 529 and attempt < len(delays):
                    time.sleep(delay)
                    continue
                raise

    resp1 = _api_create(
        client,
        model="claude-sonnet-4-6",
        max_tokens=4000,
        messages=[{"role": "user", "content": [content_block, {"type": "text", "text": PROMPT_PASO1}]}],
    )

    raw1 = resp1.content[0].text.strip()
    json_m = re.search(r"\{[\s\S]*\}", raw1)
    estructura = {}
    if json_m:
        try:
            estructura = json.loads(json_m.group(0))
        except Exception:
            estructura = {}

    desc = estructura.get("descripcion", "Red de saneamiento")[:60]
    fecha_hoy = datetime.now().strftime("%d/%m/%Y")

    # ── PASO 2: generar SVG fiel a las coordenadas extraidas ──────────────
    PROMPT_PASO2 = f"""Genera un plano tecnico SVG usando EXACTAMENTE estas coordenadas extraidas del boceto original:

{json.dumps(estructura, ensure_ascii=False, indent=2)}{ctx_extra}

ESCALA obligatoria:
  Area util del plano: x=[30,850], y=[30,560]  (820 px ancho, 530 px alto)
  svg_x = 30 + (campo_x / 100.0) * 820
  svg_y = 30 + (campo_y / 100.0) * 530
  Aplica esta formula a TODOS los campos x, y, x1, y1, x2, y2.

ESPECIFICACION SVG:
viewBox="0 0 900 660" width="900" height="660"
<rect width="900" height="660" fill="#ffffff"/>

<defs>
  <marker id="arr" markerWidth="8" markerHeight="6" refX="7" refY="3" orient="auto">
    <path d="M0,0 L8,3 L0,6 Z" fill="#1a2a3a"/>
  </marker>
</defs>

TUBERIAS: <line x1="..." y1="..." x2="..." y2="..." stroke="#1a2a3a" stroke-width="2.5" marker-end="url(#arr)"/>
  Etiqueta DN junto a punto medio de la linea, font-size="10" fill="#1a2a3a"
  Si flecha=false omite marker-end

POZOS: <g transform="translate(svg_x,svg_y)">
  <circle r="12" fill="white" stroke="#1a2a3a" stroke-width="2"/>
  <circle r="4" fill="#555"/>
  <text text-anchor="middle" dy="-18" font-family="Arial" font-size="10" fill="#1a2a3a">id</text>
</g>

BAJANTES: <g transform="translate(svg_x,svg_y)">
  <rect x="-8" y="-8" width="16" height="16" fill="white" stroke="#1a2a3a" stroke-width="2"/>
  <line x1="-6" y1="-6" x2="6" y2="6" stroke="#1a2a3a" stroke-width="1.5"/>
  <line x1="6" y1="-6" x2="-6" y2="6" stroke="#1a2a3a" stroke-width="1.5"/>
  <text text-anchor="middle" dy="-14" font-family="Arial" font-size="10" fill="#1a2a3a">id</text>
</g>

ARQUETAS: <g transform="translate(svg_x,svg_y)">
  <rect x="-10" y="-10" width="20" height="20" fill="white" stroke="#1a2a3a" stroke-width="2"/>
  <text text-anchor="middle" dy="-16" font-family="Arial" font-size="10" fill="#1a2a3a">id</text>
</g>

ETIQUETAS: <text x="svg_x" y="svg_y" font-family="Arial" font-size="11" fill="#333">texto</text>

COTAS: <line x1="..." y1="..." x2="..." y2="..." stroke="#aaa" stroke-width="0.8"/>
  + marcas perpendiculares en extremos (4px) + <text font-size="9" fill="#666">valor</text>

CAJETIN (x=688, y=555, 200x92):
<rect x="688" y="555" width="200" height="92" fill="white" stroke="#1a2a3a" stroke-width="1.5"/>
<line x1="688" y1="580" x2="888" y2="580" stroke="#1a2a3a" stroke-width="0.8"/>
<line x1="688" y1="600" x2="888" y2="600" stroke="#1a2a3a" stroke-width="0.8"/>
<line x1="688" y1="622" x2="888" y2="622" stroke="#1a2a3a" stroke-width="0.8"/>
<text x="788" y="572" text-anchor="middle" font-family="Arial" font-size="9" font-weight="bold" fill="#1a2a3a">RED DE SANEAMIENTO</text>
<text x="788" y="592" text-anchor="middle" font-family="Arial" font-size="8" fill="#333">{desc}</text>
<text x="700" y="613" font-family="Arial" font-size="8" fill="#555">Fecha: {fecha_hoy}</text>
<text x="820" y="613" font-family="Arial" font-size="8" fill="#555">Esc: S/E</text>
<text x="788" y="638" text-anchor="middle" font-family="Arial" font-size="8" fill="#555">Acometidas Europa S.L.</text>

LEYENDA (x=12, y=555, 160x92): muestra solo simbolos presentes en el plano generado.

DEVUELVE UNICAMENTE el SVG. Empieza con <svg y termina con </svg>. Sin markdown."""

    resp2 = _api_create(
        client,
        model="claude-sonnet-4-6",
        max_tokens=16000,
        messages=[{
            "role": "user",
            "content": [content_block, {"type": "text", "text": PROMPT_PASO2}],
        }],
    )

    raw = resp2.content[0].text.strip()
    svg_match = re.search(r"<svg[\s\S]*?</svg>", raw, re.IGNORECASE)
    if not svg_match:
        raise ValueError("La IA no genero un SVG valido.")

    return svg_match.group(0), estructura


def _svg_a_png(svg_text: str, dpi: int = 170) -> Path:
    """Rasteriza un SVG a PNG con PyMuPDF.

    Se usa PyMuPDF (ya presente para leer PDFs) porque WeasyPrint/cairo no estan
    operativos en este equipo: faltan las librerias GTK y su import falla.
    """
    import fitz

    tmp = Path(tempfile.mkdtemp(prefix="plano_png_"))
    svg_path = tmp / "plano.svg"
    svg_path.write_text(svg_text, encoding="utf-8")
    doc = fitz.open(str(svg_path))
    pix = doc[0].get_pixmap(dpi=dpi)
    png_path = tmp / "plano.png"
    pix.save(str(png_path))
    doc.close()
    return png_path


def _croquis_docx_desde_png(png_path: Path, titulo: str = "Croquis de red",
                            num_ref: str = "", direccion: str = "",
                            fecha_display: str = "", notas: str = "",
                            stem: str = "") -> tuple[Path, Path | None]:
    """Monta el documento del plano (DOCX + PDF) a partir del PNG del plano.

    Devuelve (docx_path, pdf_path|None). El PDF sale por Microsoft Word (COM),
    que es la via que funciona en este equipo.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT

    AZUL = RGBColor(0x1F, 0x4E, 0x79)
    GRIS = RGBColor(0x6B, 0x72, 0x80)

    doc = Document()
    sec = doc.sections[0]
    # Apaisado: el plano es mucho mas ancho que alto
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = sec.bottom_margin = Cm(1.5)
    sec.left_margin = sec.right_margin = Cm(1.8)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("ACOMETIDAS EUROPA SANEAMIENTO TECNICO S.L.")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = AZUL

    t = doc.add_paragraph()
    t.paragraph_format.space_after = Pt(6)
    tr = t.add_run((titulo or "Croquis de red").upper())
    tr.bold = True
    tr.font.size = Pt(15)
    tr.font.color.rgb = AZUL

    meta = "   |   ".join(x for x in (
        (f"Ref: {num_ref}" if num_ref else ""),
        direccion, fecha_display,
    ) if x)
    if meta:
        m = doc.add_paragraph()
        m.paragraph_format.space_after = Pt(10)
        mr = m.add_run(meta)
        mr.font.size = Pt(9)
        mr.font.color.rgb = GRIS

    # El plano debe caber en UNA pagina: se escala por ancho y por alto util
    # (si solo se fija el ancho, un plano poco apaisado desborda a la pagina 2).
    ancho_disp = sec.page_width - sec.left_margin - sec.right_margin
    alto_disp = sec.page_height - sec.top_margin - sec.bottom_margin - Cm(4.2)
    try:
        from PIL import Image as _Img
        with _Img.open(png_path) as _im:
            ratio = _im.height / _im.width
    except Exception:
        ratio = 660 / 900
    ancho = ancho_disp
    if int(ancho * ratio) > int(alto_disp):
        ancho = int(alto_disp / ratio)

    ip = doc.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.add_run().add_picture(str(png_path), width=ancho)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run("Plano tecnico generado con IA a partir del croquis aportado. "
                     "Representacion esquematica sin escala.")
    cr.italic = True
    cr.font.size = Pt(8)
    cr.font.color.rgb = GRIS

    if notas:
        np_ = doc.add_paragraph()
        np_.paragraph_format.space_before = Pt(10)
        nr1 = np_.add_run("Observaciones: ")
        nr1.bold = True
        nr1.font.size = Pt(9)
        nr1.font.color.rgb = AZUL
        nr2 = np_.add_run(notas)
        nr2.font.size = Pt(9)

    stem = stem or _safe_filename(f"Croquis_{num_ref or titulo}", maxlen=50)
    docx_path = SALIDAS_DIR / f"{stem}.docx"
    doc.save(str(docx_path))

    pdf_path = None
    try:
        pdf_path = PdfConverter().convert(docx_path, SALIDAS_DIR)
    except Exception:
        pdf_path = None
    return docx_path, pdf_path


@app.route("/api/generar_plano_ia", methods=["POST"])
def api_generar_plano_ia():
    """Analiza un boceto/croquis con IA y genera un plano tecnico SVG limpio."""
    import traceback as _tb
    tmp_dir = None
    try:
        img_file = request.files.get("imagen")
        if not img_file or not img_file.filename:
            return jsonify({"error": "No se recibio ninguna imagen."}), 400

        tmp_dir = Path(tempfile.mkdtemp(prefix="plano_ia_"))
        dest = tmp_dir / f"boceto{Path(img_file.filename).suffix.lower()}"
        img_file.save(str(dest))

        svg, estructura = _plano_svg_desde_archivo(
            dest,
            api_key=request.form.get("api_key", ""),
            contexto=request.form.get("contexto", ""),
        )
        return jsonify({"svg": svg, "estructura": estructura})
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": str(e), "traceback": _tb.format_exc()}), 500
    finally:
        if tmp_dir and tmp_dir.exists():
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)


@app.route("/api/generar_croquis", methods=["POST"])
def api_generar_croquis():
    """Genera un PDF profesional a partir de una imagen de croquis/plano."""
    import base64
    import traceback as _tb
    tmp_dir = None
    try:
        img_file  = request.files.get("imagen")
        svg_data  = request.form.get("svg_data", "").strip()   # plano IA generado
        titulo    = request.form.get("titulo", "Croquis de Red").strip()
        num_ref   = request.form.get("num_ref", "").strip()
        direccion = request.form.get("direccion", "").strip()
        fecha_raw = request.form.get("fecha", datetime.now().strftime("%Y-%m-%d")).strip()
        notas     = request.form.get("notas", "").strip()

        if not svg_data and (not img_file or not img_file.filename):
            return jsonify({"error": "No se recibio ninguna imagen."}), 400

        # PDF sin SVG: no se puede incrustar directamente, pedir que use IA primero
        if not svg_data and img_file and Path(img_file.filename).suffix.lower() == ".pdf":
            return jsonify({"error": "Para generar el PDF desde un documento PDF, usa primero 'Generar plano con IA' y activa el plano generado."}), 400

        tmp_dir = Path(tempfile.mkdtemp(prefix="croquis_"))

        # Fecha en formato largo
        try:
            dt = datetime.strptime(fecha_raw, "%Y-%m-%d")
            meses = ["enero","febrero","marzo","abril","mayo","junio",
                     "julio","agosto","septiembre","octubre","noviembre","diciembre"]
            fecha_display = f"{dt.day} de {meses[dt.month-1]} de {dt.year}"
        except ValueError:
            fecha_display = fecha_raw

        # Contenido de imagen: SVG inline o foto base64
        if svg_data:
            img_content = f'<div style="width:100%">{svg_data}</div>'
        else:
            suffix = Path(img_file.filename).suffix.lower()
            ALLOWED = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".tiff", ".tif"}
            if suffix not in ALLOWED:
                return jsonify({"error": "Formato no soportado. Usa JPG, PNG, WEBP o BMP."}), 400
            img_path = tmp_dir / f"imagen{suffix}"
            img_file.save(str(img_path))
            mime = {".jpg":"image/jpeg",".jpeg":"image/jpeg",".png":"image/png",
                    ".gif":"image/gif",".webp":"image/webp",".bmp":"image/bmp",
                    ".tiff":"image/tiff",".tif":"image/tiff"}.get(suffix, "image/jpeg")
            img_b64 = base64.b64encode(img_path.read_bytes()).decode()
            img_content = f'<img src="data:{mime};base64,{img_b64}" alt="{titulo}" style="max-width:100%;max-height:480px;object-fit:contain">'

        # HTML del documento
        notas_html = ""
        if notas:
            notas_html = f"""
            <div class="notas-box">
              <div class="notas-title">Observaciones</div>
              <div class="notas-text">{notas}</div>
            </div>"""

        meta_rows = ""
        if num_ref:
            meta_rows += f'<tr><td class="meta-k">Referencia</td><td class="meta-v">{num_ref}</td></tr>'
        if direccion:
            meta_rows += f'<tr><td class="meta-k">Ubicacion</td><td class="meta-v">{direccion}</td></tr>'
        meta_rows += f'<tr><td class="meta-k">Fecha</td><td class="meta-v">{fecha_display}</td></tr>'

        html = f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<style>
  @page {{ margin: 1.8cm 1.8cm 2cm; }}
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: 'Arial', sans-serif; color: #1a2a3a; background: white; }}

  .header {{
    background: linear-gradient(135deg, #1a3a5c 0%, #2e6da4 100%);
    color: white; padding: 18px 24px; border-radius: 8px 8px 0 0;
    display: flex; align-items: center; justify-content: space-between;
    margin-bottom: 0;
  }}
  .header-brand {{ font-size: 13px; font-weight: 700; letter-spacing: 0.05em; }}
  .header-sub {{ font-size: 9px; opacity: 0.8; margin-top: 3px; }}
  .header-doc {{ text-align: right; font-size: 9px; opacity: 0.75; }}

  .title-bar {{
    background: #f0f6ff; border: 1px solid #c8daf0; border-top: none;
    padding: 14px 24px; border-radius: 0 0 4px 4px; margin-bottom: 14px;
  }}
  .doc-title {{ font-size: 16px; font-weight: 700; color: #1a3a5c; text-transform: uppercase; }}
  .doc-subtitle {{ font-size: 10px; color: #5a7a9a; margin-top: 3px; }}

  .meta-table {{ width: 100%; border-collapse: collapse; margin-bottom: 16px; font-size: 10px; }}
  .meta-k {{ background: #e8f0f8; padding: 6px 12px; font-weight: 700;
             color: #2e6da4; width: 120px; border: 1px solid #c8daf0; }}
  .meta-v {{ padding: 6px 12px; border: 1px solid #c8daf0; color: #1a2a3a; }}

  .img-frame {{
    border: 2px solid #c8daf0; border-radius: 6px; padding: 12px;
    background: #f8fbff; text-align: center; margin-bottom: 14px;
  }}
  .img-frame img {{ max-width: 100%; max-height: 480px; object-fit: contain; }}
  .img-caption {{ font-size: 9px; color: #7a9ab8; margin-top: 8px; font-style: italic; }}

  .notas-box {{
    border-left: 4px solid #2e6da4; background: #f0f6ff;
    padding: 10px 14px; border-radius: 0 6px 6px 0; margin-bottom: 14px;
  }}
  .notas-title {{ font-size: 10px; font-weight: 700; color: #2e6da4; margin-bottom: 4px; }}
  .notas-text {{ font-size: 10px; color: #1a2a3a; line-height: 1.5; }}

  .footer {{
    position: fixed; bottom: 0; left: 1.8cm; right: 1.8cm;
    border-top: 1px solid #c8daf0; padding-top: 6px;
    display: flex; justify-content: space-between; align-items: center;
    font-size: 8px; color: #8a9ab8;
  }}
</style>
</head>
<body>

<div class="header">
  <div>
    <div class="header-brand">ACOMETIDAS EUROPA S.L.</div>
    <div class="header-sub">Saneamiento Tecnico &bull; Poceria &bull; CCTV</div>
  </div>
  <div class="header-doc">
    {'<div>Ref: ' + num_ref + '</div>' if num_ref else ''}
    <div>{fecha_display}</div>
  </div>
</div>

<div class="title-bar">
  <div class="doc-title">{titulo}</div>
  {'<div class="doc-subtitle">' + direccion + '</div>' if direccion else ''}
</div>

<table class="meta-table">
  {meta_rows}
</table>

<div class="img-frame">
  {img_content}
  <div class="img-caption">{titulo}{' &mdash; ' + direccion if direccion else ''}</div>
</div>

{notas_html}

<div class="footer">
  <span>Acometidas Europa S.L. &bull; Documento tecnico de uso interno</span>
  <span>{fecha_display}</span>
</div>

</body>
</html>"""

        html_path = tmp_dir / "croquis.html"
        html_path.write_text(html, encoding="utf-8")

        stem = _safe_filename(f"Croquis_{num_ref or titulo}", maxlen=50)
        pdf_path = SALIDAS_DIR / f"{stem}.pdf"

        # Via preferente: WeasyPrint (HTML -> PDF). En equipos sin las librerias
        # GTK su import falla, asi que se cae al camino DOCX -> PDF por Word.
        try:
            import weasyprint
            weasyprint.HTML(filename=str(html_path)).write_pdf(str(pdf_path))
        except Exception as e_weasy:
            if svg_data:
                png_path = _svg_a_png(svg_data)
            else:
                png_path = img_path
            cx_docx, cx_pdf = _croquis_docx_desde_png(
                png_path, titulo=titulo, num_ref=num_ref, direccion=direccion,
                fecha_display=fecha_display, notas=notas, stem=stem)
            if not cx_pdf:
                return jsonify({"error": f"No se pudo generar el PDF ({e_weasy}). "
                                         f"Se ha guardado el documento Word: {cx_docx.name}",
                                "docx": cx_docx.name}), 500
            return jsonify({"pdf": cx_pdf.name, "docx": cx_docx.name})

        if not pdf_path.exists():
            return jsonify({"error": "No se pudo generar el PDF."}), 500

        return jsonify({"pdf": pdf_path.name})

    except Exception as e:
        return jsonify({"error": str(e), "traceback": _tb.format_exc()}), 500
    finally:
        if tmp_dir and tmp_dir.exists():
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)


@app.route("/api/analizar", methods=["POST"])
def api_analizar():
    import traceback as _tb_mod
    _log = Path(tempfile.gettempdir()) / "debug_analizar.log"

    def _write_log(msg: str):
        print(f"[analizar] {msg}", flush=True)
        try:
            with open(_log, "a", encoding="utf-8") as _f:
                _f.write(msg + "\n")
        except Exception:
            pass

    _write_log("=== NUEVA PETICION /api/analizar ===")

    try:
        # Comprobacion temprana de que el modulo carga bien (el import real
        # que se usa esta dentro de _procesar_analisis_bg, en segundo plano).
        import core.ai_analyst  # noqa: F401
    except Exception as e:
        tb = _tb_mod.format_exc()
        _write_log("ERROR IMPORTANDO ai_analyst:\n" + tb)
        return jsonify({"error": str(e), "traceback": tb}), 500

    try:
        tipo = request.form.get("tipo_analisis", "general")
        formato = request.form.get("formato", "descriptivo")  # descriptivo | wincam | ambos
        context = request.form.get("contexto", "")
        api_key = request.form.get("api_key", "").strip()
        num_ref = request.form.get("num_ref", "").strip()
        cliente = request.form.get("cliente", "").strip()
        proyecto = request.form.get("proyecto", num_ref).strip()
        calle = request.form.get("calle", "").strip()
        poblacion = request.form.get("poblacion", "").strip()

        # Nombre de archivo de los informes: "N.o - INFORME NOMBRE DE LA CALLE N.o DE LA CALLE"
        _direccion_informe = ", ".join(x for x in (calle, poblacion) if x)
        _label_informe = _safe_filename(_direccion_informe, maxlen=60) if _direccion_informe else ""
        if num_ref and _label_informe:
            _informe_stem = f"{_safe_filename(num_ref, maxlen=20)} - INFORME {_label_informe}"
        elif num_ref:
            _informe_stem = f"{_safe_filename(num_ref, maxlen=20)} - INFORME"
        elif _label_informe:
            _informe_stem = f"INFORME {_label_informe}"
        else:
            _informe_stem = ""

        _write_log(f"tipo={tipo} formato={formato} api_key={'SET' if api_key else 'ENV'}")

        # Save uploaded files
        session_id = uuid.uuid4().hex[:8]
        session_dir = UPLOADS_DIR / session_id
        session_dir.mkdir(parents=True, exist_ok=True)
        _write_log(f"session_dir={session_dir}")

        saved_files: list[Path] = []
        for uploaded in request.files.getlist("archivos"):
            if not uploaded.filename:
                continue
            suffix = Path(uploaded.filename).suffix.lower()

            # ZIP: extraer y anadir los archivos validos que contenga
            if suffix == ".zip":
                import zipfile, shutil as _sh
                zpath = session_dir / f"{uuid.uuid4().hex[:6]}.zip"
                uploaded.save(str(zpath))
                try:
                    with zipfile.ZipFile(zpath) as zf:
                        for member in zf.namelist():
                            if member.endswith("/"):
                                continue
                            msuf = Path(member).suffix.lower()
                            if msuf not in ALLOWED_IA_EXTENSIONS:
                                continue
                            # nombre aleatorio (ignora la ruta interna -> evita zip-slip)
                            dest = session_dir / f"{uuid.uuid4().hex[:6]}{msuf}"
                            with zf.open(member) as src, open(dest, "wb") as out:
                                _sh.copyfileobj(src, out)
                            if dest.exists() and dest.stat().st_size > 0:
                                saved_files.append(dest)
                                _write_log(f"ZIP {member} -> {dest} ({dest.stat().st_size} bytes)")
                except Exception as e:
                    _write_log(f"Error extrayendo ZIP {uploaded.filename}: {e}")
                continue

            if suffix not in ALLOWED_IA_EXTENSIONS:
                _write_log(f"Extension no permitida: {suffix}")
                continue
            dest = session_dir / f"{uuid.uuid4().hex[:6]}{suffix}"
            _write_log(f"Guardando {uploaded.filename} -> {dest}")
            uploaded.save(str(dest))
            if dest.exists():
                saved_files.append(dest)
                _write_log(f"OK: {dest.stat().st_size} bytes")
            else:
                _write_log(f"WARN: no existe tras save: {dest}")

        # Croquis (optional - single file: image or PDF)
        croquis_path = None
        croquis_file = request.files.get("croquis")
        if croquis_file and croquis_file.filename:
            suf_c = Path(croquis_file.filename).suffix.lower()
            if suf_c in ALLOWED_IA_EXTENSIONS:
                dest_c = session_dir / f"croquis{suf_c}"
                croquis_file.save(str(dest_c))
                if dest_c.exists():
                    croquis_path = dest_c
                    _write_log(f"Croquis guardado: {dest_c} ({dest_c.stat().st_size} bytes)")

        _write_log(f"saved_files={len(saved_files)} croquis={'si' if croquis_path else 'no'}")
        if not saved_files:
            return jsonify({"error": "No se recibieron archivos validos o no se pudieron guardar."}), 400

        with _ANALYSIS_JOBS_LOCK:
            _ANALYSIS_JOBS[session_id] = {"status": "processing", "_ts": time.time()}
        _limpiar_jobs_antiguos()

        threading.Thread(
            target=_procesar_analisis_bg,
            kwargs=dict(
                session_id=session_id, saved_files=saved_files, croquis_path=croquis_path,
                tipo=tipo, formato=formato, context=context, api_key=api_key,
                num_ref=num_ref, cliente=cliente, proyecto=proyecto, calle=calle,
                poblacion=poblacion, informe_stem=_informe_stem, base_url=request.url_root,
            ),
            daemon=True,
        ).start()
        _write_log(f"Trabajo {session_id} lanzado en segundo plano")
        return jsonify({"_job_id": session_id, "_polling": True})

    except Exception as e:
        tb = _tb_mod.format_exc()
        _write_log("=== EXCEPCION ===\n" + tb)
        return jsonify({"error": f"{str(e)}\n\n--- TRACEBACK ---\n{tb}", "traceback": tb}), 500


def _procesar_analisis_bg(session_id, saved_files, croquis_path, tipo, formato, context,
                           api_key, num_ref, cliente, proyecto, calle, poblacion,
                           informe_stem, base_url):
    """Trabajo pesado de /api/analizar (transcodificar videos, llamar a Claude,
    generar DOCX/PDF/croquis), ejecutado en un hilo aparte para no bloquear la
    respuesta HTTP: el proxy de Railway corta la conexion si el servidor no
    manda ningun byte de vuelta durante 5 minutos, y esto puede tardar mas.
    El resultado se deja en _ANALYSIS_JOBS[session_id] para que lo recoja
    /api/analizar_status (polling desde el frontend)."""
    import traceback as _tb_mod
    _log = Path(tempfile.gettempdir()) / "debug_analizar.log"

    def _write_log(msg: str):
        print(f"[analizar:{session_id}] {msg}", flush=True)
        try:
            with open(_log, "a", encoding="utf-8") as _f:
                _f.write(msg + "\n")
        except Exception:
            pass

    with app.test_request_context(base_url=base_url):
        try:
            from core.ai_analyst import (analyze, generate_report_docx,
                                          analyze_wincam, generate_wincam_docx,
                                          attach_wincam_diagramas)

            # Copia persistente de los videos, TRANSCODIFICADOS a H.264/MP4 (ver
            # _transcodificar_video_web): las camaras CCTV suelen grabar en codecs
            # que ningun navegador reproduce, y de paso queda una version
            # comprimida lista para compartir con el cliente. videos_persistidos
            # mapea nombre_original (el que usa "_evidencia_src" del informe) ->
            # nombre_web (el .mp4 ya transcodificado y realmente servido).
            video_ext_persist = {".mp4", ".avi", ".mov", ".mkv", ".webm", ".m4v"}
            videos_persistidos: dict[str, str] = {}
            video_session_dir = VIDEOS_DIR / session_id
            for sf in saved_files:
                if sf.suffix.lower() in video_ext_persist:
                    video_session_dir.mkdir(parents=True, exist_ok=True)
                    destino_web = video_session_dir / f"{sf.stem}.mp4"
                    if _transcodificar_video_web(sf, destino_web):
                        videos_persistidos[sf.name] = destino_web.name
                        _write_log(f"Video transcodificado: {sf.name} -> {destino_web.name}")
                    else:
                        import shutil as _sh_video
                        destino_fallback = video_session_dir / sf.name
                        try:
                            _sh_video.copy2(sf, destino_fallback)
                            videos_persistidos[sf.name] = destino_fallback.name
                            _write_log(f"Video persistido SIN transcodificar (fallback): {sf.name}")
                        except Exception as e:
                            _write_log(f"No se pudo persistir video {sf.name}: {e}")

            # Enlace publico de solo-video, generado YA (independientemente del
            # formato elegido) para poder incrustarlo como texto+QR en el DOCX que
            # se genere: asi viaja siempre pegado al informe. Se guarda tambien
            # una version minima del informe compartible (por si el formato es
            # solo "descriptivo", sin secciones/observaciones); el bloque WinCam
            # de mas abajo la enriquece si aplica.
            enlace_video = None
            if videos_persistidos:
                try:
                    _token_video = compartidos.crear_o_reusar(session_id)
                    enlace_video = url_for("ver_publico", token=_token_video, _external=True)
                    INFORMES_COMPARTIDOS_DIR.mkdir(parents=True, exist_ok=True)
                    (INFORMES_COMPARTIDOS_DIR / f"Informe_WinCam_{session_id}.json").write_text(
                        json.dumps({
                            "proyecto": proyecto, "num_ref": num_ref, "cliente": cliente,
                            "nivel_urgencia_global": None, "secciones": [], "totales": {},
                            "videos": [
                                {"nombre": web_name, "url": url_for("servir_video_ia", session_id=session_id, filename=web_name)}
                                for web_name in videos_persistidos.values()
                            ],
                        }, ensure_ascii=False),
                        encoding="utf-8",
                    )
                except Exception as e:
                    _write_log(f"No se pudo generar el enlace publico: {e}")

            result = {"_formato": formato, "_session_id": session_id}
            if enlace_video:
                result["_enlace_video"] = enlace_video

            # Formato descriptivo (o ambos)
            if formato in ("descriptivo", "ambos"):
                _write_log("Llamando a analyze() descriptivo...")
                report = analyze(saved_files, tipo, context, api_key, croquis_path=croquis_path)
                _write_log("analyze() OK")
                result.update(report)
                docx_name = f"{informe_stem}.docx" if informe_stem else f"Informe_IA_{session_id}.docx"
                docx_path = SALIDAS_DIR / docx_name
                try:
                    generate_report_docx(report, docx_path, num_ref=num_ref, cliente=cliente,
                                         enlace_video=enlace_video)
                    result["_docx"] = docx_name
                except Exception as e:
                    result["_docx_error"] = str(e)
                    _write_log(f"DOCX descriptivo error: {e}")
                try:
                    pdf_path = PdfConverter().convert(docx_path, SALIDAS_DIR)
                    if pdf_path:
                        result["_pdf"] = pdf_path.name
                except Exception as e:
                    _write_log(f"PDF descriptivo error: {e}")

            # Formato WinCam (o ambos)
            if formato in ("wincam", "ambos"):
                _write_log("Llamando a analyze_wincam()...")
                wc_report = analyze_wincam(saved_files, context, api_key,
                                           proyecto=proyecto or num_ref,
                                           calle=calle, poblacion=poblacion,
                                           croquis_path=croquis_path)
                _write_log("analyze_wincam() OK")
                attach_wincam_diagramas(wc_report)
                result["_videos"] = _marcar_timestamps_video(wc_report, session_id, videos_persistidos)
                result["_wincam"] = wc_report

                # Copia reducida del informe (solo lo necesario para el visor
                # publico de solo-video: secciones/observaciones + videos) para
                # poder recuperarla despues desde /ver/<token>.
                try:
                    informe_compartible = {
                        "proyecto": wc_report.get("proyecto") or proyecto,
                        "num_ref": num_ref,
                        "cliente": cliente,
                        "nivel_urgencia_global": wc_report.get("nivel_urgencia_global"),
                        "secciones": wc_report.get("secciones") or [],
                        "totales": wc_report.get("totales") or {},
                        "videos": result.get("_videos") or [],
                    }
                    INFORMES_COMPARTIDOS_DIR.mkdir(parents=True, exist_ok=True)
                    (INFORMES_COMPARTIDOS_DIR / f"Informe_WinCam_{session_id}.json").write_text(
                        json.dumps(informe_compartible, ensure_ascii=False), encoding="utf-8"
                    )
                except Exception as e:
                    _write_log(f"No se pudo guardar el informe compartible: {e}")

                wc_name = f"{informe_stem} WINCAM.docx" if informe_stem else f"Informe_WinCam_{session_id}.docx"
                wc_path = SALIDAS_DIR / wc_name
                try:
                    generate_wincam_docx(wc_report, wc_path, num_ref=num_ref, cliente=cliente,
                                         enlace_video=enlace_video)
                    result["_docx_wincam"] = wc_name
                except Exception as e:
                    result["_wincam_docx_error"] = str(e)
                    _write_log(f"DOCX WinCam error: {e}")

            # Si no se genero el formato WinCam (p.ej. formato="descriptivo")
            # igualmente se listan los videos persistidos para poder revisarlos
            # en el visor, aunque sin marcadores de observaciones (esos requieren
            # el mapeo de fotogramas del informe WinCam).
            if "_videos" not in result and videos_persistidos:
                result["_videos"] = [
                    {"nombre": web_name, "url": url_for("servir_video_ia", session_id=session_id, filename=web_name)}
                    for web_name in videos_persistidos.values()
                ]

            # Plano tecnico del croquis aportado (DOCX + PDF). Antes el croquis solo
            # se usaba como referencia para la IA y no producia ningun documento.
            if croquis_path:
                _write_log("Generando plano tecnico del croquis...")
                try:
                    svg, _estructura = _plano_svg_desde_archivo(
                        croquis_path, api_key=api_key,
                        contexto=f"{proyecto} {calle} {poblacion}".strip())
                    titulo_plano = f"Croquis de red - {proyecto}" if proyecto else "Croquis de red"
                    direccion_plano = ", ".join(x for x in (calle, poblacion) if x)
                    fecha_plano = datetime.now().strftime("%d/%m/%Y")
                    stem_plano = f"Croquis_{session_id}"
                    png = _svg_a_png(svg)
                    cx_docx, cx_pdf = _croquis_docx_desde_png(
                        png, titulo=titulo_plano, num_ref=num_ref,
                        direccion=direccion_plano, fecha_display=fecha_plano,
                        stem=stem_plano,
                    )
                    result["_croquis_docx"] = cx_docx.name
                    if cx_pdf:
                        result["_croquis_pdf"] = cx_pdf.name
                    result["_croquis_svg"] = svg
                    result["_croquis_meta"] = {
                        "titulo": titulo_plano, "num_ref": num_ref,
                        "direccion": direccion_plano, "fecha_display": fecha_plano,
                        "stem": stem_plano,
                    }
                    _write_log(f"Plano OK: {cx_docx.name} / {cx_pdf.name if cx_pdf else 'sin PDF'}")
                except Exception as e:
                    result["_croquis_error"] = str(e)
                    _write_log(f"Plano croquis error: {e}")

            _write_log(f"Trabajo {session_id} terminado OK")
            with _ANALYSIS_JOBS_LOCK:
                _ANALYSIS_JOBS[session_id] = {"status": "done", "result": result, "_ts": time.time()}

        except Exception as e:
            tb = _tb_mod.format_exc()
            _write_log("=== EXCEPCION (segundo plano) ===\n" + tb)
            with _ANALYSIS_JOBS_LOCK:
                _ANALYSIS_JOBS[session_id] = {
                    "status": "error",
                    "error": f"{str(e)}\n\n--- TRACEBACK ---\n{tb}",
                    "traceback": tb,
                    "_ts": time.time(),
                }


@app.route("/api/analizar_status/<job_id>")
def api_analizar_status(job_id):
    """Sondeado por el frontend tras /api/analizar para recoger el resultado
    del trabajo en segundo plano (ver _procesar_analisis_bg)."""
    with _ANALYSIS_JOBS_LOCK:
        job = _ANALYSIS_JOBS.get(job_id)
    if not job:
        return jsonify({"status": "not_found", "error": "Trabajo no encontrado (puede haber caducado)."}), 404
    return jsonify(job)


@app.route("/api/regenerar_plano_svg", methods=["POST"])
def api_regenerar_plano_svg():
    """Rasteriza un SVG de plano ya editado en el navegador y regenera el DOCX/PDF,
    sin volver a analizar el croquis original con la IA."""
    import traceback as _tb
    try:
        data = request.get_json(force=True) or {}
        svg = (data.get("svg") or "").strip()
        if not svg.startswith("<svg"):
            return jsonify({"error": "SVG no valido."}), 400
        titulo = data.get("titulo") or "Croquis de red"
        num_ref = data.get("num_ref", "")
        direccion = data.get("direccion", "")
        fecha_display = data.get("fecha_display") or datetime.now().strftime("%d/%m/%Y")
        stem = data.get("stem") or _safe_filename(f"Croquis_{num_ref or titulo}", maxlen=50)

        png = _svg_a_png(svg)
        cx_docx, cx_pdf = _croquis_docx_desde_png(
            png, titulo=titulo, num_ref=num_ref, direccion=direccion,
            fecha_display=fecha_display, stem=stem)
        return jsonify({
            "croquis_docx": cx_docx.name,
            "croquis_pdf": cx_pdf.name if cx_pdf else None,
        })
    except Exception as e:
        return jsonify({"error": str(e), "traceback": _tb.format_exc()}), 500


@app.route("/api/generar_partidas_ia", methods=["POST"])
def api_generar_partidas_ia():
    """Generate budget items + technical texts using Claude, then enrich with tarifas prices."""
    import traceback as _tb
    try:
        from core.ai_analyst import generate_partidas_ia
        # Aceptar tanto JSON como FormData (cuando hay archivos adjuntos)
        ct = request.content_type or ""
        if "multipart/form-data" in ct:
            descripcion = request.form.get("descripcion", "").strip()
            tipo        = request.form.get("tipo", "obra_1")
            api_key     = request.form.get("api_key", "").strip()
            informe_raw = request.form.get("informe", "")
            try:
                informe = json.loads(informe_raw) if informe_raw else None
            except Exception:
                informe = None
            # Guardar archivos subidos temporalmente
            _ia_tmp = Path(tempfile.mkdtemp(prefix="ia_gen_"))
            saved_files: list[Path] = []
            for fobj in request.files.getlist("archivos_ia"):
                suf = Path(fobj.filename or "").suffix.lower()
                if suf in {".pdf", ".docx", ".doc", ".txt", ".md"}:
                    dest = _ia_tmp / (uuid.uuid4().hex[:6] + suf)
                    fobj.save(dest)
                    saved_files.append(dest)
        else:
            data = request.get_json(force=True)
            descripcion = data.get("descripcion", "").strip()
            tipo        = data.get("tipo", "obra_1")
            api_key     = data.get("api_key", "").strip()
            informe     = data.get("informe") or None  # existing IA report or null
            saved_files = []
            _ia_tmp = None

        tarifas = get_tarifas()
        result = generate_partidas_ia(
            descripcion, tipo, api_key, informe,
            tarifas_items=tarifas.all_items(),
            files=saved_files or None,
        )

        if "_error" in result:
            return jsonify({"error": result["_error"], "raw": result.get("_raw", "")}), 500

        # Enrich partidas: lookup por codigo exacto en tarifas; si no hay codigo o no
        # se encuentra, usa la estimacion de precio de mercado que ya aporta la IA
        # (descripcion_libre/unidad_libre/precio_estimado) en vez de dejarla a 0.
        def _enrich(partidas: list) -> list:
            enriched = []
            for p in (partidas or []):
                codigo = p.get("codigo", "").strip()
                found = tarifas.lookup(codigo) if codigo else None
                if not found and codigo:
                    # Fallback: fuzzy search on the code string itself
                    found = tarifas.search(codigo)
                if found:
                    enriched.append({
                        "codigo":          found["codigo"],
                        "descripcion":     found["descripcion"],
                        "unidad":          found["unidad"],
                        "cantidad":        float(p.get("cantidad", 1)),
                        "precio_unitario": found["precio"],
                        "tarifa_encontrada": True,
                        "nota":            p.get("nota", ""),
                    })
                else:
                    precio_est = float(p.get("precio_estimado", 0) or 0)
                    enriched.append({
                        "codigo":          "",
                        "descripcion":     p.get("descripcion_libre", "").strip() or codigo or "Partida sin catalogar",
                        "unidad":          p.get("unidad_libre", "").strip() or "ud",
                        "cantidad":        float(p.get("cantidad", 1)),
                        "precio_unitario": precio_est,
                        "tarifa_encontrada": False,
                        "nota":            p.get("nota", "") or (
                            "Precio estimado de mercado (sin referencia en catalogo)" if precio_est else ""),
                    })
            return enriched

        if tipo == "obra_2":
            result["partidas_a"] = _enrich(result.get("partidas_a", []))
            result["partidas_b"] = _enrich(result.get("partidas_b", []))
        else:
            result["partidas"] = _enrich(result.get("partidas", []))

        return jsonify(result)

    except Exception as e:
        return jsonify({"error": str(e), "traceback": _tb.format_exc()}), 500
    finally:
        if '_ia_tmp' in dir() and _ia_tmp and _ia_tmp.exists():
            import shutil as _sh
            _sh.rmtree(_ia_tmp, ignore_errors=True)


@app.route("/api/chat_informe", methods=["POST"])
def api_chat_informe():
    """Chat con el perito IA para modificar o ampliar un informe generado."""
    import traceback as _tb
    try:
        from core.ai_analyst import (SYSTEM_PROMPT, REPORT_SCHEMA,
                                      generate_report_docx, files_to_content_blocks)

        # Acepta JSON (sin adjuntos) o multipart/form-data (con archivos nuevos a analizar)
        es_multipart = bool(request.content_type and request.content_type.startswith("multipart/"))
        if es_multipart:
            report_actual = json.loads(request.form.get("report") or "{}")
            historial = json.loads(request.form.get("historial") or "[]")
            mensaje = (request.form.get("mensaje") or "").strip()
            api_key = (request.form.get("api_key") or "").strip()
            num_ref = request.form.get("num_ref", "")
            cliente = request.form.get("cliente", "")
            adjuntos = [f for f in request.files.getlist("archivos") if f.filename]
        else:
            data = request.get_json(force=True)
            report_actual = data.get("report", {})
            historial = data.get("historial", [])  # [{role, content}]
            mensaje = data.get("mensaje", "").strip()
            api_key = data.get("api_key", "").strip()
            num_ref = data.get("num_ref", "")
            cliente = data.get("cliente", "")
            adjuntos = []

        key = api_key or os.environ.get("ANTHROPIC_API_KEY", "")
        if not key:
            return jsonify({"error": "API Key no configurada."}), 400

        import anthropic as _ant
        client = _ant.Anthropic(api_key=key)

        # Guardar y procesar los archivos nuevos que el tecnico haya arrastrado al chat
        bloques_adjuntos = []
        nuevos_evidencia = []
        if adjuntos:
            chat_dir = UPLOADS_DIR / ("chat_" + uuid.uuid4().hex[:8])
            chat_dir.mkdir(parents=True, exist_ok=True)
            saved = []
            for up in adjuntos:
                suf = Path(up.filename).suffix.lower()
                if suf not in ALLOWED_IA_EXTENSIONS:
                    continue
                dest = chat_dir / f"{uuid.uuid4().hex[:6]}{suf}"
                up.save(str(dest))
                if dest.exists() and dest.stat().st_size > 0:
                    saved.append(dest)
            if saved:
                bloques_adjuntos, nuevos_evidencia = files_to_content_blocks(saved)

        # Construir historial de conversacion
        messages = []
        for h in historial:
            messages.append({"role": h["role"], "content": h["content"]})

        # Mensaje actual del usuario
        intro = (
            f"INFORME ACTUAL EN JSON:\n{json.dumps(report_actual, ensure_ascii=False, indent=2)}\n\n"
            f"INSTRUCCION DEL TECNICO:\n{mensaje}\n\n"
        )
        if bloques_adjuntos:
            intro += (
                "El tecnico adjunta los siguientes ARCHIVOS NUEVOS (imagenes, videos o documentos). "
                "Analizalos e incorpora la informacion relevante (nuevos tramos, patologias, mediciones, "
                "fotografias, datos del documento) al informe existente. No borres lo que ya es correcto; "
                "amplia y actualiza:\n"
            )
        cierre = (
            f"\nDevuelve el informe completo actualizado en el mismo formato JSON. "
            f"Si la instruccion es solo una pregunta tecnica, responde como texto en el campo "
            f"'respuesta_chat' y manten el informe sin cambios. "
            f"Responde UNICAMENTE con JSON segun el esquema:\n{REPORT_SCHEMA}"
        )
        if bloques_adjuntos:
            user_content = [{"type": "text", "text": intro}]
            user_content.extend(bloques_adjuntos)
            user_content.append({"type": "text", "text": cierre})
        else:
            user_content = intro + cierre
        messages.append({"role": "user", "content": user_content})

        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=16000,
            system=SYSTEM_PROMPT,
            messages=messages,
        )

        raw = (response.content[0].text or "").strip()

        import re
        # Quitar fences markdown (```json ... ```)
        raw_clean = re.sub(r"^```[a-zA-Z]*\n?", "", raw)
        raw_clean = re.sub(r"\n?```\s*$", "", raw_clean).strip()
        updated = _parse_json_lenient(raw_clean)
        if isinstance(updated, dict):
            # Conservar el anexo fotografico previo (Claude no reenvia las rutas) y
            # anadir las imagenes de los archivos nuevos aportados en este turno.
            prev_ev = report_actual.get("_evidencia_img") or []
            merged_ev = list(prev_ev) + [e for e in nuevos_evidencia if e not in prev_ev]
            if merged_ev:
                updated["_evidencia_img"] = merged_ev
            # Regenerar DOCX con el informe actualizado
            session_id = uuid.uuid4().hex[:8]
            # Este endpoint (regeneracion via chat) no recibe calle/poblacion por
            # separado; se usa el titulo del informe (normalmente ya incluye la
            # direccion) como mejor aproximacion al formato "N.o - INFORME DIRECCION".
            _titulo_informe = _safe_filename(updated.get("titulo", "") or "", maxlen=60)
            if num_ref and _titulo_informe:
                docx_name = f"{_safe_filename(num_ref, maxlen=20)} - INFORME {_titulo_informe}.docx"
            elif num_ref:
                docx_name = f"{_safe_filename(num_ref, maxlen=20)} - INFORME.docx"
            else:
                docx_name = f"Informe_IA_{session_id}.docx"
            docx_path = SALIDAS_DIR / docx_name
            try:
                generate_report_docx(updated, docx_path, num_ref=num_ref, cliente=cliente)
                updated["_docx"] = docx_name
                # Regenerar tambien el PDF para que quede sincronizado (con el anexo fotografico)
                try:
                    pdf_path = PdfConverter().convert(docx_path, SALIDAS_DIR)
                    if pdf_path:
                        updated["_pdf"] = pdf_path.name
                except Exception:
                    pass
            except Exception:
                pass
            updated["_assistant_msg"] = updated.get("respuesta_chat", "Informe actualizado.")
            return jsonify({"report": updated, "raw": raw})

        # No se obtuvo JSON valido
        if getattr(response, "stop_reason", None) == "max_tokens":
            try:
                with open(BASE_DIR / "debug_chat_informe.txt", "w", encoding="utf-8") as _fh:
                    _fh.write(f"stop_reason=max_tokens\nlen_raw={len(raw)}\n=== RAW ===\n{raw}")
            except Exception:
                pass
            return jsonify({"error": "La respuesta del informe se corto por longitud (limite de tokens). Pide los cambios por partes o reintenta.", "_raw": raw[:2000]}), 500
        # Respuesta conversacional (pregunta tecnica, sin cambios al informe)
        return jsonify({"report": report_actual, "raw": raw, "assistant_msg": raw_clean or raw})

    except Exception as e:
        tb = _tb.format_exc()
        return jsonify({"error": str(e), "traceback": tb}), 500


def _parse_json_lenient(text):
    """Parsea JSON de la respuesta del modelo de forma tolerante (fences, preambulo, comas colgantes)."""
    if not text:
        return None
    import re as _re
    # 1) intento directo
    try:
        return json.loads(text)
    except Exception:
        pass
    # 2) primer '{' hasta ultimo '}'
    m = _re.search(r"\{[\s\S]*\}", text)
    if not m:
        return None
    frag = m.group(0)
    try:
        return json.loads(frag)
    except Exception:
        pass
    # 3) quitar comas colgantes antes de } o ]
    frag2 = _re.sub(r",\s*([}\]])", r"\1", frag)
    try:
        return json.loads(frag2)
    except Exception:
        return None


@app.route("/api/importar_informe", methods=["POST"])
def api_importar_informe():
    """Importa un informe existente (DOCX/PDF/TXT) y lo convierte al schema JSON para edicion via chat."""
    import traceback as _tb
    import base64
    import re
    tmp_dir = None
    try:
        from core.ai_analyst import SYSTEM_PROMPT, REPORT_SCHEMA
        import anthropic as _ant

        archivo = request.files.get("archivo")
        api_key = request.form.get("api_key", "").strip()

        key = api_key or os.environ.get("ANTHROPIC_API_KEY", "")
        if not key:
            return jsonify({"error": "API Key no configurada."}), 400
        if not archivo or not archivo.filename:
            return jsonify({"error": "No se recibio ningun archivo."}), 400

        suffix = Path(archivo.filename).suffix.lower()
        if suffix not in {".pdf", ".docx", ".doc", ".txt", ".md"}:
            return jsonify({"error": "Formato no soportado. Usa PDF, DOCX o TXT."}), 400

        tmp_dir = Path(tempfile.mkdtemp(prefix="importar_"))
        dest = tmp_dir / f"informe{suffix}"
        archivo.save(str(dest))

        client = _ant.Anthropic(api_key=key)
        instruccion = (
            "Extrae toda la informacion de este informe tecnico y estructurala en el siguiente "
            "esquema JSON. Si algun campo no aparece en el documento usa null o [] segun el tipo. "
            "Devuelve UNICAMENTE el JSON sin ningun texto adicional.\n\n"
            f"Esquema requerido:\n{REPORT_SCHEMA}"
        )

        if suffix == ".pdf":
            pdf_data = dest.read_bytes()
            messages = [{
                "role": "user",
                "content": [
                    {"type": "document", "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": base64.b64encode(pdf_data).decode(),
                    }},
                    {"type": "text", "text": instruccion},
                ],
            }]
        else:
            if suffix in (".docx", ".doc"):
                try:
                    from docx import Document as _DocxDoc
                    doc = _DocxDoc(str(dest))
                    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                except Exception:
                    text = dest.read_text(encoding="utf-8", errors="replace")
            else:
                text = dest.read_text(encoding="utf-8", errors="replace")
            messages = [{"role": "user", "content": f"INFORME:\n{text}\n\n{instruccion}"}]

        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=16000,
            system=SYSTEM_PROMPT,
            messages=messages,
        )
        raw = (response.content[0].text or "").strip()
        # Quitar fences markdown (```json ... ```)
        raw = re.sub(r"^```[a-zA-Z]*\n?", "", raw)
        raw = re.sub(r"\n?```\s*$", "", raw).strip()
        report = _parse_json_lenient(raw)
        if report is not None:
            return jsonify(report)
        # Volcado de diagnostico cuando el parseo falla
        try:
            with open(BASE_DIR / "debug_importar_informe.txt", "w", encoding="utf-8") as _fh:
                _fh.write(f"stop_reason={getattr(response, 'stop_reason', None)}\n")
                _fh.write(f"len_raw={len(raw)}\n=== RAW ===\n{raw}")
        except Exception:
            pass
        if getattr(response, "stop_reason", None) == "max_tokens":
            return jsonify({"error": "El informe es demasiado largo y la respuesta se corto (limite de tokens). Divide el documento o reintenta.", "_raw": raw[:2000]}), 500
        return jsonify({"error": "No se pudo extraer JSON del informe.", "_raw": raw[:2000]}), 500

    except Exception as e:
        return jsonify({"error": str(e), "traceback": _tb.format_exc()}), 500
    finally:
        if tmp_dir and tmp_dir.exists():
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)


@app.route("/api/importar_presupuesto", methods=["POST"])
def api_importar_presupuesto():
    """Importa un presupuesto existente (DOCX/PDF) y extrae todos los campos al schema JSON."""
    import traceback as _tb
    import base64
    import re
    tmp_dir = None
    try:
        import anthropic as _ant

        archivo = request.files.get("archivo")
        api_key = request.form.get("api_key", "").strip()

        key = api_key or os.environ.get("ANTHROPIC_API_KEY", "")
        if not key:
            return jsonify({"error": "API Key no configurada."}), 400
        if not archivo or not archivo.filename:
            return jsonify({"error": "No se recibio ningun archivo."}), 400

        suffix = Path(archivo.filename).suffix.lower()
        if suffix not in {".pdf", ".docx", ".doc", ".txt"}:
            return jsonify({"error": "Formato no soportado. Usa PDF o DOCX."}), 400

        tmp_dir = Path(tempfile.mkdtemp(prefix="imp_pres_"))
        dest = tmp_dir / f"presupuesto{suffix}"
        archivo.save(str(dest))

        schema = """{
  "tipo_modelo": "obra_1|obra_2|desatasco|limpieza_aerea|fresador|robot_limpieza|fuga_agua|vaciado_fosa|cctv_bajante|inspeccion_zum|informe_desatasco|bajantes_amianto|certificado_obra|plan_seguridad|contrato_saneamiento|fontaneria|albanileria|contrato_subcontrata",
  "num_contrato": "",
  "fecha_contrato": "YYYY-MM-DD",
  "cliente_nombre": "",
  "cliente_dir": "",
  "cliente_tel": "",
  "cliente_email": "",
  "obra": "",
  "servicio": "",
  "provincia": "Madrid",
  "administracion": "",
  "admin_tel": "",
  "admin_email": "",
  "informe": "",
  "solucion": "",
  "memoria": "",
  "plazo": "30",
  "forma_pago": "",
  "partidas": [
    {"descripcion": "", "unidad": "ud", "cantidad": 1, "precio_unitario": 0}
  ],
  "partidas_b": []
}"""

        guia_tipos = (
            "GUIA DE CLASIFICACION (elige el tipo_modelo segun el TITULO y CONTENIDO real del documento, "
            "NO por palabras sueltas como 'saneamiento'):\n"
            "- obra_1: presupuesto de obra con UNA SOLA opcion y tabla de partidas (es lo MAS COMUN para presupuestos de obra civil/saneamiento).\n"
            "- obra_2: presupuesto de obra con DOS opciones A y B (dos tablas de partidas distintas).\n"
            "- desatasco: presupuesto de desatasco puntual (camion cuba, hidropresion).\n"
            "- robot_limpieza: limpieza con robot fresador/cortador en colectores.\n"
            "- limpieza_aerea: limpieza de bajantes desde cubierta (plataforma elevadora).\n"
            "- cctv_bajante: inspeccion CCTV con camara en bajantes o colectores.\n"
            "- inspeccion_zum: inspeccion ZUM (camara empujada).\n"
            "- fresador: trabajos de fresado de raices/incrustaciones.\n"
            "- fuga_agua: deteccion de fuga de agua.\n"
            "- vaciado_fosa: vaciado de fosa septica o separadora de grasas.\n"
            "- informe_desatasco: INFORME tecnico (no presupuesto) de un desatasco.\n"
            "- bajantes_amianto: sustitucion de bajantes que contienen amianto.\n"
            "- certificado_obra: CERTIFICADO final de obra (no presupuesto).\n"
            "- plan_seguridad: Plan de Seguridad y Salud (PSS).\n"
            "- contrato_saneamiento: CONTRATO de mantenimiento de saneamiento (recurring). Solo si el documento dice claramente 'CONTRATO' y describe servicios periodicos. NO uses este tipo para presupuestos de obra puntual.\n"
            "- fontaneria: presupuesto de trabajos de fontaneria.\n"
            "- albanileria: presupuesto de trabajos de albanileria.\n"
            "- contrato_subcontrata: contrato con subcontratista.\n"
            "REGLA DE DESEMPATE: si el documento tiene una tabla de partidas con precios y NO es un contrato recurrente, "
            "es casi seguro 'obra_1' (o 'obra_2' si hay dos opciones)."
        )

        instruccion = (
            "Extrae TODOS los datos de este presupuesto y devuelvelos en el siguiente esquema JSON.\n\n"
            f"{guia_tipos}\n\n"
            "Para fecha_contrato usa formato YYYY-MM-DD; si no aparece usa null. "
            "Extrae todas las partidas con descripcion, unidad, cantidad y precio unitario. "
            "Si hay dos opciones A y B, las partidas de la opcion A van en 'partidas' y las de la B en 'partidas_b', "
            "y tipo_modelo debe ser 'obra_2'. "
            "Devuelve UNICAMENTE el JSON sin texto adicional.\n\n"
            f"Esquema:\n{schema}"
        )

        client = _ant.Anthropic(api_key=key)

        if suffix == ".pdf":
            pdf_data = dest.read_bytes()
            messages = [{
                "role": "user",
                "content": [
                    {"type": "document", "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": base64.b64encode(pdf_data).decode(),
                    }},
                    {"type": "text", "text": instruccion},
                ],
            }]
        else:
            if suffix in (".docx", ".doc"):
                try:
                    from docx import Document as _DocxDoc
                    doc = _DocxDoc(str(dest))
                    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    # Also extract tables
                    for tbl in doc.tables:
                        for row in tbl.rows:
                            text += "\n" + "\t".join(c.text for c in row.cells if c.text.strip())
                except Exception:
                    text = dest.read_text(encoding="utf-8", errors="replace")
            else:
                text = dest.read_text(encoding="utf-8", errors="replace")
            messages = [{"role": "user", "content": f"PRESUPUESTO:\n{text}\n\n{instruccion}"}]

        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=8192,
            system=(
                "Eres un asistente experto en presupuestos de construccion y saneamiento. "
                "Extraes datos estructurados de documentos y los devuelves siempre en JSON valido."
            ),
            messages=messages,
        )
        raw = response.content[0].text.strip()
        match = re.search(r"\{[\s\S]*\}", raw)
        if match:
            try:
                data = json.loads(match.group())
                return jsonify(data)
            except json.JSONDecodeError:
                pass
        return jsonify({"error": "No se pudo extraer JSON del presupuesto.", "_raw": raw[:2000]}), 500

    except Exception as e:
        return jsonify({"error": str(e), "traceback": _tb.format_exc()}), 500
    finally:
        if tmp_dir and tmp_dir.exists():
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)


@app.route("/api/aplicar_modificaciones_ia", methods=["POST"])
def api_aplicar_modificaciones_ia():
    """Aplica modificaciones IA sobre un presupuesto ya importado.

    Recibe (FormData):
      - presupuesto_json: JSON string con el presupuesto actual (mismo schema que /api/importar_presupuesto)
      - instrucciones: texto libre con los cambios a aplicar
      - archivo: PDF/imagen opcional con notas de modificacion
      - api_key: opcional

    Devuelve el JSON modificado. NO modifica tipo_modelo (lo respeta tal cual viene).
    """
    import traceback as _tb
    import base64
    import re
    tmp_dir = None
    try:
        import anthropic as _ant

        pres_raw = request.form.get("presupuesto_json", "").strip()
        instrucciones = request.form.get("instrucciones", "").strip()
        api_key = request.form.get("api_key", "").strip()
        archivos = [f for f in request.files.getlist("archivo") if f and f.filename]

        key = api_key or os.environ.get("ANTHROPIC_API_KEY", "")
        if not key:
            return jsonify({"error": "API Key no configurada."}), 400
        if not pres_raw:
            return jsonify({"error": "Falta el presupuesto_json."}), 400
        if not instrucciones and not archivos:
            return jsonify({"error": "Debes indicar instrucciones (texto) o subir archivos con las modificaciones."}), 400

        try:
            presupuesto = json.loads(pres_raw)
        except json.JSONDecodeError as e:
            return jsonify({"error": f"presupuesto_json invalido: {e}"}), 400

        tipo_original = presupuesto.get("tipo_modelo", "")

        client = _ant.Anthropic(api_key=key)

        contenido = []
        if archivos:
            tmp_dir = Path(tempfile.mkdtemp(prefix="mod_pres_"))
            for idx, archivo in enumerate(archivos):
                suffix = Path(archivo.filename).suffix.lower()
                if suffix not in {".pdf", ".jpg", ".jpeg", ".png", ".gif", ".webp", ".txt", ".docx", ".doc"}:
                    return jsonify({"error": f"Formato no soportado para modificaciones: {archivo.filename}"}), 400
                dest = tmp_dir / f"mods_{idx}{suffix}"
                archivo.save(str(dest))
                data_bytes = dest.read_bytes()
                contenido.append({"type": "text", "text": f"\n[ARCHIVO ADJUNTO: {archivo.filename}]"})

                if suffix == ".pdf":
                    contenido.append({"type": "document", "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": base64.b64encode(data_bytes).decode(),
                    }})
                elif suffix in (".jpg", ".jpeg", ".png", ".gif", ".webp"):
                    media = {".jpg":"image/jpeg",".jpeg":"image/jpeg",".png":"image/png",".gif":"image/gif",".webp":"image/webp"}[suffix]
                    contenido.append({"type": "image", "source": {
                        "type": "base64",
                        "media_type": media,
                        "data": base64.b64encode(data_bytes).decode(),
                    }})
                elif suffix in (".docx", ".doc"):
                    try:
                        from docx import Document as _DocxDoc
                        doc = _DocxDoc(str(dest))
                        text_doc = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                        for tbl in doc.tables:
                            for row in tbl.rows:
                                text_doc += "\n" + "\t".join(c.text for c in row.cells if c.text.strip())
                        contenido.append({"type": "text", "text": f"(texto extraido del DOCX):\n{text_doc}"})
                    except Exception:
                        pass
                else:
                    try:
                        text_doc = dest.read_text(encoding="utf-8", errors="replace")
                        contenido.append({"type": "text", "text": text_doc})
                    except Exception:
                        pass

        prompt_mod = (
            "Aplica las siguientes modificaciones sobre el presupuesto JSON proporcionado. "
            "Devuelve el JSON COMPLETO modificado con el MISMO esquema de campos (no inventes campos nuevos).\n\n"
            "REGLAS ESTRICTAS:\n"
            "1. NO modifiques el campo 'tipo_modelo' bajo ninguna circunstancia (devuelvelo tal cual viene).\n"
            "2. Puedes modificar: partidas (anadir/quitar/cambiar precios o cantidades), partidas_b, "
            "informe, solucion, memoria, plazo, forma_pago, cliente_nombre, cliente_dir, cliente_tel, cliente_email, "
            "obra, servicio, provincia, administracion, admin_tel, admin_email, num_contrato, fecha_contrato.\n"
            "3. Si las instrucciones piden anadir partidas, anadelas al array 'partidas' (o 'partidas_b' si se indica opcion B).\n"
            "4. Si piden quitar una partida, eliminala del array.\n"
            "5. Conserva los campos no afectados por las modificaciones.\n"
            "6. Devuelve UNICAMENTE el JSON valido sin texto adicional.\n\n"
            f"PRESUPUESTO ACTUAL (JSON):\n{json.dumps(presupuesto, ensure_ascii=False, indent=2)}\n\n"
            f"INSTRUCCIONES DE MODIFICACION:\n{instrucciones or '(ver archivo adjunto)'}"
        )

        contenido.append({"type": "text", "text": prompt_mod})

        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=8192,
            system=(
                "Eres un asistente experto en presupuestos de construccion y saneamiento. "
                "Modificas presupuestos en JSON segun las instrucciones del usuario, devolviendo siempre JSON valido."
            ),
            messages=[{"role": "user", "content": contenido}],
        )
        raw = response.content[0].text.strip()
        match = re.search(r"\{[\s\S]*\}", raw)
        if match:
            try:
                data = json.loads(match.group())
                # Forzar tipo_modelo original (seguridad por si la IA lo cambia)
                if tipo_original:
                    data["tipo_modelo"] = tipo_original
                return jsonify(data)
            except json.JSONDecodeError:
                pass
        return jsonify({"error": "No se pudo extraer JSON modificado.", "_raw": raw[:2000]}), 500

    except Exception as e:
        return jsonify({"error": str(e), "traceback": _tb.format_exc()}), 500
    finally:
        if tmp_dir and tmp_dir.exists():
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)


@app.route("/api/subir_evidencia", methods=["POST"])
def api_subir_evidencia():
    """Recibe videos/imagenes para el anexo fotografico de un informe.
    Extrae fotogramas de los videos, guarda todas las imagenes en una carpeta
    de sesion (UPLOADS_DIR/evid_xxx) y devuelve un token para adjuntarlas al
    documento en el momento de generar (via campo oculto 'evidencia_token')."""
    import traceback as _tb
    import shutil as _sh
    try:
        from core.ai_analyst import files_to_content_blocks
        archivos = [f for f in request.files.getlist("archivos") if f.filename]
        if not archivos:
            return jsonify({"error": "No se recibieron archivos."}), 400

        token = "evid_" + uuid.uuid4().hex[:10]
        sess = UPLOADS_DIR / token
        sess.mkdir(parents=True, exist_ok=True)

        saved = []
        for up in archivos:
            suf = Path(up.filename).suffix.lower()
            if suf not in ALLOWED_IA_EXTENSIONS:
                continue
            dest = sess / f"src_{uuid.uuid4().hex[:6]}{suf}"
            up.save(str(dest))
            if dest.exists() and dest.stat().st_size > 0:
                saved.append(dest)
        if not saved:
            return jsonify({"error": "Ningun archivo valido (usa video o imagen)."}), 400

        # Extrae fotogramas de los videos y recoge las imagenes -> lista de rutas
        _, evidencia = files_to_content_blocks(saved)

        # Copia todas las imagenes (frames + fotos) al folder de sesion, numeradas
        fotos = []
        for i, p in enumerate(evidencia):
            src = Path(p)
            if not src.exists():
                continue
            dst = sess / f"foto_{i:03d}{src.suffix.lower()}"
            try:
                if src.resolve() != dst.resolve():
                    _sh.copyfile(src, dst)
                fotos.append(dst.name)
            except Exception:
                pass

        return jsonify({"token": token, "num_fotos": len(fotos)})
    except Exception as e:
        return jsonify({"error": str(e), "traceback": _tb.format_exc()}), 500


@app.route("/api/tarifa/<code>")
def api_tarifa(code: str):
    item = get_tarifas().lookup(code)
    if not item:
        return jsonify({}), 404
    return jsonify(item)


@app.route("/api/siguiente_numero")
def api_siguiente_numero():
    """Devuelve el siguiente numero de presupuesto para un tipo dado.
    Parametros GET: tipo, num_base (opcional, para revision)
    Responde: {numero, sin_contador} donde sin_contador=True indica que el usuario
    debe introducir el mismo numero que el presupuesto asociado.
    """
    from utils.base_datos import TIPOS_SIN_CONTADOR
    tipo = request.args.get("tipo", "")
    num_base_str = request.args.get("num_base", "")
    num_base = int(num_base_str) if num_base_str.isdigit() else None
    sin_contador = tipo in TIPOS_SIN_CONTADOR
    try:
        numero = siguiente_numero(tipo, num_base)
    except Exception as e:
        return jsonify({"numero": "", "sin_contador": sin_contador, "error": str(e)})
    return jsonify({"numero": numero, "sin_contador": sin_contador})


@app.route("/api/clientes")
def api_clientes():
    """Autocomplete de clientes. Param GET: q (texto a buscar)"""
    from utils.base_datos import buscar_clientes
    q = request.args.get("q", "").strip()
    return jsonify(buscar_clientes(q, limit=15))


@app.route("/api/tarifas")
def api_tarifas_list():
    return jsonify(get_tarifas().all_codes())


@app.route("/api/tarifas_full")
def api_tarifas_full():
    return jsonify(get_tarifas().all_items())


@app.route("/api/analizar_solicitud", methods=["POST"])
def api_analizar_solicitud():
    """Extract budget form data from uploaded documents using AI."""
    files_in = request.files.getlist("archivos")
    if not files_in:
        return jsonify({"error": "No se recibieron archivos."}), 400

    api_key = request.form.get("api_key", "").strip() or os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API Key de Anthropic no configurada."}), 400

    ALLOWED = {
        ".jpg", ".jpeg", ".png", ".gif", ".webp",
        ".pdf", ".txt", ".eml", ".msg", ".html", ".htm",
    }
    tmp_dir = Path(tempfile.mkdtemp(prefix="solicitud_"))
    saved: list[Path] = []
    for fobj in files_in:
        suf = Path(fobj.filename or "").suffix.lower()
        if suf not in ALLOWED:
            continue
        dest = tmp_dir / f"{uuid.uuid4().hex}{suf}"
        fobj.save(dest)
        saved.append(dest)

    if not saved:
        return jsonify({"error": "Ningun archivo tiene formato valido (jpg, png, pdf, txt, eml)."}), 400

    try:
        result = extract_solicitud_data(saved, api_key=api_key)
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        import shutil
        shutil.rmtree(tmp_dir, ignore_errors=True)


@app.route("/api/importar_subcontrata", methods=["POST"])
def api_importar_subcontrata():
    """Extract all partidas from a subcontractor budget document using AI."""
    files_in = request.files.getlist("archivos")
    if not files_in:
        return jsonify({"error": "No se recibieron archivos."}), 400

    api_key = request.form.get("api_key", "").strip() or os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API Key de Anthropic no configurada."}), 400

    ALLOWED = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".pdf", ".txt", ".html", ".htm", ".csv"}
    tmp_dir = Path(tempfile.mkdtemp(prefix="subcontrata_"))
    saved: list[Path] = []
    for fobj in files_in:
        suf = Path(fobj.filename or "").suffix.lower()
        if suf not in ALLOWED:
            continue
        dest = tmp_dir / f"{uuid.uuid4().hex}{suf}"
        fobj.save(dest)
        saved.append(dest)

    if not saved:
        return jsonify({"error": "Formato no valido. Sube PDF, imagen o texto."}), 400

    try:
        result = extract_partidas_from_subcontrata(saved, api_key=api_key)
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        import shutil
        shutil.rmtree(tmp_dir, ignore_errors=True)


_MATCH_STOP = {"de", "del", "la", "el", "los", "las", "y", "en", "con", "por",
               "para", "a", "o", "un", "una", "al", "su", "sus", "capitulo", "capítulo"}


def _match_norm(s: str) -> str:
    import unicodedata
    s = unicodedata.normalize("NFKD", s or "").encode("ascii", "ignore").decode()
    return s.lower()


def _match_tokens(s: str) -> list[str]:
    import re as _re
    s = _match_norm(s)
    # Quitar prefijo tipo "CAPITULO 1 - " o "FONTANERIA - " del inicio
    s = _re.sub(r"^[^-\n]{2,45}-\s*", "", s)
    words = _re.findall(r"[a-z0-9]+", s)
    return [w for w in words if len(w) >= 3 and w not in _MATCH_STOP]


def _best_tarifa_match(desc: str, items: list[dict]):
    """Devuelve (item, score 0..1) de la tarifa mas parecida a la descripcion, o (None, 0)."""
    import difflib
    q = _match_tokens(desc)
    if not q:
        return None, 0.0
    qset = set(q)
    qnorm = _match_norm(desc)
    best, best_score = None, 0.0
    for it in items:
        dt = set(_match_tokens(it.get("descripcion", "")))
        if not dt:
            continue
        overlap = len(qset & dt) / len(qset)
        ratio = difflib.SequenceMatcher(None, qnorm, _match_norm(it.get("descripcion", ""))).ratio()
        score = 0.7 * overlap + 0.3 * ratio
        if score > best_score:
            best, best_score = it, score
    return best, best_score


@app.route("/api/importar_presupuesto_completo", methods=["POST"])
def api_importar_presupuesto_completo():
    """Extrae TODAS las partidas (multi-oficio) de un informe/presupuesto aportado.

    Devuelve el resultado en el mismo formato que consume stepRevisarIA en el frontend
    (informe_tecnico, solucion_adoptar, memoria_tecnica, partidas[]), de modo que el
    tecnico pueda editar/anadir/eliminar partidas antes de generar el presupuesto final.
    """
    files_in = request.files.getlist("archivos")
    if not files_in:
        return jsonify({"error": "No se recibieron archivos."}), 400

    api_key = request.form.get("api_key", "").strip() or os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API Key de Anthropic no configurada."}), 400
    descripcion = request.form.get("descripcion", "").strip()

    ALLOWED = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".pdf",
               ".txt", ".md", ".html", ".htm", ".csv", ".xml",
               ".doc", ".docx", ".xls", ".xlsx"}
    tmp_dir = Path(tempfile.mkdtemp(prefix="presup_completo_"))
    saved: list[Path] = []
    for fobj in files_in:
        suf = Path(fobj.filename or "").suffix.lower()
        if suf not in ALLOWED:
            continue
        dest = tmp_dir / f"{uuid.uuid4().hex}{suf}"
        fobj.save(dest)
        saved.append(dest)

    if not saved:
        return jsonify({"error": "Formato no valido. Sube PDF, imagen, Word, Excel o texto."}), 400

    try:
        result = extract_full_presupuesto(saved, api_key=api_key, descripcion=descripcion)
        if "_error" in result:
            return jsonify({"error": result["_error"], "raw": result.get("_raw", "")}), 500

        # Para partidas SIN precio (informe sin valorar): buscar la tarifa mas parecida
        # en TARIFAS.xlsx. Si hay match razonable, se usa su precio y se marca "Tarifa".
        # Si no, se deja en 0 (en blanco) para que el tecnico lo rellene.
        UMBRAL = 0.5
        try:
            items_tarifa = get_tarifas().all_items()
        except Exception:
            items_tarifa = []
        n_tarifa = 0
        for p in result.get("partidas", []):
            if p.get("precio_unitario", 0) or not items_tarifa:
                continue
            match, score = _best_tarifa_match(p.get("descripcion", ""), items_tarifa)
            if match and match.get("precio", 0) > 0 and score >= UMBRAL:
                p["precio_unitario"] = match["precio"]
                p["codigo"] = match["codigo"]
                p["tarifa_encontrada"] = True
                p["importe"] = round(p.get("cantidad", 1) * match["precio"], 2)
                nota_prev = (p.get("nota") or "").strip()
                p["nota"] = (nota_prev + " " if nota_prev else "") + f"Precio tomado de tarifa {match['codigo']} (similitud {round(score*100)}%)"
                n_tarifa += 1
        result["_partidas_con_tarifa"] = n_tarifa
        return jsonify(result)
    except Exception as e:
        import traceback as _tb
        return jsonify({"error": str(e), "traceback": _tb.format_exc()}), 500
    finally:
        import shutil
        shutil.rmtree(tmp_dir, ignore_errors=True)


# --- Subcontratas guardadas (config/subcontratas.json) ---------------------
SUBCONTRATAS_PATH = BASE_DIR / "config" / "subcontratas.json"
_SUBC_KEYS = ["empresa", "cif", "domicilio", "telefono", "email", "rep_nombre",
              "rep_dni", "rep_domicilio", "notario", "notario_loc",
              "fecha_escritura", "protocolo", "registro"]


def _load_subcontratas() -> list:
    try:
        with open(SUBCONTRATAS_PATH, encoding="utf-8") as fh:
            data = json.load(fh)
        return data if isinstance(data, list) else []
    except Exception:
        return []


def _save_subcontratas(items: list) -> None:
    SUBCONTRATAS_PATH.parent.mkdir(exist_ok=True)
    with open(SUBCONTRATAS_PATH, "w", encoding="utf-8") as fh:
        json.dump(items, fh, indent=2, ensure_ascii=False)


@app.route("/api/subcontratas", methods=["GET"])
def api_subcontratas_list():
    return jsonify(_load_subcontratas())


@app.route("/api/subcontratas", methods=["POST"])
def api_subcontratas_save():
    body = request.get_json(silent=True) or {}
    datos = {k: str(body.get(k, "") or "").strip() for k in _SUBC_KEYS}
    if not datos["empresa"]:
        return jsonify({"error": "Falta el nombre de la empresa."}), 400
    items = _load_subcontratas()
    cif = datos["cif"].upper()
    name = datos["empresa"].lower()
    replaced = False
    for i, it in enumerate(items):
        same_cif = cif and (it.get("cif", "").strip().upper() == cif)
        same_name = (not cif) and (it.get("empresa", "").strip().lower() == name)
        if same_cif or same_name:
            items[i] = datos
            replaced = True
            break
    if not replaced:
        items.append(datos)
    _save_subcontratas(items)
    return jsonify({"ok": True, "guardada": datos["empresa"], "count": len(items)})


@app.route("/api/subcontratas/<path:empresa>", methods=["DELETE"])
def api_subcontratas_delete(empresa):
    key = empresa.strip().lower()
    items = [it for it in _load_subcontratas() if it.get("empresa", "").strip().lower() != key]
    _save_subcontratas(items)
    return jsonify({"ok": True, "count": len(items)})


@app.route("/api/extraer_datos_subcontrata", methods=["POST"])
def api_extraer_datos_subcontrata():
    """Extrae los datos identificativos de la subcontrata de un documento (PDF/imagen/DOCX/texto)."""
    files_in = request.files.getlist("archivos")
    if not files_in:
        return jsonify({"error": "No se recibieron archivos."}), 400
    api_key = request.form.get("api_key", "").strip() or os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "API Key de Anthropic no configurada."}), 400

    ALLOWED = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".pdf",
               ".txt", ".html", ".htm", ".csv", ".doc", ".docx"}
    tmp_dir = Path(tempfile.mkdtemp(prefix="subc_datos_"))
    saved: list[Path] = []
    for fobj in files_in:
        suf = Path(fobj.filename or "").suffix.lower()
        if suf not in ALLOWED:
            continue
        dest = tmp_dir / f"{uuid.uuid4().hex}{suf}"
        fobj.save(dest)
        saved.append(dest)

    if not saved:
        return jsonify({"error": "Formato no valido. Sube PDF, imagen, DOCX o texto."}), 400

    try:
        result = extract_subcontrata_datos(saved, api_key=api_key)
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        import shutil
        shutil.rmtree(tmp_dir, ignore_errors=True)


@app.route("/generar", methods=["POST"])
def generar():
    f = request.form
    tipo = f.get("tipo_modelo", "obra_1")

    common_data, num_contrato, fecha_larga, obra, servicio = _build_common_data(f)

    items = items_a = items_b = subc_items = None
    total_sin_iva = total_a = total_b = total_estimado = 0.0

    if tipo == "contrato_subcontrata":
        subc_items = _parse_items(f, "item")
        total_sin_iva = sum(i["importe"] for i in (subc_items or []))
        # Fechas contrato
        fecha_contrato_raw = f.get("fecha_contrato", datetime.now().strftime("%Y-%m-%d"))
        meses_esp = ["enero", "febrero", "marzo", "abril", "mayo", "junio",
                     "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
        try:
            dt = datetime.strptime(fecha_contrato_raw, "%Y-%m-%d")
            dia_semana = ["lunes","martes","miercoles","jueves","viernes","sabado","domingo"][dt.weekday()]
            fecha_contrato_larga = f"{dia_semana}, {dt.day} de {meses_esp[dt.month-1]} de {dt.year}"
            fecha_aceptacion = f"{dt.day} de {meses_esp[dt.month-1]} de {dt.year}"
        except ValueError:
            fecha_contrato_larga = fecha_contrato_raw
            fecha_aceptacion = fecha_contrato_raw
        # Fecha inicio trabajos
        fecha_inicio_raw = f.get("fecha_inicio_trabajos", "")
        if fecha_inicio_raw:
            try:
                dt2 = datetime.strptime(fecha_inicio_raw, "%Y-%m-%d")
                fecha_inicio_str = f"{dt2.day} de {meses_esp[dt2.month-1].upper()} de {dt2.year}"
            except ValueError:
                fecha_inicio_str = fecha_inicio_raw
        else:
            fecha_inicio_str = ""
        # Importe en letra
        letra = numero_a_letras(total_sin_iva) + " EUROS"
        # Build data dict
        data = {
            **common_data,
            "[[FECHA_CONTRATO]]":          fecha_contrato_larga,
            "[[FECHA_ACEPTACION]]":        fecha_aceptacion,
            "[[PRESUPUESTO_NUM]]":         f.get("num_contrato", "").strip(),
            "[[OBRA_DIRECCION]]":          obra,
            "[[IMPORTE_TOTAL]]":           fmt_euro_plain(total_sin_iva),
            "[[IMPORTE_LETRA]]":           letra,
            "[[FORMA_PAGO]]":              f.get("forma_pago", "100% a la finalizacion de los trabajos mediante pagare a 90 dias").strip(),
            "[[PLAZO_EJECUCION]]":         f.get("plazo", "30").strip(),
            "[[FECHA_INICIO_TRABAJOS]]":   fecha_inicio_str,
            # Subcontratista
            "[[SUBC_EMPRESA]]":            f.get("subc_empresa", "").strip(),
            "[[SUBC_CIF]]":               f.get("subc_cif", "").strip(),
            "[[SUBC_DOMICILIO]]":          f.get("subc_domicilio", "").strip(),
            "[[SUBC_REP_NOMBRE]]":         f.get("subc_rep_nombre", "").strip(),
            "[[SUBC_REP_DNI]]":           f.get("subc_rep_dni", "").strip(),
            "[[SUBC_REP_DOMICILIO]]":      f.get("subc_rep_domicilio", "").strip(),
            "[[SUBC_NOTARIO]]":           f.get("subc_notario", "").strip(),
            "[[SUBC_NOTARIO_LOC]]":        f.get("subc_notario_loc", "").strip(),
            "[[SUBC_FECHA_ESCRITURA]]":    f.get("subc_fecha_escritura", "").strip(),
            "[[SUBC_PROTOCOLO]]":         f.get("subc_protocolo", "").strip(),
            "[[SUBC_REGISTRO]]":          f.get("subc_registro", "").strip(),
            "[[SUBC_TELEFONO]]":          f.get("subc_telefono", "").strip(),
            "[[SUBC_EMAIL]]":             f.get("subc_email", "").strip(),
        }

    elif tipo in ("obra_1", "fontaneria", "albanileria"):
        items = _parse_items(f, "item")
        if not items:
            return redirect(url_for("index"))
        total_sin_iva = sum(i["importe"] for i in items)
        data = {**common_data, **_build_obra_data(f, total_sin_iva, fecha_larga)}

    elif tipo == "obra_2":
        items_a = _parse_items(f, "item_a")
        items_b = _parse_items(f, "item_b")
        if not items_a and not items_b:
            return redirect(url_for("index"))
        total_a = sum(i["importe"] for i in (items_a or []))
        total_b = sum(i["importe"] for i in (items_b or []))
        total_sin_iva = max(total_a, total_b)
        data = {
            **common_data,
            **_build_obra_data(f, total_sin_iva, fecha_larga),
            **_build_obra_2_extra(f, total_a, total_b),
        }

    else:
        data, total_estimado = _build_service_data(f, tipo, common_data)
        # Partidas adicionales opcionales en modelos de servicio
        extra_items = _parse_items(f, "item")
        if extra_items:
            items = extra_items
            extra_total = sum(i["importe"] for i in extra_items)
            total_estimado += extra_total
            data["[[TOTAL_PARTIDAS_ADICIONALES]]"] = fmt_euro_plain(extra_total)

    # File naming - sanitized to avoid Windows path length / accent issues
    # Formato pedido: "N.o - NOMBRE DE LA CALLE N.o DE LA CALLE" (direccion tal cual)
    num_safe = _safe_filename(num_contrato.replace("/", "-"), maxlen=20)
    label = _safe_filename(obra, maxlen=60)
    stem = f"{num_safe} - {label}" if label else f"{num_safe} - PRESUPUESTO"
    folder_name = stem
    output_dir = SALIDAS_DIR / folder_name
    output_dir.mkdir(parents=True, exist_ok=True)

    generated = {}
    errors = []

    # DOCX
    template_file = TEMPLATES_DIR / TEMPLATE_FILES.get(tipo, "MODELO_MAESTRO_1OPCION.docx")
    if template_file.exists():
        try:
            docx_path = output_dir / f"{stem}.docx"
            DocxGenerator(template_file).generate(
                docx_path, data,
                items=items, items_a=items_a, items_b=items_b,
                subc_items=subc_items,
            )
            titulo_banner = f.get("titulo_servicio", "").strip()
            if titulo_banner:
                _inject_banner_title(docx_path, titulo_banner)
            generated["docx"] = docx_path.relative_to(SALIDAS_DIR).as_posix()
        except Exception as e:
            errors.append(f"DOCX: {e}")
    else:
        errors.append(f"Plantilla no encontrada: {template_file.name}")

    # Excel + HTML for obra types and new trades
    if tipo in ("obra_1", "obra_2", "fontaneria", "albanileria"):
        items_for_excel = items or items_a or []
        try:
            xlsx_path = output_dir / f"{stem}.xlsx"
            ExcelGenerator().generate(xlsx_path, data, items_for_excel)
            generated["xlsx"] = xlsx_path.relative_to(SALIDAS_DIR).as_posix()
        except Exception as e:
            errors.append(f"Excel: {e}")

        try:
            html_path = output_dir / f"{stem}_Visor.html"
            HtmlGenerator().generate(html_path, data, items_for_excel)
            generated["html"] = html_path.relative_to(SALIDAS_DIR).as_posix()
        except Exception as e:
            errors.append(f"HTML: {e}")

    # Anexo fotografico opcional: fotogramas de video + fotos subidas por el usuario.
    # El token viene de /api/subir_evidencia (campo oculto 'evidencia_token').
    evid_token = f.get("evidencia_token", "").strip()
    if (evid_token and "docx" in generated
            and evid_token.startswith("evid_") and evid_token[5:].isalnum()):
        try:
            from core.ai_analyst import _append_photo_annex
            from docx import Document as _DocxDoc
            sess = UPLOADS_DIR / evid_token
            img_ext = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
            fotos = sorted(str(p) for p in sess.glob("foto_*") if p.suffix.lower() in img_ext)
            if fotos:
                _dpath = SALIDAS_DIR / generated["docx"]
                _doc = _DocxDoc(str(_dpath))
                _append_photo_annex(_doc, fotos, titulo="Anexo Fotografico")
                _doc.save(str(_dpath))
        except Exception as e:
            errors.append(f"Anexo fotografico: {e}")

    # PDF
    try:
        docx_for_pdf = SALIDAS_DIR / generated["docx"] if "docx" in generated else None
        html_for_pdf = SALIDAS_DIR / generated["html"] if "html" in generated else None
        pdf_path = PdfConverter().convert(
            docx_for_pdf or output_dir / f"{stem}.docx",
            output_dir,
            html_fallback_path=html_for_pdf,
        )
        if pdf_path:
            generated["pdf"] = pdf_path.relative_to(SALIDAS_DIR).as_posix()
    except Exception as e:
        errors.append(f"PDF: {e}")

    # ── Guardar en disco externo y registrar en Sheets (background, no bloquea) ─
    import threading
    archivos_generados = [
        SALIDAS_DIR / v for v in generated.values()
        if isinstance(v, str) and (SALIDAS_DIR / v).exists()
    ]
    _registro_kwargs = dict(
        tipo=tipo, num_contrato=num_contrato,
        fecha=datetime.now().strftime("%d/%m/%Y"),
        cliente=data.get("[[CLIENTE_NOMBRE]]", ""),
        obra=obra or servicio, importe=total_sin_iva,
        carpeta=folder_name,
    )

    def _background_save():
        try:
            carpeta_red = guardar_en_red(tipo, num_contrato, obra or servicio, archivos_generados)
            if carpeta_red:
                _registro_kwargs["carpeta"] = carpeta_red.name
            sheets_registrar(**_registro_kwargs)
        except Exception as ex:
            import logging
            logging.getLogger(__name__).warning("Background save error: %s", ex)

    threading.Thread(target=_background_save, daemon=True).start()

    return render_template(
        "resultado.html",
        tipo_modelo=tipo,
        num_contrato=num_contrato,
        obra=obra,
        servicio=servicio,
        cliente=data.get("[[CLIENTE_NOMBRE]]", ""),
        fecha=fecha_larga,
        items=items or [],
        items_a=items_a or [],
        items_b=items_b or [],
        total_sin_iva=fmt_euro(total_sin_iva),
        total_con_iva=fmt_euro(total_sin_iva * 1.21),
        total_a=fmt_euro(total_a),
        total_b=fmt_euro(total_b),
        total_estimado=fmt_euro(total_estimado),
        generated=generated,
        errors=errors,
        folder_name=folder_name,
    )


@app.route("/api/estado_disco")
def api_estado_disco():
    return jsonify(disco_estado())


@app.route("/configuracion", methods=["GET", "POST"])
def configuracion():
    from utils.app_config import load as cfg_load, save as cfg_save
    from utils.output_saver import estado_disco

    msg = ""
    if request.method == "POST":
        cfg = cfg_load()
        cfg["output_base_path"] = request.form.get("output_base_path", "").strip()
        cfg["output_subfolder"] = request.form.get("output_subfolder", "").strip()
        cfg_save(cfg)
        msg = "Configuracion guardada."

    cfg = cfg_load()
    disco = estado_disco()
    return render_template("configuracion.html", cfg=cfg, disco=disco, msg=msg)


@app.route("/video_ia/<session_id>/<path:filename>")
def servir_video_ia(session_id: str, filename: str):
    """Sirve (con soporte de range requests, para el <video> del navegador) un
    video CCTV original subido para analisis IA, persistido en VIDEOS_DIR.
    Se usa tanto para revisarlo dentro de la app como desde el visor publico."""
    full = VIDEOS_DIR / session_id / filename
    if not full.exists() or not full.is_file():
        abort(404)
    try:
        full.resolve().relative_to(VIDEOS_DIR.resolve())
    except ValueError:
        abort(403)
    return send_file(full, conditional=True)


@app.route("/api/compartir_informe", methods=["POST"])
def api_compartir_informe():
    """Genera (o reutiliza) un enlace publico de solo-video para un informe
    CCTV ya generado, para poder enviarselo al cliente final."""
    data = request.get_json(force=True) or {}
    session_id = (data.get("session_id") or "").strip()
    if not session_id or "/" in session_id or ".." in session_id:
        return jsonify({"error": "session_id invalido."}), 400
    informe_path = INFORMES_COMPARTIDOS_DIR / f"Informe_WinCam_{session_id}.json"
    if not informe_path.exists():
        return jsonify({"error": "No se encontro ese informe."}), 404
    token = compartidos.crear_o_reusar(session_id)
    return jsonify({"url": url_for("ver_publico", token=token, _external=True)})


@app.route("/ver/<token>")
def ver_publico(token):
    """Visor publico de solo-video (sin cuenta) para el cliente final."""
    session_id = compartidos.resolver(token)
    if not session_id:
        abort(404)
    informe_path = INFORMES_COMPARTIDOS_DIR / f"Informe_WinCam_{session_id}.json"
    if not informe_path.exists():
        abort(404)
    try:
        datos = json.loads(informe_path.read_text(encoding="utf-8"))
    except Exception:
        abort(404)
    for v in datos.get("videos") or []:
        v["url"] = url_for("servir_video_ia", session_id=session_id, filename=v["nombre"])
    return render_template("ver_publico.html", datos=datos)


@app.route("/descargar/<path:rel_path>")
def descargar(rel_path: str):
    full = SALIDAS_DIR / rel_path
    if not full.exists() or not full.is_file():
        abort(404)
    try:
        full.resolve().relative_to(SALIDAS_DIR.resolve())
    except ValueError:
        abort(403)
    return send_file(full, as_attachment=True, download_name=full.name)


# ── Entry point ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import webbrowser, threading, time

    def open_browser():
        time.sleep(1.2)
        webbrowser.open("http://localhost:5000")

    threading.Thread(target=open_browser, daemon=True).start()
    print("\n  Generador de Presupuestos + Analisis IA - Grupo Europa")
    print("  Abriendo en el navegador: http://localhost:5000")
    print("  (Para parar: pulsa Ctrl+C)\n")
    app.run(host="0.0.0.0", port=5000, debug=True, use_reloader=False, threaded=True)
