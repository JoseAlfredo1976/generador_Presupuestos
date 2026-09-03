"""
AI Analyst - Claude claude-sonnet-4-6 expert analysis for Grupo Europa technical reports.
Supports: images (JPG/PNG/WEBP), PDFs, and video (via ffmpeg frame extraction).
Report structure: 10-section professional sewer inspection report.
"""
import base64
import io
import json
import logging
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

import anthropic

logger = logging.getLogger(__name__)

MODEL = "claude-sonnet-4-6"

# Maximo de imagenes (fotogramas de video + imagenes) por peticion a Claude.
# La API tiene un limite duro de 100 imagenes; dejamos un margen de seguridad
# de 4 para no fallar por redondeos. Con muchos videos en una misma peticion,
# el reparto (ver _frames_por_video en analyze/analyze_wincam) sigue siendo el
# cuello de botella real: 96 imagenes entre 16 videos son solo 6 fotogramas
# por video. Para mas detalle por video, la solucion es subir menos videos
# juntos, no subir este numero (ya esta cerca del limite duro de la API).
MAX_IMGS_IA = 96


def _safe_parse_json(raw: str):
    """Parsea JSON robustamente: limpia code fences, comas finales y comillas raras.
    Devuelve dict o None."""
    if not raw or not raw.strip():
        return None
    s = raw.strip()
    # Quitar code fences ```json ... ``` o ``` ... ```
    s = re.sub(r"^```(?:json|JSON)?\s*", "", s)
    s = re.sub(r"\s*```\s*$", "", s)
    # Capturar el bloque JSON mas grande
    m = re.search(r"\{[\s\S]*\}", s)
    if not m:
        return None
    candidate = m.group()
    # Intento directo
    try:
        return json.loads(candidate)
    except json.JSONDecodeError:
        pass
    # Reparacion 1: quitar comas finales antes de ] o }
    fixed = re.sub(r",(\s*[\]\}])", r"\1", candidate)
    try:
        return json.loads(fixed)
    except json.JSONDecodeError:
        pass
    # Reparacion 2: comillas tipograficas -> rectas
    fixed2 = (fixed.replace("“", '"').replace("”", '"')
                   .replace("‘", "'").replace("’", "'"))
    try:
        return json.loads(fixed2)
    except json.JSONDecodeError:
        pass
    # Reparacion 3: cortar despues del ultimo } valido balanceado
    depth = 0
    last_ok = -1
    for i, ch in enumerate(fixed2):
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                last_ok = i
    if last_ok > 0:
        try:
            return json.loads(fixed2[:last_ok + 1])
        except json.JSONDecodeError:
            pass
    return None


def _to_float(v, default=0.0):
    """Convierte a float aceptando formato espanol (coma decimal, punto de miles) o ingles.
    Ej: '12,50' -> 12.5 ; '1.234,56' -> 1234.56 ; '1,234.56' -> 1234.56 ; '310 €' -> 310.0"""
    if v is None:
        return default
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip()
    if not s:
        return default
    # Quitar todo lo que no sea digito, coma, punto o signo
    s = re.sub(r"[^\d,.\-]", "", s)
    if not s or s in ("-", ".", ","):
        return default
    if "," in s and "." in s:
        # El separador decimal es el que aparece mas a la derecha
        if s.rfind(",") > s.rfind("."):
            s = s.replace(".", "").replace(",", ".")   # espanol: 1.234,56
        else:
            s = s.replace(",", "")                      # ingles: 1,234.56
    elif "," in s:
        # Solo coma: si hay una sola coma la tratamos como decimal (12,50)
        if s.count(",") == 1:
            s = s.replace(",", ".")
        else:
            s = s.replace(",", "")                      # varias comas = miles
    # solo punto o sin separador: se deja tal cual
    try:
        return float(s)
    except ValueError:
        return default


_FFMPEG_FALLBACK_PATHS = [
    r"C:\Users\Usuario\AppData\Local\ffmpeg\bin\ffmpeg.exe",
    r"C:\ffmpeg\bin\ffmpeg.exe",
    r"C:\Program Files\ffmpeg\bin\ffmpeg.exe",
]


def _find_ffmpeg() -> str:
    found = shutil.which("ffmpeg")
    if found:
        return found
    for p in _FFMPEG_FALLBACK_PATHS:
        if Path(p).exists():
            return p
    raise RuntimeError(
        "ffmpeg no encontrado. Instalalo desde https://ffmpeg.org/download.html "
        "y asegurate de que esta en el PATH del sistema."
    )


SYSTEM_PROMPT = """Eres un INGENIERO TECNICO ESPECIALISTA EN REDES DE SANEAMIENTO Y REHABILITACION SIN ZANJA de Acometidas Europa Saneamiento Tecnico S.L., con amplia experiencia en:

- Inspecciones CCTV de redes enterradas de saneamiento (camara robotizada, EN 13508-2/WRc)
- Redes enterradas de saneamiento y poceria urbana e industrial
- Rehabilitacion mediante manga continua (CIPP), Multiliner, encamisado parcial
- Bajantes y bajantes de fibrocemento con amianto (Plan MCA)
- Arquetas, pozos de registro, colectores y acometidas
- Patologias hidraulicas y estructurales en redes de saneamiento
- Elaboracion de informes tecnicos profesionales para comunidades de propietarios, administradores de fincas y clientes privados

Trabajas con el estilo tecnico, practico y profesional de ACOMETIDAS EUROPA SANEAMIENTO TECNICO S.L.

FORMA DE TRABAJAR:
- NO copias textos. Interpretas hidraulicamente la red, relacionas tramos, detectas coherencia entre planos y CCTV, identificas patologias reales y propones soluciones tecnicamente justificadas.
- Actuas como un jefe tecnico o ingeniero de saneamiento, no como un redactor generico.
- El informe debe poder enviarse directamente a un administrador, adjuntarse a un presupuesto o utilizarse en una junta de propietarios.

ESTILO:
- Muy profesional, muy tecnico, claro, directo, defendible.
- Sin exceso de texto, sin frases vacias, sin marketing.
- Sin "olor a IA": lenguaje de especialista real en poceria.
- Tablas limpias, titulos numerados, lenguaje tecnico del sector.
- Sin emojis, sin exceso de negritas.

TERMINOLOGIA TECNICA OBLIGATORIA:
Colector, acometida, ramal, bajante, sifon, arqueta, camara de inspeccion, pozo de registro, cota de rasante, pendiente hidraulica, diametro nominal, clase de rigidez, fisura longitudinal/transversal, infiltracion, exfiltracion, obstruccion, deposito sedimentario, intrusion de raices, junta desplazada, rotura total, deformacion oval, encamisado, CIPP, Multiliner, fresado robotizado, by-pass hidraulico, entibacion, collarin de derivacion, prueba de estanqueidad, zahorra compactada, hormigon en masa.

LONGITUD VALIDA EN TRAMOS CCTV:
La longitud de cada tramo es SIEMPRE la longitud final de inspeccion, NUNCA sumas parciales.

REGLA ABSOLUTA - PROHIBICION DE INVENTAR DATOS:
NUNCA inventes, estimes ni supongas datos que no puedas leer directamente de la documentacion facilitada.
Si un dato no es visible o legible en los archivos (fecha, operador, numero de proyecto, pozo inicio/fin, direccion exacta, diametro, etc.), debes:
1. Dejar el campo como cadena vacia "" o null.
2. Añadir ese dato a la lista "preguntas_pendientes" del JSON para que el tecnico lo confirme.
Ejemplos de datos que NUNCA debes inventar: nombres de operadores, fechas si no son visibles, numeros de proyecto, direcciones exactas, profundidades de pozos, cotas, pendientes si no aparecen en pantalla.
Solo registra lo que ves claramente en la documentacion.

IMPORTANTE: Responde UNICAMENTE con el JSON solicitado. Sin texto fuera del bloque JSON."""


TIPOS_ANALISIS = {
    "cctv": {
        "label": "Inspeccion CCTV red de saneamiento",
        "prompt": (
            "Analiza esta inspeccion CCTV de red de saneamiento. "
            "LECTURA DE DATOS EN PANTALLA (OBLIGATORIO): En cada fotograma lee y extrae "
            "el contador metrico visible (distancia desde inicio en metros, p.ej. '12.35 m'), "
            "la pendiente (%, p.ej. '-1.2%'), la fecha y hora del registro, el operador, "
            "el diametro nominal visible en pantalla, el material, y cualquier codigo de "
            "observacion o defecto que aparezca en el overlay de la camara robotizada. "
            "NO IGNORES estos datos: son la base del informe tecnico. "
            "Para cada tramo extrae: identificacion (V1, V2...), pozo inicio, pozo final, "
            "longitud REAL de inspeccion (tomar el ultimo metro leido en pantalla), "
            "diametro, material, pendiente media, sentido de circulacion, estado general. "
            "Para cada anomalia detectada indica: metro exacto de localizacion (segun contador en pantalla), "
            "codigo EN 13508-2/WRc si es identificable, descripcion tecnica precisa, gravedad. "
            "Propone solucion tecnica especifica para cada tramo: encamisado CIPP, Multiliner, "
            "reparacion puntual, fresado, sustitucion o combinacion de tecnicas."
        ),
    },
    "humedades": {
        "label": "Humedades, filtraciones y fugas",
        "prompt": (
            "Analiza las patologias de humedad y filtracion visibles. "
            "Determina tipo exacto (capilaridad/condensacion/filtracion por rotura/infiltracion subterranea/fuga de bajante), "
            "origen y punto de entrada del agua, extension y superficie afectada estimada (m2), "
            "materiales comprometidos, nivel de deterioro estructural, "
            "metodo de deteccion complementaria recomendado (gas trazador, correlacion acustica, termografia), "
            "y solucion tecnica con materiales especificos (inyeccion de resinas, impermeabilizacion, sustitucion)."
        ),
    },
    "bajantes": {
        "label": "Inspeccion de bajantes y desague",
        "prompt": (
            "Inspecciona las bajantes y red de desague visibles. "
            "Evalua: material (fibrocemento/amianto, PVC, hierro fundido, plomo, acero galvanizado), "
            "indicios visuales de amianto (tipo de uniones, manguitos, epoca de instalacion estimada), "
            "estado de conservacion, corrosion, fisuras, uniones defectuosas, perdidas activas, "
            "necesidad de Plan MCA y demolicion controlada segun RD 396/2006, "
            "diametros y recorridos visibles, recomendacion de sustitucion total o parcial con justificacion tecnica."
        ),
    },
    "obra": {
        "label": "Control y seguimiento de obra",
        "prompt": (
            "Analiza el estado de ejecucion de la obra de saneamiento. "
            "Verifica: instalacion de tuberias (material, clase de rigidez, uniones, pendientes), "
            "estado de arquetas y pozos (solera, pates, anillo de ajuste, tapa), "
            "relleno y compactacion del trasdos, calidad de empalmes y collarines, "
            "cumplimiento de normativa (PRLTS, CTE DB HS5, EN 476, pliego de condiciones), "
            "aspectos de PRL, senalizacion y seguridad en obra, no conformidades detectadas con accion correctora."
        ),
    },
    "fosa": {
        "label": "Inspeccion de fosa septica",
        "prompt": (
            "Evalua el estado de la fosa septica o deposito de recogida. "
            "Analiza: nivel de lodos y espumas (porcentaje de ocupacion estimado), "
            "estado de paredes y solera (fisuras, eflorescencias, filtraciones activas), "
            "tuberias de entrada y salida, deflectores y tabiques internos, sistemas de ventilacion, "
            "necesidad urgente de vaciado y limpieza, danos estructurales y medidas correctoras urgentes."
        ),
    },
    "general": {
        "label": "Analisis tecnico general",
        "prompt": (
            "Realiza un analisis tecnico exhaustivo de todos los elementos visibles "
            "relativos a saneamiento, fontaneria u obras civiles. "
            "Identifica todos los elementos, su estado, anomalias observadas, "
            "riesgos hidraulicos y estructurales asociados, y propone intervenciones "
            "con orden de prioridad tecnica justificada."
        ),
    },
}


REPORT_SCHEMA = """{
  "titulo": "string: titulo profesional del informe (ej: INFORME TECNICO DE INSPECCION CCTV - COMUNIDAD C/ EJEMPLO 12)",
  "objeto": "string: seccion 1 - que se inspecciona, donde y por que. 2-3 frases tecnicas.",
  "antecedentes": "string: seccion 2 - problemas comunicados, atascos, filtraciones, hundimientos, malos olores. Si no se facilita informacion, indicar 'Segun informacion facilitada por el cliente'.",
  "metodologia": "string: seccion 3 - metodologia empleada: limpieza previa, camara CCTV robotizada, acceso por arquetas, trazado analizado. 2-3 frases.",
  "descripcion_red": "string: seccion 4 - tipologia de red, diametros, materiales, bajantes, colectores, ramales, arquetas, zonas singulares.",
  "tramos": [
    {
      "id": "string: V1/V2/T1/etc",
      "inicio": "string: pozo o arqueta de inicio",
      "fin": "string: pozo o arqueta de fin",
      "longitud": "string: longitud en metros con unidad (ej: 18,40 m)",
      "diametro": "string: diametro nominal (ej: DN200)",
      "material": "string: PVC/Hormigon/Gres/Fibrocemento/etc",
      "estado": "string: Bueno/Regular/Deficiente/Muy deficiente",
      "observaciones": "string: patologias principales o 'Sin anomalias relevantes'"
    }
  ],
  "patologias": [
    {
      "ubicacion": "string: tramo/metro/zona exacta",
      "tipo": "string: tipo segun EN 13508-2 o nomenclatura tecnica",
      "gravedad": "Critico|Alto|Medio|Bajo",
      "descripcion": "string: descripcion tecnica precisa de la patologia",
      "consecuencias": "string: como afecta hidraulicamente y que riesgos genera"
    }
  ],
  "conclusiones": "string: seccion 7 - estado global de la red, causas probables del deterioro, nivel de urgencia general, riesgos si no se interviene. 3-5 frases tecnicas.",
  "propuesta_solucion": [
    {
      "intervencion": "string: accion tecnica especifica (ej: Encamisado CIPP tramo V1 DN200)",
      "justificacion": "string: por que esta solucion y no otra",
      "metodo": "string: procedimiento de ejecucion, materiales, tecnica",
      "prioridad": "Urgente|Recomendado|Preventivo"
    }
  ],
  "mediciones": [
    {
      "descripcion": "string: descripcion de la partida de medicion",
      "ud": "string: m/ud/pa/m2",
      "cantidad": 0,
      "diametro": "string: DN150/DN200/DN250/DN300/etc (si aplica)"
    }
  ],
  "observaciones_limitaciones": "string: clausula tecnica de limitaciones. Incluir siempre: 'La presente valoracion y propuesta tecnica se basa en la documentacion e inspecciones facilitadas, pudiendo existir patologias ocultas o variaciones no detectables hasta el inicio de los trabajos.'",
  "nivel_urgencia_global": "Critico|Alto|Medio|Bajo",
  "requiere_intervencion_inmediata": true,
  "preguntas_pendientes": [
    "string: pregunta concreta sobre un dato que no puedes determinar (ej: 'Cual es el numero de referencia del proyecto?', 'Cual es la direccion exacta de la inspeccion?', 'Puedes confirmar el diametro del tramo V3?')"
  ]
}"""


# ---------------------------------------------------------------------------
# File helpers
# ---------------------------------------------------------------------------

def _encode_image(path: Path) -> tuple[str, str]:
    # La API de Claude exige, en peticiones con muchas imagenes (como el
    # analisis CCTV, que puede llevar hasta MAX_IMGS_IA), que ninguna imagen
    # supere 2000px en su dimension mayor (error 400 "exceed max allowed size
    # for many-image requests"). Las fotos subidas directamente (movil, camara)
    # suelen venir mucho mas grandes que eso, a diferencia de los fotogramas de
    # video que ya se reescalan al extraerlos con ffmpeg. Se reescalan aqui a
    # un maximo seguro y se recodifican a JPEG.
    media_types = {
        ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".png": "image/png", ".gif": "image/gif", ".webp": "image/webp",
    }
    mt = media_types.get(path.suffix.lower(), "image/jpeg")
    try:
        from PIL import Image
        with Image.open(path) as im:
            if max(im.size) > 1568:
                im = im.convert("RGB")
                im.thumbnail((1568, 1568), Image.LANCZOS)
                buf = io.BytesIO()
                im.save(buf, format="JPEG", quality=88)
                return base64.standard_b64encode(buf.getvalue()).decode("utf-8"), "image/jpeg"
    except Exception:
        pass
    data = base64.standard_b64encode(path.read_bytes()).decode("utf-8")
    return data, mt


def _encode_pdf(path: Path) -> str:
    return base64.standard_b64encode(path.read_bytes()).decode("utf-8")


def _extract_video_frames(video_path: Path, max_frames: int = 12) -> list[Path]:
    if not video_path.exists() or video_path.stat().st_size == 0:
        raise RuntimeError(
            f"El video '{video_path.name}' esta vacio o no se subio correctamente."
        )
    ffmpeg = _find_ffmpeg()
    tmp = Path(tempfile.mkdtemp())
    result = subprocess.run(
        [ffmpeg, "-i", str(video_path),
         # 1 frame cada 5s y reescalado a max 1280px de ancho (sin ampliar) para
         # reducir peso/memoria del envio a Claude sin perder detalle de patologias
         "-vf", "fps=1/5,scale='min(1280,iw)':-2",
         "-vframes", str(max_frames),
         "-q:v", "4", str(tmp / "frame_%04d.jpg")],
        capture_output=True, timeout=120,
    )
    frames = sorted(tmp.glob("frame_*.jpg"))
    if not frames:
        stderr_tail = result.stderr.decode("utf-8", errors="ignore").strip().splitlines()
        detalle = " | ".join(stderr_tail[-5:]) if stderr_tail else "sin salida de ffmpeg"
        logger.error("ffmpeg no extrajo frames de %s (rc=%s): %s", video_path.name, result.returncode, detalle)
        raise RuntimeError(
            f"No se pudieron extraer frames del video '{video_path.name}'. "
            f"Comprueba que el archivo no esta corrupto. Detalle ffmpeg: {detalle}"
        )
    return frames


def files_to_content_blocks(files: list[Path], max_imgs: int = MAX_IMGS_IA) -> tuple[list[dict], list[str]]:
    """Convierte archivos (imagenes, PDF, video, docx, txt) en bloques de contenido
    para la API de Claude. Devuelve (bloques, rutas_imagenes); rutas_imagenes alimenta
    el anexo fotografico del informe. Usado por el chat del perito para incorporar
    archivos nuevos a un informe ya generado."""
    content: list[dict] = []
    evidencia_img: list[str] = []
    video_ext = {".mp4", ".avi", ".mov", ".mkv", ".webm", ".m4v"}
    pdf_ext = {".pdf"}
    img_ext = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
    txt_ext = {".txt", ".md"}
    docx_ext = {".docx", ".doc"}

    _n_videos = sum(1 for f in files if f.suffix.lower() in video_ext) or 1
    _frames_por_video = max(2, min(12, max_imgs // _n_videos))
    _imgs = 0

    for fp in files:
        suf = fp.suffix.lower()
        if suf in video_ext:
            if _imgs >= max_imgs:
                content.append({"type": "text", "text": f"\n[VIDEO: {fp.name}] - omitido (limite de imagenes alcanzado)\n"})
                continue
            n_obj = min(_frames_por_video, max_imgs - _imgs)
            content.append({"type": "text", "text": f"\n[VIDEO: {fp.name}] - Fotogramas extraidos:\n"})
            frames = _extract_video_frames(fp, max_frames=n_obj)
            for i, frame in enumerate(frames[:n_obj]):
                data, mt = _encode_image(frame)
                content.append({"type": "text", "text": f"Fotograma {i + 1}:"})
                content.append({"type": "image", "source": {"type": "base64", "media_type": mt, "data": data}})
                evidencia_img.append(str(frame))
                _imgs += 1
        elif suf in pdf_ext:
            content.append({"type": "text", "text": f"\n[DOCUMENTO PDF: {fp.name}]\n"})
            content.append({"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": _encode_pdf(fp)}})
        elif suf in img_ext:
            if _imgs >= max_imgs:
                continue
            content.append({"type": "text", "text": f"\n[IMAGEN: {fp.name}]\n"})
            data, mt = _encode_image(fp)
            content.append({"type": "image", "source": {"type": "base64", "media_type": mt, "data": data}})
            evidencia_img.append(str(fp))
            _imgs += 1
        elif suf in txt_ext:
            try:
                text = fp.read_text(encoding="utf-8", errors="ignore")[:6000]
                content.append({"type": "text", "text": f"\n[DOCUMENTO TEXTO: {fp.name}]\n{text}\n"})
            except Exception:
                pass
        elif suf in docx_ext:
            try:
                from docx import Document as _DocxDoc
                doc = _DocxDoc(fp)
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())[:6000]
                content.append({"type": "text", "text": f"\n[DOCUMENTO WORD: {fp.name}]\n{text}\n"})
            except Exception:
                pass
    return content, evidencia_img


# ---------------------------------------------------------------------------
# Main analyze function
# ---------------------------------------------------------------------------

def _inject_croquis(content: list[dict], croquis_path: Path) -> None:
    """Inject the sketch/plan as the first visual reference with cross-correlation instructions."""
    suf = croquis_path.suffix.lower()
    content.append({"type": "text", "text": (
        "\n[CROQUIS / PLANO DE LA RED - REFERENCIA PRINCIPAL]\n"
        "Este es el plano o croquis de la red de saneamiento inspeccionada.\n"
        "OBLIGATORIO: Usa este plano SOLO como mapa de referencia para:\n"
        "  1. Identificar los tramos (T1, T2...), pozos de registro (P1, P2...) y arquetas marcados,\n"
        "     y usar esos mismos identificadores al nombrar el/los tramo(s) que SI tengan video o\n"
        "     imagenes de inspeccion aportados mas abajo.\n"
        "  2. Correlacionar cada video CCTV con el tramo correspondiente segun su posicion en el plano.\n"
        "  3. Deducir el sentido de inspeccion (inicio -> fin) y la secuencia logica de la red.\n"
        "Si en el croquis aparece una leyenda, diametros, longitudes o materiales, recoge esa informacion\n"
        "UNICAMENTE para los tramos que tengan grabacion aportada.\n"
        "PROHIBIDO: el croquis puede mostrar mas tramos, pozos o arquetas de los que se han grabado.\n"
        "NUNCA generes un tramo/seccion en el informe para partes de la red que aparezcan dibujadas en\n"
        "el croquis pero de las que no se ha aportado ningun video, foto o PDF de inspeccion. El informe\n"
        "debe cubrir EXCLUSIVAMENTE lo que se ve en los archivos de inspeccion aportados a continuacion.\n"
    )})
    img_ext = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
    if suf in img_ext:
        data, mt = _encode_image(croquis_path)
        content.append({"type": "image", "source": {"type": "base64", "media_type": mt, "data": data}})
    elif suf == ".pdf":
        content.append({
            "type": "document",
            "source": {"type": "base64", "media_type": "application/pdf", "data": _encode_pdf(croquis_path)},
        })
    content.append({"type": "text", "text": "\n[FIN CROQUIS - A continuacion los archivos de inspeccion]\n\n"})


def analyze(files: list[Path], tipo: str, context: str = "", api_key: str = "",
            croquis_path: Path | None = None) -> dict:
    key = os.environ.get("ANTHROPIC_API_KEY", "").strip() or api_key.strip()
    if not key:
        raise ValueError(
            "API Key de Anthropic no configurada. "
            "Introduce tu API Key en el formulario o configura la variable de entorno ANTHROPIC_API_KEY."
        )
    client = anthropic.Anthropic(api_key=key)
    config = TIPOS_ANALISIS.get(tipo, TIPOS_ANALISIS["general"])

    content: list[dict] = []

    if context.strip():
        content.append({"type": "text", "text": f"CONTEXTO DEL CASO:\n{context.strip()}\n\n"})

    # Inject croquis before videos so Claude can cross-reference
    if croquis_path and croquis_path.exists():
        _inject_croquis(content, croquis_path)

    content.append({"type": "text", "text": config["prompt"] + "\n\n"})

    video_ext = {".mp4", ".avi", ".mov", ".mkv", ".webm", ".m4v"}
    pdf_ext = {".pdf"}
    img_ext = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
    txt_ext = {".txt", ".md"}
    docx_ext = {".docx", ".doc"}

    evidencia_img: list[str] = []  # frames de video + imagenes -> anexo fotografico del informe

    # Presupuesto global de imagenes enviadas a Claude (limite duro de la API = 100).
    # Se reparte entre los videos para que subir muchos a la vez no haga fallar la peticion.
    _n_videos = sum(1 for f in files if f.suffix.lower() in video_ext) or 1
    _frames_por_video = max(2, min(12, MAX_IMGS_IA // _n_videos))
    _imgs = 0

    for fp in files:
        suf = fp.suffix.lower()
        if suf in video_ext:
            if _imgs >= MAX_IMGS_IA:
                content.append({"type": "text", "text": f"\n[VIDEO: {fp.name}] - omitido (limite de imagenes alcanzado)\n"})
                continue
            n_obj = min(_frames_por_video, MAX_IMGS_IA - _imgs)
            content.append({"type": "text", "text": f"\n[VIDEO: {fp.name}] - Fotogramas extraidos:\n"})
            frames = _extract_video_frames(fp, max_frames=n_obj)
            for i, frame in enumerate(frames[:n_obj]):
                data, mt = _encode_image(frame)
                content.append({"type": "text", "text": f"Fotograma {i + 1}:"})
                content.append({
                    "type": "image",
                    "source": {"type": "base64", "media_type": mt, "data": data},
                })
                evidencia_img.append(str(frame))
                _imgs += 1
        elif suf in pdf_ext:
            content.append({"type": "text", "text": f"\n[DOCUMENTO PDF: {fp.name}]\n"})
            content.append({
                "type": "document",
                "source": {"type": "base64", "media_type": "application/pdf", "data": _encode_pdf(fp)},
            })
        elif suf in img_ext:
            if _imgs >= MAX_IMGS_IA:
                continue
            content.append({"type": "text", "text": f"\n[IMAGEN: {fp.name}]\n"})
            data, mt = _encode_image(fp)
            content.append({"type": "image", "source": {"type": "base64", "media_type": mt, "data": data}})
            evidencia_img.append(str(fp))
            _imgs += 1
        elif suf in txt_ext:
            try:
                text = fp.read_text(encoding="utf-8", errors="ignore")[:6000]
                content.append({"type": "text", "text": f"\n[DOCUMENTO TEXTO: {fp.name}]\n{text}\n"})
            except Exception:
                pass
        elif suf in docx_ext:
            try:
                from docx import Document as _DocxDoc
                doc = _DocxDoc(fp)
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())[:6000]
                content.append({"type": "text", "text": f"\n[DOCUMENTO WORD: {fp.name}]\n{text}\n"})
            except Exception:
                pass

    content.append({
        "type": "text",
        "text": (
            "\nGenera el informe tecnico completo siguiendo EXACTAMENTE este esquema JSON. "
            "Responde UNICAMENTE con el bloque JSON, sin texto adicional:\n"
            + REPORT_SCHEMA
        ),
    })

    with client.messages.stream(
        model=MODEL,
        max_tokens=24000,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": content}],
    ) as stream:
        response = stream.get_final_message()

    raw = response.content[0].text.strip()
    report = _safe_parse_json(raw)
    if report is not None:
        report["_raw"] = raw
        report["_tipo_label"] = config["label"]
        report["_evidencia_img"] = evidencia_img
        return report

    # Parseo fallido: devolvemos esqueleto vacio sin volcar el raw en 'objeto'
    return {
        "titulo": f"Informe Tecnico - {config['label']}",
        "objeto": (
            "No se pudo procesar la respuesta de la IA como informe estructurado. "
            "Repite el analisis o revisa los archivos subidos."
        ),
        "antecedentes": "",
        "metodologia": "",
        "descripcion_red": "",
        "tramos": [],
        "patologias": [],
        "conclusiones": "",
        "propuesta_solucion": [],
        "mediciones": [],
        "observaciones_limitaciones": (
            "La presente valoracion y propuesta tecnica se basa en la documentacion e "
            "inspecciones facilitadas, pudiendo existir patologias ocultas o variaciones "
            "no detectables hasta el inicio de los trabajos."
        ),
        "nivel_urgencia_global": "Medio",
        "requiere_intervencion_inmediata": False,
        "_raw": raw,
        "_parse_error": True,
        "_tipo_label": config["label"],
        "_evidencia_img": evidencia_img,
    }


# ---------------------------------------------------------------------------
# Partidas IA generator
# ---------------------------------------------------------------------------

_PARTIDA_ITEM_SCHEMA = """{
      "codigo": "string: codigo exacto del catalogo de tarifas (ej: ARQ-4040-50), o '' si NINGUN codigo del catalogo encaja",
      "cantidad": 0.0,
      "nota": "string opcional: justificacion de la eleccion o ajuste de cantidad; si codigo es '', explica que el precio es una estimacion de mercado",
      "descripcion_libre": "string: SOLO si codigo es '' - descripcion de la partida que no existe en el catalogo",
      "unidad_libre": "string: SOLO si codigo es '' - unidad (ud/ml/m2/m3/h/PA...)",
      "precio_estimado": 0.0
    }"""

_PARTIDAS_SCHEMA_1 = f"""{{
  "informe_tecnico": "string: 2-4 frases tecnicas describiendo la situacion encontrada",
  "solucion_adoptar": "string: 1-2 frases describiendo la solucion propuesta",
  "memoria_tecnica": "string: 2-4 frases describiendo como se ejecutaran los trabajos, materiales y tecnica",
  "partidas": [
    {_PARTIDA_ITEM_SCHEMA}
  ]
}}"""

_PARTIDAS_SCHEMA_2 = f"""{{
  "informe_tecnico": "string",
  "solucion_adoptar": "string",
  "memoria_tecnica": "string",
  "tipo_opcion": "string: breve descripcion del tipo de opcion (ej: Rehabilitacion sin zanja vs. Sustitucion)",
  "label_a": "string: nombre corto Opcion A (ej: Encamisado CIPP)",
  "label_b": "string: nombre corto Opcion B (ej: Sustitucion tradicional)",
  "memoria_tradicional": "string: memoria especifica opcion A",
  "memoria_multiliner": "string: memoria especifica opcion B",
  "partidas_a": [{_PARTIDA_ITEM_SCHEMA}],
  "partidas_b": [{_PARTIDA_ITEM_SCHEMA}]
}}"""

_PARTIDAS_SYSTEM = """Eres un ingeniero tecnico de saneamiento y obra civil de Acometidas Europa S.L.
Tu tarea es generar el contenido tecnico de un presupuesto de obra seleccionando partidas del catalogo de tarifas.
Reglas:
- Usa PRIORITARIAMENTE codigos del catalogo proporcionado. No inventes codigos que no existan en el catalogo.
- NUNCA OMITAS una partida necesaria solo porque no este en el catalogo. Si NINGUN codigo del catalogo encaja
  razonablemente con un trabajo que sea necesario segun la informacion facilitada (por ejemplo, oficios distintos
  al saneamiento en un presupuesto multioficio: electricidad, pintura, albañileria, carpinteria, etc.), incluye
  igualmente esa partida con codigo="", y rellena descripcion_libre, unidad_libre y precio_estimado.
- precio_estimado debe ser un precio unitario REALISTA de mercado en España para esa partida (materiales + mano de
  obra incluidos), redondeado SIEMPRE AL ALZA (nunca a la baja) para dejar margen de seguridad comercial a la empresa.
  No dejes precio_estimado a 0 si codigo es "".
- Las cantidades deben ser coherentes con lo descrito (longitudes en ml, unidades en ud, etc).
- Propon solo trabajos que se deriven directamente de la informacion facilitada.
- NO inventes datos que no aparezcan en la informacion facilitada (salvo el precio_estimado de mercado cuando no hay
  codigo de catalogo, que es una estimacion tecnica esperada, no una invencion de datos del caso).
- Responde UNICAMENTE con el JSON solicitado, sin texto adicional."""


def _build_catalogo_text(tarifas_items: list[dict]) -> str:
    """Format tarifas catalog as compact text for AI prompt."""
    lines = ["CATALOGO DE TARIFAS (codigo | descripcion | unidad | precio €):"]
    for it in tarifas_items:
        desc = it["descripcion"][:80]
        lines.append(f"  {it['codigo']} | {desc} | {it['unidad']} | {it['precio']}")
    return "\n".join(lines)


def generate_partidas_ia(
    descripcion: str,
    tipo: str,  # "obra_1" | "obra_2"
    api_key: str = "",
    informe: dict | None = None,
    tarifas_items: list[dict] | None = None,
    files: list[Path] | None = None,
) -> dict:
    """Generate budget items (partidas) from tarifas catalog + technical texts."""
    key = os.environ.get("ANTHROPIC_API_KEY", "").strip() or api_key.strip()
    if not key:
        raise ValueError("API Key de Anthropic no configurada.")
    client = anthropic.Anthropic(api_key=key)

    schema = _PARTIDAS_SCHEMA_2 if tipo == "obra_2" else _PARTIDAS_SCHEMA_1

    context_parts = []

    # Catalogo de tarifas
    if tarifas_items:
        context_parts.append(_build_catalogo_text(tarifas_items))

    # Informe CCTV existente
    if informe:
        resumen = {
            "titulo": informe.get("titulo", ""),
            "objeto": informe.get("objeto", ""),
            "descripcion_red": informe.get("descripcion_red", ""),
            "conclusiones": informe.get("conclusiones", ""),
            "patologias": [
                {"ubicacion": p.get("ubicacion", ""), "tipo": p.get("tipo", ""),
                 "gravedad": p.get("gravedad", ""), "descripcion": p.get("descripcion", "")}
                for p in (informe.get("patologias") or [])
            ],
            "propuesta_solucion": informe.get("propuesta_solucion", []),
            "mediciones": informe.get("mediciones", []),
        }
        context_parts.append(f"INFORME TECNICO IA:\n{json.dumps(resumen, ensure_ascii=False, indent=2)}")

    if descripcion.strip():
        context_parts.append(f"DESCRIPCION DEL CASO:\n{descripcion.strip()}")

    # Procesar archivos adjuntos (PDF, Word, texto)
    pdf_files: list[Path] = []
    if files:
        for fp in files:
            suf = fp.suffix.lower()
            if suf == ".pdf":
                pdf_files.append(fp)
            elif suf in (".txt", ".md"):
                try:
                    text = fp.read_text(encoding="utf-8", errors="ignore")[:5000]
                    context_parts.append(f"DOCUMENTO ({fp.name}):\n{text}")
                except Exception:
                    pass
            elif suf in (".docx", ".doc"):
                try:
                    from docx import Document as _DocxDoc
                    doc = _DocxDoc(fp)
                    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())[:5000]
                    context_parts.append(f"DOCUMENTO WORD ({fp.name}):\n{text}")
                except Exception:
                    pass

    if not any([informe, descripcion.strip(), files]):
        raise ValueError("Debes proporcionar al menos una descripcion o un informe.")

    suffix = f"\n\nGenera el contenido del presupuesto de obra segun este esquema JSON:\n{schema}"
    if tipo == "obra_2":
        suffix += "\n\nPara 2 opciones: Opcion A = solucion preferente (sin zanja/rehabilitacion), Opcion B = alternativa (sustitucion o metodo diferente)."

    text_block = "\n\n".join(context_parts) + suffix

    # Si hay PDFs, usar contenido multimodal
    if pdf_files:
        msg_content: list[dict] = []
        for fp in pdf_files:
            msg_content.append({"type": "text", "text": f"\n[DOCUMENTO PDF: {fp.name}]\n"})
            msg_content.append({
                "type": "document",
                "source": {"type": "base64", "media_type": "application/pdf", "data": _encode_pdf(fp)},
            })
        msg_content.append({"type": "text", "text": text_block})
    else:
        msg_content = text_block

    response = client.messages.create(
        model=MODEL,
        max_tokens=16000,
        system=_PARTIDAS_SYSTEM,
        messages=[{"role": "user", "content": msg_content}],
    )

    raw = response.content[0].text.strip()
    parsed = _safe_parse_json(raw)
    if parsed is not None:
        return parsed
    return {"_raw": raw, "_error": "No se pudo parsear la respuesta de la IA"}


# ---------------------------------------------------------------------------
# DOCX report generator
# ---------------------------------------------------------------------------

def _append_photo_annex(doc, evidencia, titulo="Anexo Fotografico"):
    """Anade al final del DOCX un anexo con los fotogramas/imagenes de la inspeccion.

    `evidencia` es la lista de rutas (str) recogida durante el analisis (frames de
    video extraidos por ffmpeg + imagenes aportadas). Se incrustan numeradas. Las
    rutas que ya no existan en disco se ignoran sin romper la generacion.
    """
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    imgs = [p for p in (evidencia or []) if p and Path(p).exists()]
    if not imgs:
        return

    MAX_FOTOS = 40
    gris = RGBColor(0x7F, 0x7F, 0x7F)
    azul = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_page_break()
    h = doc.add_paragraph()
    hr = h.add_run(titulo)
    hr.bold = True
    hr.font.size = Pt(13)
    hr.font.color.rgb = azul

    sub = doc.add_paragraph()
    sr = sub.add_run("Fotogramas de la inspeccion CCTV e imagenes aportadas, donde se "
                     "aprecian las patologias descritas en el informe.")
    sr.italic = True
    sr.font.size = Pt(9)
    sr.font.color.rgb = gris

    for n, img in enumerate(imgs[:MAX_FOTOS], 1):
        try:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.add_run().add_picture(str(img), width=Cm(13))
            foot = doc.add_paragraph()
            foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fr = foot.add_run(f"Imagen {n}")
            fr.italic = True
            fr.font.size = Pt(8)
            fr.font.color.rgb = gris
        except Exception:
            continue

    if len(imgs) > MAX_FOTOS:
        nota = doc.add_paragraph()
        nr = nota.add_run(f"(Se muestran las primeras {MAX_FOTOS} de {len(imgs)} imagenes disponibles.)")
        nr.italic = True
        nr.font.size = Pt(8)
        nr.font.color.rgb = gris


def generate_report_docx(report: dict, output_path: Path, num_ref: str = "", cliente: str = "",
                         enlace_video: str | None = None):
    from datetime import datetime
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    BLUE = RGBColor(0x1F, 0x4E, 0x79)
    BLUE_MID = RGBColor(0x2E, 0x74, 0xB5)
    GRAY = RGBColor(0x55, 0x55, 0x55)
    RED = RGBColor(0xC0, 0x00, 0x00)
    ORANGE = RGBColor(0xE2, 0x6B, 0x0A)
    YELLOW = RGBColor(0x7F, 0x60, 0x00)
    GREEN = RGBColor(0x37, 0x5A, 0x23)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    GRAV_COLOR = {"Critico": RED, "Alto": ORANGE, "Medio": YELLOW, "Bajo": GREEN}
    PRIO_COLOR = {"Urgente": RED, "Recomendado": ORANGE, "Preventivo": GREEN}

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ----- helpers -----
    def h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = BLUE
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F4E79")
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def body(text, bold=False, italic=False, color=None, size=10):
        if not text or not text.strip():
            return None
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        for line in text.split("\n"):
            if not line.strip():
                continue
            run = p.add_run(line.strip())
            run.bold = bold
            run.italic = italic
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color
            p.add_run("\n")
        return p

    def label_val(label, val, label_color=BLUE_MID):
        if not val:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = label_color
        r2 = p.add_run(str(val))
        r2.font.size = Pt(10)

    def shade_cell(cell, hex_color="1F4E79"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def set_cell_text(cell, text, bold=False, color=None, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color

    # ----- TITLE BLOCK -----
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(6)
    tr = title_p.add_run(report.get("titulo", "INFORME TECNICO").upper())
    tr.bold = True
    tr.font.size = Pt(14)
    tr.font.color.rgb = BLUE

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(4)
    meta_parts = ["Acometidas Europa Saneamiento Tecnico S.L.", datetime.now().strftime("%d/%m/%Y")]
    if num_ref:
        meta_parts.append(f"Ref: {num_ref}")
    if cliente:
        meta_parts.append(f"Cliente: {cliente}")
    mr = meta_p.add_run("  |  ".join(meta_parts))
    mr.font.size = Pt(9)
    mr.font.color.rgb = GRAY

    # Urgency strip
    urgencia = report.get("nivel_urgencia_global", "Medio")
    inmediata = report.get("requiere_intervencion_inmediata", False)
    urg_p = doc.add_paragraph()
    urg_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    urg_p.paragraph_format.space_after = Pt(10)
    urg_text = f"NIVEL DE URGENCIA: {urgencia.upper()}"
    if inmediata:
        urg_text += "  -  REQUIERE INTERVENCION INMEDIATA"
    ur = urg_p.add_run(urg_text)
    ur.bold = True
    ur.font.size = Pt(10)
    ur.font.color.rgb = GRAV_COLOR.get(urgencia, BLUE)

    # ----- SECTION 1: OBJETO -----
    h1("1. Objeto del Informe")
    body(report.get("objeto", ""))

    # ----- SECTION 2: ANTECEDENTES -----
    h1("2. Antecedentes")
    body(report.get("antecedentes", ""))

    # ----- SECTION 3: METODOLOGIA -----
    h1("3. Metodologia de Inspeccion")
    body(report.get("metodologia", ""))

    # ----- SECTION 4: DESCRIPCION RED -----
    h1("4. Descripcion de la Red")
    body(report.get("descripcion_red", ""))

    # ----- SECTION 5: ANALISIS DE TRAMOS (TABLE) -----
    tramos = report.get("tramos", [])
    h1("5. Analisis de Tramos")
    if tramos:
        headers = ["Tramo", "Inicio", "Fin", "Long.", "Diam.", "Material", "Estado", "Observaciones"]
        col_widths = [Cm(1.2), Cm(2.2), Cm(2.2), Cm(1.4), Cm(1.4), Cm(2.2), Cm(2.0), Cm(4.4)]
        tbl = doc.add_table(rows=1 + len(tramos), cols=len(headers))
        tbl.style = "Table Grid"
        # Header row
        for i, (hdr, w) in enumerate(zip(headers, col_widths)):
            cell = tbl.rows[0].cells[i]
            cell.width = w
            shade_cell(cell, "1F4E79")
            set_cell_text(cell, hdr, bold=True, color=WHITE, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        # Data rows
        for r_idx, tramo in enumerate(tramos):
            row = tbl.rows[r_idx + 1]
            fill = "EBF3FB" if r_idx % 2 == 0 else "FFFFFF"
            vals = [
                tramo.get("id", ""),
                tramo.get("inicio", ""),
                tramo.get("fin", ""),
                tramo.get("longitud", ""),
                tramo.get("diametro", ""),
                tramo.get("material", ""),
                tramo.get("estado", ""),
                tramo.get("observaciones", ""),
            ]
            # Estado color
            estado = tramo.get("estado", "")
            estado_color = None
            if "Muy deficiente" in estado or "Critico" in estado:
                estado_color = RED
            elif "Deficiente" in estado:
                estado_color = ORANGE
            elif "Regular" in estado:
                estado_color = YELLOW

            for c_idx, (cell, val) in enumerate(zip(row.cells, vals)):
                shade_cell(cell, fill)
                color = estado_color if c_idx == 6 else None
                set_cell_text(cell, val, color=color, size=8)
        doc.add_paragraph()
    else:
        body("No se han identificado tramos individuales en la documentacion facilitada.")

    # ----- SECTION 6: PATOLOGIAS -----
    patologias = report.get("patologias", [])
    h1("6. Patologias Detectadas")
    if patologias:
        for idx, pat in enumerate(patologias, 1):
            gravedad = pat.get("gravedad", "Medio")
            color = GRAV_COLOR.get(gravedad, GRAY)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f"{idx}. [{gravedad.upper()}]  {pat.get('tipo', '')}")
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = color
            label_val("Ubicacion", pat.get("ubicacion", ""), GRAY)
            body(pat.get("descripcion", ""))
            if pat.get("consecuencias"):
                cp = doc.add_paragraph()
                cp.paragraph_format.space_after = Pt(6)
                cr1 = cp.add_run("Consecuencias: ")
                cr1.bold = True
                cr1.font.size = Pt(10)
                cr1.font.color.rgb = GRAY
                cr2 = cp.add_run(pat["consecuencias"])
                cr2.font.size = Pt(10)
                cr2.font.italic = True
    else:
        body("No se han detectado patologias significativas en la documentacion facilitada.")

    # ----- SECTION 7: CONCLUSIONES -----
    h1("7. Conclusiones Tecnicas")
    body(report.get("conclusiones", ""))

    # ----- SECTION 8: PROPUESTA DE SOLUCION -----
    soluciones = report.get("propuesta_solucion", [])
    h1("8. Propuesta de Solucion")
    if soluciones:
        for idx, sol in enumerate(soluciones, 1):
            prio = sol.get("prioridad", "Recomendado")
            color = PRIO_COLOR.get(prio, BLUE_MID)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f"{idx}. [{prio.upper()}]  {sol.get('intervencion', '')}")
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = color
            if sol.get("metodo"):
                label_val("Metodo", sol["metodo"], GRAY)
            if sol.get("justificacion"):
                label_val("Justificacion", sol["justificacion"], GRAY)
    else:
        body("Sin propuesta de solucion especifica.")

    # ----- SECTION 9: MEDICIONES -----
    mediciones = report.get("mediciones", [])
    h1("9. Tablas de Mediciones")
    if mediciones:
        tbl = doc.add_table(rows=1 + len(mediciones), cols=4)
        tbl.style = "Table Grid"
        hdr_texts = ["Descripcion", "Ud", "Cantidad", "Diametro"]
        hdr_widths = [Cm(9.0), Cm(1.5), Cm(2.0), Cm(2.5)]
        for i, (ht, hw) in enumerate(zip(hdr_texts, hdr_widths)):
            cell = tbl.rows[0].cells[i]
            cell.width = hw
            shade_cell(cell, "1F4E79")
            set_cell_text(cell, ht, bold=True, color=WHITE, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        for r_idx, med in enumerate(mediciones):
            row = tbl.rows[r_idx + 1]
            fill = "EBF3FB" if r_idx % 2 == 0 else "FFFFFF"
            vals = [
                med.get("descripcion", ""),
                med.get("ud", ""),
                str(med.get("cantidad", "")),
                med.get("diametro", "-"),
            ]
            for cell, val in zip(row.cells, vals):
                shade_cell(cell, fill)
                set_cell_text(cell, val, size=8)
        doc.add_paragraph()
    else:
        body("Mediciones no disponibles.")

    # ----- SECTION 10: OBSERVACIONES Y LIMITACIONES -----
    h1("10. Observaciones y Limitaciones")
    limitaciones = report.get(
        "observaciones_limitaciones",
        "La presente valoracion y propuesta tecnica se basa en la documentacion e "
        "inspecciones facilitadas, pudiendo existir patologias ocultas o variaciones "
        "no detectables hasta el inicio de los trabajos."
    )
    body(limitaciones, italic=True, color=GRAY)

    # ----- FOOTER -----
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        "Acometidas Europa Saneamiento Tecnico S.L.  |  Tel: 91 386 21 12  |  info@acometidaseuropa.com"
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = GRAY

    # Anexo fotografico: fotogramas/imagenes de la inspeccion
    _append_photo_annex(doc, report.get("_evidencia_img"),
                        titulo="11. Anexo Fotografico")

    # -- Enlace al video de la inspeccion (visor publico, sin necesidad de
    # cuenta): se incrusta aqui para que viaje siempre pegado al documento.
    if enlace_video:
        h1("Video de la inspeccion")
        fila = doc.add_table(rows=1, cols=2)
        fila.autofit = False
        fila.columns[0].width = Cm(3.6)
        fila.columns[1].width = Cm(13.4)
        celda_qr, celda_txt = fila.rows[0].cells
        try:
            import qrcode
            qr_img = qrcode.make(enlace_video, border=1)
            qr_buf = io.BytesIO()
            qr_img.save(qr_buf, format="PNG")
            celda_qr.paragraphs[0].add_run().add_picture(qr_buf, width=Cm(3.2))
        except Exception:
            set_cell_text(celda_qr, "(codigo QR no disponible)", size=8, color=GRAY)
        set_cell_text(celda_txt,
                      "Escanea este codigo con la camara del movil, o pulsa el enlace "
                      "de abajo, para ver el video de la inspeccion:",
                      size=9)
        p_link = celda_txt.add_paragraph()
        p_link.paragraph_format.space_before = Pt(6)
        _add_hyperlink(p_link, enlace_video, enlace_video, size_pt=9, bold=True)

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Helpers de maquetacion DOCX (bordes, margenes, miniaturas cuadradas, fichas)
# ---------------------------------------------------------------------------

_THUMB_DIR: Path | None = None


def _square_thumb(src: Path | str, px: int = 900) -> Path | None:
    """Devuelve una miniatura CUADRADA (recorte centrado) de la imagen origen.

    Los fotogramas de CCTV son 16:9; recortando al centro se conserva la zona
    util (la tuberia) y todas las fotos del informe quedan del mismo tamano.
    Si Pillow no esta disponible o la imagen falla, devuelve None y el llamante
    inserta la original.
    """
    global _THUMB_DIR
    try:
        from PIL import Image
    except Exception:
        return None
    try:
        src = Path(src)
        if _THUMB_DIR is None:
            _THUMB_DIR = Path(tempfile.mkdtemp(prefix="fotos_informe_"))
        dst = _THUMB_DIR / f"sq_{abs(hash(str(src)))}_{px}.jpg"
        if dst.exists():
            return dst
        with Image.open(src) as im:
            im = im.convert("RGB")
            w, h = im.size
            lado = min(w, h)
            left = (w - lado) // 2
            top = (h - lado) // 2
            im = im.crop((left, top, left + lado, top + lado))
            if lado > px:
                im = im.resize((px, px), Image.LANCZOS)
            im.save(dst, "JPEG", quality=88, optimize=True)
        return dst
    except Exception:
        return None


def _cell_borders(cell, color="D6DCE4", sz=6, sides=("top", "left", "bottom", "right")):
    """Bordes finos por celda (sz en 1/8 de punto)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:tcBorders"))
    if old is not None:
        tcPr.remove(old)
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        if side in sides:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        else:
            el.set(qn("w:val"), "nil")
        borders.append(el)
    tcPr.append(borders)


def _add_hyperlink(paragraph, url: str, text: str, color_hex: str = "0563C1", size_pt=9, bold=False):
    """Inserta un hyperlink de verdad (clicable en Word/PDF) en el parrafo dado.

    python-docx no trae un metodo add_hyperlink: hay que registrar la relacion
    externa y montar el <w:hyperlink> a mano.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_hex)
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    if bold:
        rPr.append(OxmlElement("w:b"))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _cell_margins(cell, top=70, left=110, bottom=70, right=110):
    """Padding interior de celda en twips (1 cm = 567)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:tcMar"))
    if old is not None:
        tcPr.remove(old)
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _shade(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:shd"))
    if old is not None:
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _row_keep_together(row):
    """Evita que una fila se parta entre paginas."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    trPr.append(el)


def _repeat_header(row):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def _para_bottom_border(p, color="1F4E79", sz=8):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _page_field(paragraph, instr="PAGE"):
    """Inserta un campo de Word (PAGE / NUMPAGES) que se numera solo."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    r = paragraph.add_run()
    f1 = OxmlElement("w:fldChar")
    f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = f" {instr} "
    f2 = OxmlElement("w:fldChar")
    f2.set(qn("w:fldCharType"), "end")
    r._r.append(f1)
    r._r.append(it)
    r._r.append(f2)
    return r


def _photo_cards(doc, fotos, cols=3, img_cm=4.6, total_cm=17.0):
    """Rejilla de fichas fotograficas: imagen cuadrada en recuadro + etiqueta.

    `fotos` es una lista de dicts: {"path": ruta, "titulo": "FOTO 3", "detalle": "..."}
    """
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    fotos = [f for f in (fotos or []) if f.get("path") and Path(f["path"]).exists()]
    if not fotos:
        return

    AZUL = RGBColor(0x1F, 0x4E, 0x79)
    GRIS = RGBColor(0x6B, 0x72, 0x80)
    col_cm = total_cm / cols
    filas = (len(fotos) + cols - 1) // cols

    tbl = doc.add_table(rows=filas, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for r_i in range(filas):
        _row_keep_together(tbl.rows[r_i])
        for c_i in range(cols):
            cell = tbl.rows[r_i].cells[c_i]
            cell.width = Cm(col_cm)
            idx = r_i * cols + c_i
            if idx >= len(fotos):
                _cell_borders(cell, sides=())
                continue
            foto = fotos[idx]
            _cell_borders(cell, color="C9D5E3", sz=6)
            _cell_margins(cell, top=100, left=100, bottom=90, right=100)
            _shade(cell, "FFFFFF")

            cell.text = ""
            pi = cell.paragraphs[0]
            pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pi.paragraph_format.space_after = Pt(4)
            sq = _square_thumb(foto["path"])
            try:
                if sq:
                    pi.add_run().add_picture(str(sq), width=Cm(img_cm), height=Cm(img_cm))
                else:
                    pi.add_run().add_picture(str(foto["path"]), width=Cm(img_cm))
            except Exception:
                continue

            pt = cell.add_paragraph()
            pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pt.paragraph_format.space_after = Pt(1)
            rt = pt.add_run(foto.get("titulo", ""))
            rt.bold = True
            rt.font.size = Pt(7.5)
            rt.font.color.rgb = AZUL

            det = foto.get("detalle", "")
            if det:
                pd = cell.add_paragraph()
                pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pd.paragraph_format.space_after = Pt(0)
                rd = pd.add_run(det)
                rd.font.size = Pt(6.5)
                rd.font.color.rgb = GRIS

    # Separacion tras la rejilla
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)


def _diagrama_tramo(sec: dict, s_idx: int, px_por_metro: float = 40) -> bytes | None:
    """Esquema vertical del tramo (estilo WinCam): pozos inicio/fin, marcas de
    posicion de cada observacion codificada y flecha con el sentido del flujo.

    Devuelve el PNG como bytes, o None si el tramo no tiene longitud valida
    (nada que dibujar).
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return None

    longitud = float(sec.get("longitud_m", 0) or 0)
    if longitud <= 0:
        return None

    obs_list = sorted(
        (o for o in (sec.get("observaciones_tabla") or [])
         if isinstance(o.get("posicion_m"), (int, float))),
        key=lambda o: o["posicion_m"],
    )

    GRAV_RGB = {"A": (0xC0, 0x00, 0x00), "B": (0xE2, 0x6B, 0x0A),
                "C": (0xBF, 0x90, 0x00), "D": (0x54, 0x82, 0x35)}
    AZUL = (0x1F, 0x4E, 0x79)
    GRIS_BARRA = (0xA9, 0xB4, 0xC0)
    GRIS_POZO = (0xB0, 0xB6, 0xBE)
    GRIS_TXT = (0x33, 0x33, 0x33)

    margen_top, margen_bottom, margen_lat = 50, 50, 20
    alto_barra = max(180, min(900, int(round(longitud * px_por_metro))))
    alto = margen_top + alto_barra + margen_bottom
    ancho = 360
    img = Image.new("RGB", (ancho, alto), "white")
    d = ImageDraw.Draw(img)

    fontdir = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts"
    try:
        font_lbl = ImageFont.truetype(str(fontdir / "segoeuib.ttf"), 14)
        font_txt = ImageFont.truetype(str(fontdir / "segoeui.ttf"), 12)
        font_pos = ImageFont.truetype(str(fontdir / "segoeuib.ttf"), 12)
    except Exception:
        font_lbl = font_txt = font_pos = ImageFont.load_default()

    cx = margen_lat + 55
    y_top = margen_top
    y_bot = margen_top + alto_barra
    r = 17

    # Tuberia
    d.rectangle([cx - 6, y_top, cx + 6, y_bot], fill=GRIS_BARRA)

    # Pozos (circulos) en los extremos, con su identificacion
    d.ellipse([cx - r, y_top - r, cx + r, y_top + r], fill=GRIS_POZO, outline=(0x80, 0x86, 0x90), width=2)
    d.ellipse([cx - r, y_bot - r, cx + r, y_bot + r], fill=GRIS_POZO, outline=(0x80, 0x86, 0x90), width=2)

    nombre_ini = str(sec.get("pozo_inicio") or "-")[:6]
    nombre_fin = str(sec.get("pozo_fin") or "-")[:6]
    for label, ycc in ((nombre_ini, y_top), (nombre_fin, y_bot)):
        bbox = d.textbbox((0, 0), label, font=font_lbl)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        d.text((cx - tw / 2 - bbox[0], ycc - th / 2 - bbox[1]), label, font=font_lbl, fill="white")

    # Marcas de posicion de cada observacion, a la derecha de la tuberia
    for o in obs_list:
        pos = max(0.0, min(longitud, float(o["posicion_m"])))
        y = y_top + (pos / longitud) * alto_barra
        d.line([cx - 11, y, cx + 11, y], fill=(0x55, 0x5B, 0x66), width=2)
        pos_txt = f"{pos:.2f}".replace(".", ",")
        d.text((cx + 20, y - 8), pos_txt, font=font_pos, fill=GRIS_TXT)
        codigo = (o.get("codigo") or "").strip()
        if codigo:
            bbox = d.textbbox((0, 0), pos_txt, font=font_pos)
            grav = (o.get("gravedad") or "").strip().upper()
            color = GRAV_RGB.get(grav, GRIS_TXT)
            d.text((cx + 20 + (bbox[2] - bbox[0]) + 10, y - 8), codigo, font=font_txt, fill=color)

    # Flecha de sentido del flujo, a la izquierda de la tuberia
    contrario = (sec.get("sentido_flujo") or "").strip().lower() == "contrario_inspeccion"
    fx = cx - 32
    fy_a, fy_b = (y_bot - 24, y_top + 24) if contrario else (y_top + 24, y_bot - 24)
    d.line([fx, fy_a, fx, fy_b], fill=AZUL, width=4)
    head = 9
    punta_abajo = fy_b > fy_a
    if punta_abajo:
        d.polygon([(fx - head, fy_b - head), (fx + head, fy_b - head), (fx, fy_b + head)], fill=AZUL)
    else:
        d.polygon([(fx - head, fy_b + head), (fx + head, fy_b + head), (fx, fy_b - head)], fill=AZUL)

    buf = io.BytesIO()
    img.save(buf, "PNG")
    return buf.getvalue()


def attach_wincam_diagramas(report: dict) -> None:
    """Genera el esquema de cada tramo y lo adjunta como PNG en base64
    (clave '_diagrama_png_b64' de cada seccion). Se llama una sola vez tras
    analyze_wincam() para que tanto la vista previa web como generate_wincam_docx
    reutilicen el mismo diagrama sin recalcularlo dos veces.
    """
    for s_idx, sec in enumerate(report.get("secciones", []) or []):
        diag_bytes = _diagrama_tramo(sec, s_idx)
        if diag_bytes:
            sec["_diagrama_png_b64"] = base64.b64encode(diag_bytes).decode()


# ---------------------------------------------------------------------------
# WinCam-style schema and generator
# ---------------------------------------------------------------------------

WINCAM_SCHEMA = """{
  "proyecto": "string: nombre del proyecto",
  "num_proyecto": "string: numero de proyecto o referencia",
  "fecha": "string: fecha de inspeccion en formato DD/MM/YYYY",
  "operador": "string: nombre del operador si visible, sino ''",
  "calle": "string: calle o ubicacion inspeccionada",
  "poblacion": "string: municipio",
  "motivo_inspeccion": "string: motivo (Control general del estado / Preentrega / Mantenimiento / etc)",
  "tipo_red": "string: Red mixta (fecales/pluviales) / Red separativa fecales / Red separativa pluviales",
  "secciones": [
    {
      "num": 1,
      "nombre": "string: V1",
      "pozo_inicio": "string: identificacion pozo/arqueta inicio",
      "pozo_fin": "string: identificacion pozo/arqueta fin",
      "longitud_m": 0.00,
      "diametro_mm": 0,
      "material": "string: PVC/Hormigon/Gres/Fibrocemento/PEAD",
      "tipo_red": "string: igual que arriba o especifico del tramo",
      "sentido_flujo": "string: 'coincide_inspeccion' si el agua circula de pozo_inicio a pozo_fin (mismo sentido que el avance de la camara), 'contrario_inspeccion' si circula al reves, o 'desconocido' si no hay evidencia suficiente",
      "fotos": [1],
      "observaciones_tabla": [
        {
          "posicion_m": 0.00,
          "codigo": "string: ICNI/IFIN/BAB/BAC/BAG/BAH/BAI/BAJ/DAA/DAB/etc EN13508-2",
          "descripcion": "string: descripcion tecnica del hallazgo o evento",
          "gravedad": "string: A/B/C/D segun WRc (A=urgente,B=corto,C=preventivo,D=info) o '' si es ICNI/IFIN",
          "foto": 0
        }
      ]
    }
  ],
  "totales": {
    "longitud_total_m": 0.00,
    "num_secciones": 0,
    "num_fotos": 0,
    "por_diametro": [
      { "dn_mm": 0, "longitud_m": 0.00, "num_tramos": 0 }
    ]
  },
  "nivel_urgencia_global": "Critico|Alto|Medio|Bajo",
  "requiere_intervencion_inmediata": false,
  "preguntas_pendientes": [
    "string: dato que no puedes leer en la documentacion y necesitas que el tecnico confirme"
  ]
}"""

WINCAM_PROMPT = (
    "Genera un informe CCTV en formato WinCam Viewer con la siguiente estructura exacta. "
    "REGLA CRITICA sobre 'secciones': el numero de tramos del array 'secciones' debe coincidir "
    "con el numero de videos/archivos de inspeccion CCTV aportados (uno por tramo grabado), NUNCA "
    "con el numero de tramos que aparezcan dibujados en el croquis. Si se aporta 1 solo video, el "
    "informe debe tener EXACTAMENTE 1 seccion, aunque el croquis muestre una red con muchos mas "
    "tramos, pozos o arquetas: esos tramos sin grabacion NO se reportan, no se les inventan "
    "observaciones ni longitudes. El croquis se usa solo para nombrar/ubicar el tramo grabado. "
    "PROHIBIDO INVENTAR: solo registra datos que puedas leer directamente en los archivos. "
    "Si un dato no es visible (operador, fecha, numero de proyecto, pozo inicio/fin, "
    "direccion, diametro exacto), dejalo como cadena vacia y añadelo a 'preguntas_pendientes'. "
    "Para cada tramo/seccion extrae UNICAMENTE lo que sea legible: numero de seccion, "
    "nombre (V1/V2...), pozos inicio/fin si aparecen en pantalla, "
    "longitud exacta (ultimo metro leido en el contador de pantalla, NUNCA estimada), "
    "diametro en mm si visible, material si visible. "
    "Para la tabla de observaciones usa codigos EN 13508-2 SOLO si la anomalia es claramente visible: "
    "ICNI=inicio inspeccion, IFIN=fin inspeccion, "
    "BAB=fisura longitudinal, BAC=fisura circunferencial, BAD=rotura, BAE=deformacion oval, "
    "BAF=erosion/desgaste, BAG=corrosion, BAH=junta abierta, BAI=desplazamiento de junta, "
    "BAJ=hundimiento, DAA=depositos sedimentarios, DAB=incrustaciones, DAC=intrusion de raices, "
    "DAD=obstruccion total, DAG=infiltracion activa, DAH=exfiltracion, "
    "DAJ=acometida visible, FAA=cambio de material, FAB=cambio de diametro. "
    "Posicion en metros segun contador de pantalla. Gravedad WRc: A=urgente, B=corto plazo, "
    "C=preventivo, D=informativo. Si no puedes determinar la gravedad, dejala vacia. "
    "Calcula totales por diametro solo con los datos leidos. "
    "Determina 'sentido_flujo' de cada tramo a partir del croquis, la pendiente visible, "
    "flechas o marcas de flujo en pantalla, o el contexto aportado: usa 'coincide_inspeccion' "
    "si el agua circula de pozo_inicio a pozo_fin, 'contrario_inspeccion' si es al reves, "
    "o 'desconocido' si no hay evidencia suficiente. No lo inventes. "
    "Lista en 'preguntas_pendientes' TODOS los datos que no hayas podido leer.\n"
    "FOTOS: cada imagen que recibes va precedida de su etiqueta [FOTO n]. "
    "En cada seccion rellena 'fotos' con la lista de numeros n de las fotos que "
    "pertenecen a ESE tramo (las del video correspondiente), en orden de avance. "
    "En cada observacion rellena 'foto' con el numero n de la foto que muestra ese "
    "hallazgo concreto, o 0 si ninguna lo muestra. Nunca inventes numeros de foto: "
    "usa solo los que has recibido."
)


def analyze_wincam(files: list[Path], context: str = "", api_key: str = "",
                   proyecto: str = "", calle: str = "", poblacion: str = "",
                   croquis_path: Path | None = None) -> dict:
    """Analyze files and return WinCam-style structured report."""
    key = os.environ.get("ANTHROPIC_API_KEY", "").strip() or api_key.strip()
    if not key:
        raise ValueError("API Key de Anthropic no configurada.")
    client = anthropic.Anthropic(api_key=key)

    content: list[dict] = []
    if context.strip():
        content.append({"type": "text", "text": f"CONTEXTO:\n{context.strip()}\n\n"})
    if proyecto:
        content.append({"type": "text", "text": f"Nombre del proyecto: {proyecto}\nCalle: {calle}\nPoblacion: {poblacion}\n\n"})

    # Inject croquis before videos so Claude can cross-reference tramos
    if croquis_path and croquis_path.exists():
        _inject_croquis(content, croquis_path)

    content.append({"type": "text", "text": WINCAM_PROMPT + "\n\n"})

    video_ext = {".mp4", ".avi", ".mov", ".mkv", ".webm", ".m4v"}
    pdf_ext = {".pdf"}
    img_ext = {".jpg", ".jpeg", ".png", ".gif", ".webp"}

    evidencia_img: list[str] = []  # frames de video + imagenes -> fichas fotograficas
    evidencia_src: list[str] = []  # archivo de origen de cada foto (mismo indice)

    # Presupuesto global de imagenes (ver MAX_IMGS_IA): reparte fotogramas entre videos.
    _n_videos = sum(1 for f in files if f.suffix.lower() in video_ext) or 1
    _frames_por_video = max(2, min(12, MAX_IMGS_IA // _n_videos))
    _imgs = 0

    for fp in files:
        suf = fp.suffix.lower()
        if suf in video_ext:
            if _imgs >= MAX_IMGS_IA:
                content.append({"type": "text", "text": f"\n[VIDEO CCTV: {fp.name}] - omitido (limite de imagenes alcanzado)\n"})
                continue
            n_obj = min(_frames_por_video, MAX_IMGS_IA - _imgs)
            content.append({"type": "text", "text": f"\n[VIDEO CCTV: {fp.name}]\n"})
            frames = _extract_video_frames(fp, max_frames=n_obj)
            for i, frame in enumerate(frames[:n_obj]):
                data, mt = _encode_image(frame)
                # Numeracion GLOBAL: la IA referencia estas fotos por su numero
                # en 'fotos' de cada seccion y en 'foto' de cada observacion.
                content.append({"type": "text", "text": f"[FOTO {_imgs + 1}] fotograma {i + 1} de {fp.name}:"})
                content.append({"type": "image", "source": {"type": "base64", "media_type": mt, "data": data}})
                evidencia_img.append(str(frame))
                evidencia_src.append(fp.name)
                _imgs += 1
        elif suf in pdf_ext:
            content.append({"type": "text", "text": f"\n[PDF: {fp.name}]\n"})
            content.append({"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": _encode_pdf(fp)}})
        elif suf in img_ext:
            if _imgs >= MAX_IMGS_IA:
                continue
            data, mt = _encode_image(fp)
            content.append({"type": "text", "text": f"[FOTO {_imgs + 1}] imagen {fp.name}:"})
            content.append({"type": "image", "source": {"type": "base64", "media_type": mt, "data": data}})
            evidencia_img.append(str(fp))
            evidencia_src.append(fp.name)
            _imgs += 1

    content.append({"type": "text", "text": f"\nResponde UNICAMENTE con JSON segun este esquema:\n{WINCAM_SCHEMA}"})

    with client.messages.stream(
        model=MODEL,
        max_tokens=24000,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": content}],
    ) as stream:
        response = stream.get_final_message()

    raw = response.content[0].text.strip()
    parsed = _safe_parse_json(raw)
    if parsed is not None:
        parsed["_evidencia_img"] = evidencia_img
        parsed["_evidencia_src"] = evidencia_src
        return parsed
    return {"proyecto": proyecto, "secciones": [], "totales": {}, "_raw": raw,
            "_parse_error": True, "_evidencia_img": evidencia_img,
            "_evidencia_src": evidencia_src}


def generate_wincam_docx(report: dict, output_path: Path, num_ref: str = "", cliente: str = "",
                         enlace_video: str | None = None):
    """Genera el informe de inspeccion CCTV (estilo WinCam) en DOCX.

    Maquetacion: portada, indice + resumen, una pagina por tramo con su ficha,
    su tabla de observaciones y DEBAJO las fotografias de ese tramo en fichas
    cuadradas etiquetadas, y cierre con informacion de proyecto y leyenda.
    """
    from datetime import datetime
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    BLUE = RGBColor(0x1F, 0x4E, 0x79)
    BLUE_MID = RGBColor(0x2E, 0x74, 0xB5)
    GRAY = RGBColor(0x33, 0x33, 0x33)
    LGRAY = RGBColor(0x6B, 0x72, 0x80)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    RED = RGBColor(0xC0, 0x00, 0x00)
    ORANGE = RGBColor(0xE2, 0x6B, 0x0A)
    YELLOW = RGBColor(0xBF, 0x90, 0x00)
    GREEN = RGBColor(0x54, 0x82, 0x35)

    HDR_BG = "1F4E79"
    ZEBRA = "F4F8FC"
    SOFT = "EEF4FB"
    BORDER = "D6DCE4"

    GRAV_BG = {"A": "C00000", "B": "E26B0A", "C": "BF9000", "D": "548235"}
    GRAV_LBL = {"A": "Urgente", "B": "Corto plazo", "C": "Preventivo", "D": "Informativo"}

    doc = Document()

    # Tipografia base del documento
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(9)
    normal.font.color.rgb = GRAY
    try:
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")
    except Exception:
        pass
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(1.6)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)
        sec.different_first_page_header_footer = True

    fecha = report.get("fecha", "") or datetime.now().strftime("%d/%m/%Y")
    proyecto = report.get("proyecto", "")
    calle = report.get("calle", "")
    poblacion = report.get("poblacion", "")
    secciones = report.get("secciones", [])
    totales = report.get("totales", {})

    # Fotografias disponibles, numeradas igual que las vio la IA ([FOTO n])
    fotos_disp = [p for p in (report.get("_evidencia_img") or [])]
    fotos_src = report.get("_evidencia_src") or []

    def cell_text(cell, text, bold=False, color=None, size=8.5,
                  align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("" if text is None else str(text))
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color

    def head_row(tbl, headers, widths=None, size=8):
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            if widths:
                cell.width = widths[i]
            _shade(cell, HDR_BG)
            _cell_borders(cell, color=HDR_BG, sz=6)
            _cell_margins(cell, top=70, bottom=70)
            cell_text(cell, h, bold=True, color=WHITE, size=size,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        _repeat_header(tbl.rows[0])

    def body_cell(cell, idx_fila, width=None):
        if width:
            cell.width = width
        _shade(cell, ZEBRA if idx_fila % 2 == 0 else "FFFFFF")
        _cell_borders(cell, color=BORDER, sz=4)
        _cell_margins(cell, top=55, bottom=55)

    def titulo_bloque(texto, sub=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(sub and 1 or 8)
        r = p.add_run(texto.upper())
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = BLUE
        _para_bottom_border(p, "1F4E79", 8)
        if sub:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(8)
            sr = sp.add_run(sub)
            sr.italic = True
            sr.font.size = Pt(8)
            sr.font.color.rgb = LGRAY

    # ── Encabezado y pie automaticos (paginas 2 en adelante) ──────────────
    sec0 = doc.sections[0]
    hp = sec0.header.paragraphs[0]
    hp.text = ""
    hr1 = hp.add_run("ACOMETIDAS EUROPA SANEAMIENTO TECNICO S.L.")
    hr1.bold = True
    hr1.font.size = Pt(8)
    hr1.font.color.rgb = BLUE
    hr2 = hp.add_run(f"    {proyecto}" + (f"    Ref: {num_ref}" if num_ref else "") + f"    {fecha}")
    hr2.font.size = Pt(7.5)
    hr2.font.color.rgb = LGRAY
    _para_bottom_border(hp, "C9D5E3", 6)

    fpar = sec0.footer.paragraphs[0]
    fpar.text = ""
    fpar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr1 = fpar.add_run("Informe de inspeccion CCTV   |   Pagina ")
    fr1.font.size = Pt(7.5)
    fr1.font.color.rgb = LGRAY
    _page_field(fpar, "PAGE").font.size = Pt(7.5)
    fr2 = fpar.add_run(" de ")
    fr2.font.size = Pt(7.5)
    fr2.font.color.rgb = LGRAY
    _page_field(fpar, "NUMPAGES").font.size = Pt(7.5)

    # ── PORTADA ───────────────────────────────────────────────────────────
    banda = doc.add_table(rows=1, cols=1)
    banda.autofit = False
    bc = banda.rows[0].cells[0]
    bc.width = Cm(17)
    _shade(bc, HDR_BG)
    _cell_borders(bc, color=HDR_BG, sz=6)
    _cell_margins(bc, top=340, left=280, bottom=340, right=280)
    bc.text = ""
    bp1 = bc.paragraphs[0]
    bp1.paragraph_format.space_after = Pt(2)
    br1 = bp1.add_run("ACOMETIDAS EUROPA SANEAMIENTO TECNICO S.L.")
    br1.bold = True
    br1.font.size = Pt(9)
    br1.font.color.rgb = WHITE
    bp2 = bc.add_paragraph()
    bp2.paragraph_format.space_after = Pt(4)
    br2 = bp2.add_run("INFORME DE INSPECCION CCTV")
    br2.bold = True
    br2.font.size = Pt(26)
    br2.font.color.rgb = WHITE
    bp3 = bc.add_paragraph()
    bp3.paragraph_format.space_after = Pt(0)
    br3 = bp3.add_run("Red de saneamiento   |   Codificacion EN 13508-2   |   Clasificacion WRc")
    br3.font.size = Pt(9)
    br3.font.color.rgb = RGBColor(0xC9, 0xD9, 0xEC)

    doc.add_paragraph()

    port_rows = [
        ("Proyecto", proyecto),
        ("Referencia", num_ref or "-"),
        ("Cliente", cliente or "-"),
        ("Calle / Ubicacion", calle or "-"),
        ("Poblacion", poblacion or "-"),
        ("Fecha de inspeccion", fecha),
        ("Motivo", report.get("motivo_inspeccion", "Control general del estado")),
        ("Tipo de red", report.get("tipo_red", "-")),
        ("Operador", report.get("operador", "") or "-"),
    ]
    pt_tbl = doc.add_table(rows=len(port_rows), cols=2)
    pt_tbl.autofit = False
    for i, (lbl, val) in enumerate(port_rows):
        c0, c1 = pt_tbl.rows[i].cells
        c0.width = Cm(5.5)
        c1.width = Cm(11.5)
        for c in (c0, c1):
            _cell_borders(c, color=BORDER, sz=4, sides=("bottom",))
            _cell_margins(c, top=70, left=0, bottom=70, right=60)
        cell_text(c0, lbl.upper(), bold=True, color=BLUE, size=8)
        cell_text(c1, val, size=10)

    doc.add_paragraph()

    urg = report.get("nivel_urgencia_global", "")
    if urg:
        urg_bg = {"Critico": "C00000", "Alto": "E26B0A", "Medio": "BF9000", "Bajo": "548235"}.get(urg, "1F4E79")
        ub = doc.add_table(rows=1, cols=1)
        ub.autofit = False
        uc = ub.rows[0].cells[0]
        uc.width = Cm(17)
        _shade(uc, urg_bg)
        _cell_borders(uc, color=urg_bg, sz=6)
        _cell_margins(uc, top=130, left=200, bottom=130, right=200)
        txt = f"NIVEL DE URGENCIA GLOBAL: {urg.upper()}"
        if report.get("requiere_intervencion_inmediata"):
            txt += "   |   REQUIERE INTERVENCION INMEDIATA"
        cell_text(uc, txt, bold=True, color=WHITE, size=10.5,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.paragraph_format.space_before = Pt(24)
    pr = pie.add_run("Doctor Severo Ochoa, 35, 5-D - 28100 Alcobendas (Madrid)   |   "
                     "913 862 112   |   info@saneamientotecnico.es")
    pr.font.size = Pt(7.5)
    pr.font.color.rgb = LGRAY

    doc.add_page_break()

    # ── INDICE ────────────────────────────────────────────────────────────
    titulo_bloque("Contenido")
    toc = doc.add_table(rows=3 + len(secciones), cols=2)
    toc.autofit = False
    head_row(toc, ["Apartado", "Detalle"], widths=[Cm(7.0), Cm(10.0)])
    idx_rows = [("Resumen de secciones", "Longitudes, diametros y totales")]
    for s_i, sec in enumerate(secciones):
        idx_rows.append((
            f"Tramo {sec.get('num', s_i + 1)} - {sec.get('nombre', '')}",
            f"{sec.get('pozo_inicio', '')} a {sec.get('pozo_fin', '')}   |   "
            f"{float(sec.get('longitud_m', 0) or 0):.2f} m".replace(".", ",")
            + (f"   |   DN {sec.get('diametro_mm')}" if sec.get("diametro_mm") else "")
        ))
    idx_rows.append(("Informacion de proyecto", "Datos del contratista y leyenda de codigos"))
    for r_idx, (lbl, det) in enumerate(idx_rows):
        row = toc.rows[r_idx + 1]
        for c_i, c in enumerate(row.cells):
            body_cell(c, r_idx, width=[Cm(7.0), Cm(10.0)][c_i])
        cell_text(row.cells[0], lbl, bold=True, size=8.5, color=BLUE)
        cell_text(row.cells[1], det, size=8.5)

    doc.add_page_break()

    # ── RESUMEN DE SECCIONES ──────────────────────────────────────────────
    titulo_bloque("Resumen de secciones")

    # Indicadores clave
    long_total = float(totales.get("longitud_total_m", 0) or 0)
    kpis = [
        (f"{long_total:.2f} m".replace(".", ","), "Longitud inspeccionada"),
        (str(totales.get("num_secciones", len(secciones))), "Tramos inspeccionados"),
        (str(totales.get("num_fotos", 0) or len(fotos_disp)), "Fotografias del informe"),
    ]
    kpi_tbl = doc.add_table(rows=1, cols=3)
    kpi_tbl.autofit = False
    for i, (val, lbl) in enumerate(kpis):
        c = kpi_tbl.rows[0].cells[i]
        c.width = Cm(17 / 3)
        _shade(c, SOFT)
        _cell_borders(c, color="C9D5E3", sz=4)
        _cell_margins(c, top=150, left=140, bottom=150, right=140)
        c.text = ""
        p1 = c.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(1)
        r1 = p1.add_run(val)
        r1.bold = True
        r1.font.size = Pt(16)
        r1.font.color.rgb = BLUE
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(lbl.upper())
        r2.font.size = Pt(7)
        r2.font.color.rgb = LGRAY

    doc.add_paragraph()

    # Tabla resumen de tramos
    res_w = [Cm(1.2), Cm(2.4), Cm(4.3), Cm(2.6), Cm(2.8), Cm(1.9), Cm(1.8)]
    tbl_res = doc.add_table(rows=1 + len(secciones), cols=7)
    tbl_res.autofit = False
    head_row(tbl_res, ["N.", "Tramo", "Pozos (inicio - fin)", "Material",
                       "Long. insp.", "DN (mm)", "Anomalias"], widths=res_w)
    for r_idx, sec in enumerate(secciones):
        row = tbl_res.rows[r_idx + 1]
        obs = sec.get("observaciones_tabla", []) or []
        n_anom = sum(1 for o in obs if (o.get("gravedad") or "") in GRAV_BG)
        vals = [
            sec.get("num", r_idx + 1),
            sec.get("nombre", ""),
            f"{sec.get('pozo_inicio', '') or '-'}  a  {sec.get('pozo_fin', '') or '-'}",
            sec.get("material", "") or "-",
            f"{float(sec.get('longitud_m', 0) or 0):.2f} m".replace(".", ","),
            sec.get("diametro_mm", "") or "-",
            str(n_anom),
        ]
        for c_idx, val in enumerate(vals):
            body_cell(row.cells[c_idx], r_idx, width=res_w[c_idx])
            cell_text(row.cells[c_idx], val, size=8,
                      bold=(c_idx == 1),
                      color=BLUE if c_idx == 1 else None,
                      align=WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, 4, 5, 6) else WD_ALIGN_PARAGRAPH.LEFT)

    # Totales por diametro
    por_dn = totales.get("por_diametro", [])
    if por_dn:
        doc.add_paragraph()
        dn_tbl = doc.add_table(rows=1 + len(por_dn), cols=3)
        dn_tbl.autofit = False
        dn_w = [Cm(5.6), Cm(5.7), Cm(5.7)]
        head_row(dn_tbl, ["Diametro", "Longitud", "N. de tramos"], widths=dn_w)
        for i, dn_item in enumerate(por_dn):
            row = dn_tbl.rows[i + 1]
            vals = [f"DN {dn_item.get('dn_mm', '')}",
                    f"{float(dn_item.get('longitud_m', 0) or 0):.2f} m".replace(".", ","),
                    str(dn_item.get("num_tramos", ""))]
            for c_idx, val in enumerate(vals):
                body_cell(row.cells[c_idx], i, width=dn_w[c_idx])
                cell_text(row.cells[c_idx], val, size=8.5,
                          bold=(c_idx == 0),
                          align=WD_ALIGN_PARAGRAPH.CENTER)

        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        tr = tp.add_run(f"Total inspeccionado: {long_total:.2f} m".replace(".", ","))
        tr.bold = True
        tr.font.size = Pt(10)
        tr.font.color.rgb = BLUE

    pend = report.get("preguntas_pendientes") or []
    if pend:
        doc.add_paragraph()
        np_ = doc.add_paragraph()
        np_.paragraph_format.space_after = Pt(2)
        nr = np_.add_run("Datos pendientes de confirmar por el tecnico")
        nr.bold = True
        nr.font.size = Pt(9)
        nr.font.color.rgb = BLUE_MID
        for q in pend[:12]:
            qp = doc.add_paragraph()
            qp.paragraph_format.space_after = Pt(1)
            qp.paragraph_format.left_indent = Cm(0.4)
            qr = qp.add_run(f"-  {q}")
            qr.font.size = Pt(8)
            qr.font.color.rgb = LGRAY

    doc.add_page_break()

    # ── Reparto de fotografias por tramo ──────────────────────────────────
    # La IA asigna cada foto a su tramo ("fotos") y a la observacion concreta
    # que ilustra ("foto"). Si no lo hace, se reparten por video de origen o,
    # en ultimo caso, equitativamente en orden.
    n_fotos = len(fotos_disp)
    asignadas: set[int] = set()
    sec_fotos: list[list[int]] = []
    foto_obs: dict[int, dict] = {}

    def _idx_foto(v):
        try:
            k = int(v) - 1
        except (TypeError, ValueError):
            return None
        return k if 0 <= k < n_fotos else None

    for sec in secciones:
        ids: list[int] = []
        for o in sec.get("observaciones_tabla", []) or []:
            k = _idx_foto(o.get("foto"))
            if k is not None:
                foto_obs.setdefault(k, o)
        for v in (sec.get("fotos") or []):
            k = _idx_foto(v)
            if k is not None and k not in asignadas:
                ids.append(k)
                asignadas.add(k)
        for o in sec.get("observaciones_tabla", []) or []:
            k = _idx_foto(o.get("foto"))
            if k is not None and k not in asignadas:
                ids.append(k)
                asignadas.add(k)
        sec_fotos.append(sorted(ids))

    if secciones and n_fotos and not asignadas:
        grupos: list[list[int]] = []
        if len(fotos_src) == n_fotos:
            orden = []
            for s in fotos_src:
                if s not in orden:
                    orden.append(s)
            if len(orden) == len(secciones):
                grupos = [[i for i, s in enumerate(fotos_src) if s == nm] for nm in orden]
        if not grupos:
            base, resto = divmod(n_fotos, len(secciones))
            k = 0
            for i in range(len(secciones)):
                n = base + (1 if i < resto else 0)
                grupos.append(list(range(k, k + n)))
                k += n
        sec_fotos = grupos
        asignadas = set(range(n_fotos))

    def _fichas_foto(indices, nombre_tramo):
        fichas = []
        for k in indices:
            o = foto_obs.get(k)
            titulo = f"FOTO {k + 1}   |   {nombre_tramo}"
            detalle = ""
            if o:
                pos = o.get("posicion_m")
                partes = []
                if isinstance(pos, (int, float)):
                    partes.append(f"{float(pos):.2f} m".replace(".", ","))
                if o.get("codigo"):
                    partes.append(str(o["codigo"]))
                desc = (o.get("descripcion") or "").strip()
                if desc:
                    partes.append(desc if len(desc) <= 70 else desc[:67] + "...")
                detalle = "  |  ".join(partes)
            fichas.append({"path": fotos_disp[k], "titulo": titulo, "detalle": detalle})
        return fichas

    # ── PAGINAS POR SECCION ───────────────────────────────────────────────
    for s_idx, sec in enumerate(secciones):
        nombre = sec.get("nombre", f"V{s_idx + 1}")
        titulo_bloque(f"Tramo {sec.get('num', s_idx + 1)} - {nombre}",
                      sub=f"{sec.get('pozo_inicio', '') or '-'} a {sec.get('pozo_fin', '') or '-'}"
                          f"   |   {calle}" + (f", {poblacion}" if poblacion else ""))

        # Ficha del tramo: 12 campos en tarjetas de 4 columnas
        campos = [
            ("Fecha", fecha),
            ("N. de tramo", str(sec.get("num", s_idx + 1))),
            ("Nombre del tramo", nombre),
            ("Tipo de red", sec.get("tipo_red", "") or report.get("tipo_red", "") or "-"),
            ("Pozo inicio", sec.get("pozo_inicio", "") or "-"),
            ("Pozo final", sec.get("pozo_fin", "") or "-"),
            ("Longitud inspeccionada", f"{float(sec.get('longitud_m', 0) or 0):.2f} m".replace(".", ",")),
            ("Diametro", f"{sec.get('diametro_mm', '')} mm" if sec.get("diametro_mm") else "-"),
            ("Material", sec.get("material", "") or "-"),
            ("Calle", calle or "-"),
            ("Poblacion", poblacion or "-"),
            ("Motivo", report.get("motivo_inspeccion", "Control general del estado")),
        ]
        ficha = doc.add_table(rows=3, cols=4)
        ficha.autofit = False
        for i, (lbl, val) in enumerate(campos):
            cell = ficha.rows[i // 4].cells[i % 4]
            cell.width = Cm(4.25)
            _shade(cell, SOFT if i // 4 % 2 == 0 else "FFFFFF")
            _cell_borders(cell, color=BORDER, sz=4)
            _cell_margins(cell, top=70, left=110, bottom=70, right=90)
            cell.text = ""
            p1 = cell.paragraphs[0]
            p1.paragraph_format.space_after = Pt(0)
            r1 = p1.add_run(lbl.upper())
            r1.bold = True
            r1.font.size = Pt(6.5)
            r1.font.color.rgb = LGRAY
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_after = Pt(0)
            r2 = p2.add_run(str(val))
            r2.font.size = Pt(8.5)
            r2.font.color.rgb = GRAY

        doc.add_paragraph()

        # Tabla de observaciones
        obs_list = sec.get("observaciones_tabla", []) or []
        if obs_list:
            ot = doc.add_paragraph()
            ot.paragraph_format.space_after = Pt(3)
            otr = ot.add_run("Observaciones codificadas (EN 13508-2)")
            otr.bold = True
            otr.font.size = Pt(9)
            otr.font.color.rgb = BLUE_MID

            obs_w = [Cm(2.2), Cm(1.9), Cm(9.1), Cm(2.3), Cm(1.5)]
            obs_tbl = doc.add_table(rows=1 + len(obs_list), cols=5)
            obs_tbl.autofit = False
            head_row(obs_tbl, ["Posicion", "Codigo", "Descripcion / Observacion",
                               "Gravedad", "Foto"], widths=obs_w)

            for o_idx, obs in enumerate(obs_list):
                row = obs_tbl.rows[o_idx + 1]
                _row_keep_together(row)
                codigo = obs.get("codigo", "") or ""
                grav = (obs.get("gravedad", "") or "").strip().upper()
                pos = obs.get("posicion_m", 0)
                desc = obs.get("descripcion", "") or ""
                k_foto = _idx_foto(obs.get("foto"))

                for c_idx, c in enumerate(row.cells):
                    body_cell(c, o_idx, width=obs_w[c_idx])

                pos_txt = f"{float(pos):.2f} m".replace(".", ",") if isinstance(pos, (int, float)) else str(pos)
                cell_text(row.cells[0], pos_txt, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

                # Codigo: pastilla de color segun gravedad
                row.cells[1].text = ""
                cp = row.cells[1].paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_after = Pt(0)
                cr = cp.add_run(codigo)
                cr.bold = True
                cr.font.size = Pt(8)
                if grav in GRAV_BG:
                    cr.font.color.rgb = WHITE
                    _shade(row.cells[1], GRAV_BG[grav])
                else:
                    cr.font.color.rgb = BLUE

                cell_text(row.cells[2], desc, size=8)

                if grav in GRAV_BG:
                    grav_color = {"A": RED, "B": ORANGE, "C": YELLOW, "D": GREEN}[grav]
                    cell_text(row.cells[3], f"{grav} - {GRAV_LBL[grav]}", bold=True,
                              color=grav_color, size=7.5, align=WD_ALIGN_PARAGRAPH.CENTER)
                else:
                    cell_text(row.cells[3], "-", size=8, color=LGRAY,
                              align=WD_ALIGN_PARAGRAPH.CENTER)

                cell_text(row.cells[4], f"Foto {k_foto + 1}" if k_foto is not None else "-",
                          size=7.5, color=BLUE_MID if k_foto is not None else LGRAY,
                          align=WD_ALIGN_PARAGRAPH.CENTER)

        # Esquema del tramo: perfil, marcas de posicion y sentido del flujo,
        # debajo de la tabla de observaciones y antes de las fotografias.
        # Reutiliza el diagrama ya calculado por attach_wincam_diagramas() si
        # existe (misma imagen que ve la vista previa web); si no, lo genera aqui.
        diag_b64 = sec.get("_diagrama_png_b64")
        diag_bytes = base64.b64decode(diag_b64) if diag_b64 else _diagrama_tramo(sec, s_idx)
        if diag_bytes:
            dt = doc.add_paragraph()
            dt.paragraph_format.space_before = Pt(8)
            dt.paragraph_format.space_after = Pt(3)
            dtr = dt.add_run(f"Esquema del tramo {nombre} y sentido del flujo")
            dtr.bold = True
            dtr.font.size = Pt(9)
            dtr.font.color.rgb = BLUE_MID

            dp = doc.add_paragraph()
            dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dp.paragraph_format.space_after = Pt(2)
            dp.add_run().add_picture(io.BytesIO(diag_bytes), height=Cm(9.0))

            if (sec.get("sentido_flujo") or "").strip().lower() == "desconocido":
                dn = doc.add_paragraph()
                dn.alignment = WD_ALIGN_PARAGRAPH.CENTER
                dn.paragraph_format.space_after = Pt(4)
                dnr = dn.add_run("* Sentido del flujo no confirmado en la inspeccion; se muestra de pozo inicio a pozo final.")
                dnr.italic = True
                dnr.font.size = Pt(7)
                dnr.font.color.rgb = LGRAY

        # Fotografias de ESTE tramo, justo debajo de su tabla
        fichas = _fichas_foto(sec_fotos[s_idx] if s_idx < len(sec_fotos) else [], nombre)
        if fichas:
            fp_ = doc.add_paragraph()
            fp_.paragraph_format.space_before = Pt(8)
            fp_.paragraph_format.space_after = Pt(4)
            fpr = fp_.add_run(f"Reportaje fotografico del tramo {nombre}")
            fpr.bold = True
            fpr.font.size = Pt(9)
            fpr.font.color.rgb = BLUE_MID
            _photo_cards(doc, fichas, cols=3, img_cm=4.4)

        if s_idx < len(secciones) - 1:
            doc.add_page_break()

    # ── FOTOGRAFIAS NO ASIGNADAS A NINGUN TRAMO ───────────────────────────
    sobrantes = [k for k in range(n_fotos) if k not in asignadas]
    if sobrantes:
        doc.add_page_break()
        titulo_bloque("Anexo fotografico",
                      sub="Imagenes de la inspeccion no vinculadas a un tramo concreto.")
        _photo_cards(doc, _fichas_foto(sobrantes, "General"), cols=3, img_cm=4.4)

    # ── ULTIMA PAGINA: INFO PROYECTO + LEYENDA ────────────────────────────
    doc.add_page_break()
    titulo_bloque("Informacion de proyecto")

    info_rows = [
        ("Nombre de proyecto", proyecto),
        ("Referencia", num_ref or "-"),
        ("Fecha de inspeccion", fecha),
        ("Calle / Ubicacion", calle or "-"),
        ("Poblacion", poblacion or "-"),
        ("Cliente", cliente or "-"),
        ("Contratista", "Acometidas Europa Saneamiento Tecnico S.L."),
        ("Direccion", "Doctor Severo Ochoa, 35, 5-D - 28100 Alcobendas, Madrid"),
        ("Telefono", "913 862 112"),
        ("E-mail", "info@saneamientotecnico.es"),
    ]
    info_tbl = doc.add_table(rows=len(info_rows), cols=2)
    info_tbl.autofit = False
    for i, (lbl, val) in enumerate(info_rows):
        c0, c1 = info_tbl.rows[i].cells
        c0.width = Cm(5.5)
        c1.width = Cm(11.5)
        for c in (c0, c1):
            _cell_borders(c, color=BORDER, sz=4, sides=("bottom",))
            _cell_margins(c, top=60, left=0, bottom=60, right=60)
        cell_text(c0, lbl.upper(), bold=True, color=BLUE, size=7.5)
        cell_text(c1, val, size=9)

    doc.add_paragraph()
    titulo_bloque("Leyenda de clasificacion WRc")
    leg_w = [Cm(2.2), Cm(3.6), Cm(11.2)]
    leg = doc.add_table(rows=5, cols=3)
    leg.autofit = False
    head_row(leg, ["Grado", "Nivel", "Criterio de actuacion"], widths=leg_w)
    leg_rows = [
        ("A", "Urgente", "Defecto grave o riesgo de colapso. Intervencion inmediata."),
        ("B", "Corto plazo", "Defecto significativo. Reparacion programada a corto plazo."),
        ("C", "Preventivo", "Defecto leve. Mantenimiento preventivo y seguimiento."),
        ("D", "Informativo", "Observacion sin afeccion estructural ni funcional."),
    ]
    for i, (g, nivel, crit) in enumerate(leg_rows):
        row = leg.rows[i + 1]
        for c_idx, c in enumerate(row.cells):
            body_cell(c, i, width=leg_w[c_idx])
        _shade(row.cells[0], GRAV_BG[g])
        cell_text(row.cells[0], g, bold=True, color=WHITE, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row.cells[1], nivel, bold=True, size=8.5,
                  color={"A": RED, "B": ORANGE, "C": YELLOW, "D": GREEN}[g])
        cell_text(row.cells[2], crit, size=8.5)

    nota = doc.add_paragraph()
    nota.paragraph_format.space_before = Pt(10)
    nr = nota.add_run("Informe generado a partir de la inspeccion CCTV realizada por "
                      "Acometidas Europa Saneamiento Tecnico S.L. Los codigos de las "
                      "observaciones siguen la norma EN 13508-2 y la clasificacion de "
                      "gravedad el metodo WRc.")
    nr.italic = True
    nr.font.size = Pt(7.5)
    nr.font.color.rgb = LGRAY

    # -- Enlace al video de la inspeccion (visor publico, sin necesidad de
    # cuenta): se incrusta aqui para que viaje siempre pegado al documento.
    if enlace_video:
        titulo_bloque("Video de la inspeccion")
        fila = doc.add_table(rows=1, cols=2)
        fila.autofit = False
        fila.columns[0].width = Cm(3.6)
        fila.columns[1].width = Cm(13.4)
        celda_qr, celda_txt = fila.rows[0].cells
        try:
            import qrcode
            qr_img = qrcode.make(enlace_video, border=1)
            qr_buf = io.BytesIO()
            qr_img.save(qr_buf, format="PNG")
            celda_qr.paragraphs[0].add_run().add_picture(qr_buf, width=Cm(3.2))
        except Exception:
            cell_text(celda_qr, "(codigo QR no disponible)", size=8, color=LGRAY)
        cell_text(celda_txt,
                  "Escanea este codigo con la camara del movil, o pulsa el enlace de "
                  "abajo, para ver el video de la inspeccion con las anomalias senaladas:",
                  size=9)
        p_link = celda_txt.add_paragraph()
        p_link.paragraph_format.space_before = Pt(6)
        _add_hyperlink(p_link, enlace_video, enlace_video, size_pt=9, bold=True)

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Extract solicitud data from uploaded documents / emails
# ---------------------------------------------------------------------------

_SOLICITUD_SYSTEM = """Eres un asistente administrativo de Acometidas Europa Saneamiento Tecnico S.L.
Tu tarea es leer documentos (correos electronicos, PDFs, capturas de pantalla, memorias) enviados por clientes o administradores de fincas y extraer los datos necesarios para rellenar un presupuesto.

REGLAS:
- Extrae SOLO lo que esta escrito claramente en el documento. No inventes nada.
- Si un dato no aparece, deja el campo como cadena vacia "".
- El campo "servicio" debe ser una descripcion corta del trabajo solicitado (ej: "Desatasco y limpieza de red", "Inspeccion CCTV bajantes").
- El campo "obra" es la direccion completa donde se realizara el trabajo.
- El campo "cliente_nombre" es el nombre de la comunidad, empresa o persona que solicita.
- El campo "administracion" es el nombre de la administracion o administrador de fincas si aparece.
- El campo "provincia" extrae solo la ciudad o provincia (ej: "Madrid", "Barcelona").
- Si hay numero de referencia/expediente del cliente, ponlo en "ref_cliente".
- Responde UNICAMENTE con el JSON. Sin texto fuera del JSON."""

_SOLICITUD_SCHEMA = """{
  "cliente_nombre": "string - nombre del cliente o comunidad",
  "cliente_dir": "string - direccion del cliente",
  "cliente_tel": "string - telefono del cliente",
  "cliente_email": "string - email del cliente",
  "administracion": "string - nombre del administrador o administracion",
  "obra": "string - direccion completa donde se realizara el trabajo",
  "servicio": "string - descripcion corta del servicio solicitado",
  "provincia": "string - ciudad o provincia",
  "ref_cliente": "string - numero de referencia del cliente si existe",
  "notas": "string - cualquier dato relevante no incluido en los campos anteriores"
}"""


def extract_solicitud_data(files: list[Path], api_key: str = "") -> dict:
    """Extract budget form data from uploaded documents/emails using Claude AI."""
    key = os.environ.get("ANTHROPIC_API_KEY", "").strip() or api_key.strip()
    if not key:
        raise ValueError("API Key de Anthropic no configurada.")
    client = anthropic.Anthropic(api_key=key)

    content: list[dict] = []
    content.append({"type": "text", "text": (
        "Lee el/los siguiente(s) documento(s) y extrae los datos del presupuesto.\n"
        f"Devuelve UNICAMENTE un JSON con este esquema:\n{_SOLICITUD_SCHEMA}\n\n"
        "DOCUMENTOS:\n"
    )})

    img_ext = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
    txt_ext = {".txt", ".eml", ".msg", ".html", ".htm", ".csv"}

    for fp in files:
        suf = fp.suffix.lower()
        if suf in img_ext:
            data, mt = _encode_image(fp)
            content.append({"type": "image", "source": {"type": "base64", "media_type": mt, "data": data}})
        elif suf == ".pdf":
            content.append({
                "type": "document",
                "source": {"type": "base64", "media_type": "application/pdf", "data": _encode_pdf(fp)},
            })
        elif suf in txt_ext:
            try:
                text = fp.read_text(encoding="utf-8", errors="replace")
            except Exception:
                text = fp.read_text(errors="replace")
            content.append({"type": "text", "text": f"\n--- DOCUMENTO: {fp.name} ---\n{text[:8000]}\n"})
        else:
            # Try to read as text anyway
            try:
                text = fp.read_text(encoding="utf-8", errors="replace")
                content.append({"type": "text", "text": f"\n--- ARCHIVO: {fp.name} ---\n{text[:4000]}\n"})
            except Exception:
                pass  # Skip unreadable files

    response = client.messages.create(
        model=MODEL,
        max_tokens=1024,
        system=_SOLICITUD_SYSTEM,
        messages=[{"role": "user", "content": content}],
    )

    raw = response.content[0].text.strip()
    # Strip markdown fences if present
    raw = re.sub(r"^```[a-z]*\n?", "", raw, flags=re.MULTILINE)
    raw = re.sub(r"```$", "", raw, flags=re.MULTILINE)
    return json.loads(raw.strip())


# ---------------------------------------------------------------------------
# Extract partidas from a subcontractor's budget document
# ---------------------------------------------------------------------------

_SUBCONTRATA_SYSTEM = """Eres un asistente tecnico de Acometidas Europa Saneamiento Tecnico S.L.
Tu tarea es leer presupuestos de subcontratas o de otras empresas (PDF, imagen o texto) y extraer TODAS las partidas o lineas de trabajo con sus datos economicos.

REGLAS:
- Extrae TODAS las partidas/lineas que aparezcan en el documento, sin omitir ninguna.
- Si una linea no tiene precio unitario pero si importe total y cantidad, calcula: precio_unitario = importe / cantidad.
- Si solo hay importe total sin cantidad, pon cantidad=1 y precio_unitario=importe_total.
- Normaliza las unidades: "m.l." -> "ml", "m2" -> "m2", "Ud." -> "ud", "PA" -> "PA", etc.
- Si el documento tiene capitulos o secciones, incluye el nombre del capitulo como prefijo en la descripcion (ej: "CAPITULO 1 - Excavacion y demolicion").
- Los precios SIEMPRE en formato numerico (float), sin simbolo de euro.
- Si no puedes leer un valor claramente, ponlo a 0.
- Responde UNICAMENTE con el JSON. Sin texto fuera del JSON."""

_SUBCONTRATA_SCHEMA = """{
  "empresa": "string - nombre de la empresa subcontratista si aparece",
  "referencia": "string - numero de presupuesto o referencia si aparece",
  "partidas": [
    {
      "descripcion": "string - descripcion completa de la partida",
      "unidad": "string - ud, ml, m2, m3, PA, kg, h, etc.",
      "cantidad": 0.0,
      "precio_unitario": 0.0,
      "importe": 0.0
    }
  ],
  "total": 0.0,
  "notas": "string - observaciones relevantes del presupuesto"
}"""


def extract_partidas_from_subcontrata(files: list[Path], api_key: str = "") -> dict:
    """Extract all budget line items from a subcontractor's budget document."""
    key = os.environ.get("ANTHROPIC_API_KEY", "").strip() or api_key.strip()
    if not key:
        raise ValueError("API Key de Anthropic no configurada.")
    client = anthropic.Anthropic(api_key=key)

    content: list[dict] = []
    content.append({"type": "text", "text": (
        "Lee este presupuesto de subcontrata y extrae TODAS las partidas con sus datos economicos.\n"
        f"Devuelve UNICAMENTE un JSON con este esquema:\n{_SUBCONTRATA_SCHEMA}\n\n"
        "PRESUPUESTO:\n"
    )})

    img_ext = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
    txt_ext = {".txt", ".eml", ".html", ".htm", ".csv", ".xml"}

    for fp in files:
        suf = fp.suffix.lower()
        if suf in img_ext:
            data, mt = _encode_image(fp)
            content.append({"type": "image", "source": {"type": "base64", "media_type": mt, "data": data}})
        elif suf == ".pdf":
            content.append({
                "type": "document",
                "source": {"type": "base64", "media_type": "application/pdf", "data": _encode_pdf(fp)},
            })
        elif suf in txt_ext:
            try:
                text = fp.read_text(encoding="utf-8", errors="replace")
            except Exception:
                text = fp.read_text(errors="replace")
            content.append({"type": "text", "text": f"\n--- DOCUMENTO: {fp.name} ---\n{text[:200000]}\n"})
        else:
            try:
                text = fp.read_text(encoding="utf-8", errors="replace")
                content.append({"type": "text", "text": f"\n--- ARCHIVO: {fp.name} ---\n{text[:200000]}\n"})
            except Exception:
                pass

    response = client.messages.create(
        model=MODEL,
        max_tokens=16000,
        system=_SUBCONTRATA_SYSTEM,
        messages=[{"role": "user", "content": content}],
    )

    raw = response.content[0].text.strip()
    result = _safe_parse_json(raw)
    if result is None:
        raise ValueError("No se pudo interpretar la respuesta de la IA (JSON invalido o cortado).")

    # Recalculate importe for each partida to ensure consistency (formato espanol tolerado)
    for p in result.get("partidas", []):
        cant = _to_float(p.get("cantidad"), 0.0)
        precio = _to_float(p.get("precio_unitario"), 0.0)
        if precio == 0:
            imp = _to_float(p.get("importe"), 0.0)
            if imp and cant:
                precio = round(imp / cant, 2)
            elif imp and not cant:
                cant = 1.0
                precio = imp
        if cant == 0:
            cant = 1.0
        p["cantidad"] = cant
        p["precio_unitario"] = precio
        p["importe"] = round(cant * precio, 2)

    return result


# ---------------------------------------------------------------------------
# Presupuesto/informe COMPLETO multi-oficio -> extraccion fiel de TODAS las partidas
# ---------------------------------------------------------------------------

_PRESUP_COMPLETO_SYSTEM = """Eres un aparejador tecnico de Acometidas Europa que prepara presupuestos de obra a partir de documentos aportados (informes tecnicos, presupuestos de otras empresas, mediciones, hojas de calculo).
Tu tarea es EXTRAER TODAS Y CADA UNA de las partidas que aparezcan en el/los documento(s), de CUALQUIER oficio (albanileria, fontaneria, electricidad, saneamiento, poceria, pintura, carpinteria, climatizacion, etc.), SIN OMITIR NINGUNA y SIN RESUMIR NI AGRUPAR varias partidas en una sola.

REGLAS CRITICAS:
- Extrae TODAS las partidas, una por una, en el mismo orden en que aparecen. OMITIR partidas es un error grave.
- Respeta la descripcion, la unidad, la cantidad y el precio unitario tal como figuran en el documento. NO inventes precios ni cantidades.
- Si una linea tiene importe total y cantidad pero no precio unitario, calcula precio_unitario = importe / cantidad.
- Si solo hay importe total sin cantidad, pon cantidad = 1 y precio_unitario = importe.
- Si el documento agrupa por capitulos u oficios, antepon el nombre del capitulo/oficio a la descripcion (ej: "FONTANERIA - Sustitucion de bajante de PVC DN110").
- Normaliza unidades: "m.l."/"ml." -> "ml", "m2"/"M2" -> "m2", "Ud."/"uds" -> "ud", "PA"/"P.A." -> "PA".
- Precios y cantidades SIEMPRE como numero (float), sin simbolo de euro, con punto decimal.
- Si un valor no se lee con seguridad ponlo a 0, pero NUNCA elimines la partida por ello.
- Responde UNICAMENTE con el JSON solicitado, sin texto adicional ni markdown."""

_PRESUP_COMPLETO_SCHEMA = """{
  "informe_tecnico": "string: 1-3 frases resumiendo el objeto de la obra (deducido del documento)",
  "solucion_adoptar": "string: 1-2 frases con la solucion global propuesta",
  "memoria_tecnica": "string: 1-3 frases sobre como se ejecutaran los trabajos",
  "partidas": [
    {
      "descripcion": "string: descripcion completa de la partida (con prefijo de oficio/capitulo si aplica)",
      "unidad": "string: ud, ml, m2, m3, PA, kg, h, etc.",
      "cantidad": 0.0,
      "precio_unitario": 0.0,
      "importe": 0.0
    }
  ]
}"""


def _text_from_office_file(fp: Path) -> str:
    """Extrae texto plano de DOCX/DOC o XLSX/XLS para pasarlo al modelo."""
    suf = fp.suffix.lower()
    if suf in (".docx", ".doc"):
        try:
            from docx import Document as _DocxDoc
            doc = _DocxDoc(str(fp))
            partes = [p.text for p in doc.paragraphs if p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    celdas = [c.text.strip() for c in row.cells]
                    if any(celdas):
                        partes.append(" | ".join(celdas))
            return "\n".join(partes)
        except Exception:
            return fp.read_text(encoding="utf-8", errors="replace")
    if suf in (".xlsx", ".xls"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(fp), read_only=True, data_only=True)
            lineas = []
            for ws in wb.worksheets:
                lineas.append(f"--- HOJA: {ws.title} ---")
                for row in ws.iter_rows(values_only=True):
                    celdas = [str(c) for c in row if c is not None and str(c).strip()]
                    if celdas:
                        lineas.append(" | ".join(celdas))
            return "\n".join(lineas)
        except Exception:
            return ""
    return fp.read_text(encoding="utf-8", errors="replace")


def extract_full_presupuesto(files: list[Path], api_key: str = "", descripcion: str = "") -> dict:
    """Extrae TODAS las partidas (multi-oficio) de un informe/presupuesto aportado.

    Robusto frente a documentos largos: si la respuesta se corta por limite de tokens,
    continua la generacion automaticamente y concatena hasta completar el JSON.
    """
    key = os.environ.get("ANTHROPIC_API_KEY", "").strip() or api_key.strip()
    if not key:
        raise ValueError("API Key de Anthropic no configurada.")
    client = anthropic.Anthropic(api_key=key)

    img_ext = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
    txt_ext = {".txt", ".md", ".eml", ".html", ".htm", ".csv", ".xml"}
    office_ext = {".docx", ".doc", ".xlsx", ".xls"}

    MAX_TEXT = 200000  # tope de seguridad por documento

    content: list[dict] = []
    intro = ("Lee el/los siguiente(s) documento(s) y extrae TODAS las partidas de obra "
             "con sus datos economicos, de todos los oficios, sin omitir ninguna.\n")
    if descripcion.strip():
        intro += f"\nContexto aportado por el tecnico: {descripcion.strip()}\n"
    intro += f"\nDevuelve UNICAMENTE un JSON con este esquema:\n{_PRESUP_COMPLETO_SCHEMA}\n\nDOCUMENTOS:\n"
    content.append({"type": "text", "text": intro})

    for fp in files:
        suf = fp.suffix.lower()
        if suf in img_ext:
            data, mt = _encode_image(fp)
            content.append({"type": "image", "source": {"type": "base64", "media_type": mt, "data": data}})
        elif suf == ".pdf":
            content.append({
                "type": "document",
                "source": {"type": "base64", "media_type": "application/pdf", "data": _encode_pdf(fp)},
            })
        elif suf in office_ext:
            text = _text_from_office_file(fp)[:MAX_TEXT]
            content.append({"type": "text", "text": f"\n--- DOCUMENTO: {fp.name} ---\n{text}\n"})
        elif suf in txt_ext:
            try:
                text = fp.read_text(encoding="utf-8", errors="replace")
            except Exception:
                text = fp.read_text(errors="replace")
            content.append({"type": "text", "text": f"\n--- DOCUMENTO: {fp.name} ---\n{text[:MAX_TEXT]}\n"})

    messages = [{"role": "user", "content": content}]
    raw_parts: list[str] = []
    truncado = False
    for _ in range(6):  # 1 + hasta 5 continuaciones
        response = client.messages.create(
            model=MODEL,
            max_tokens=16000,
            system=_PRESUP_COMPLETO_SYSTEM,
            messages=messages,
        )
        chunk = response.content[0].text or ""
        raw_parts.append(chunk)
        if getattr(response, "stop_reason", None) != "max_tokens":
            truncado = False
            break
        truncado = True
        # Continuar exactamente donde se corto
        messages.append({"role": "assistant", "content": chunk})
        messages.append({"role": "user", "content": (
            "La respuesta se corto. Continua el JSON EXACTAMENTE donde lo dejaste, "
            "sin repetir ni un solo caracter de lo ya escrito, sin markdown y sin explicaciones. "
            "Sigue enumerando las partidas que falten hasta cerrar el JSON."
        )})

    raw = "".join(raw_parts)
    parsed = _safe_parse_json(raw)
    if parsed is None:
        return {"_error": "No se pudo parsear la respuesta de la IA.", "_raw": raw[:2000],
                "_truncado": truncado}

    # Normalizar partidas (acepta numeros en formato espanol: "12,50", "1.234,56")
    partidas = []
    for p in (parsed.get("partidas") or []):
        cant = _to_float(p.get("cantidad"), 0.0)
        precio = _to_float(p.get("precio_unitario"), 0.0)
        # Si no hay precio unitario pero si importe total y cantidad, deducirlo
        if precio == 0:
            imp = _to_float(p.get("importe") or p.get("total"), 0.0)
            if imp and cant:
                precio = round(imp / cant, 2)
            elif imp and not cant:
                cant = 1.0
                precio = imp
        if cant == 0:
            cant = 1.0
        partidas.append({
            "codigo": "",
            "descripcion": (p.get("descripcion") or "").strip(),
            "unidad": (p.get("unidad") or "ud").strip() or "ud",
            "cantidad": cant,
            "precio_unitario": precio,
            "importe": round(cant * precio, 2),
            "tarifa_encontrada": False,
            "nota": "",
        })
    parsed["partidas"] = partidas
    parsed["_n_partidas"] = len(partidas)
    parsed["_truncado"] = truncado
    return parsed


_SUBC_DATOS_FIELDS = ["empresa", "cif", "domicilio", "telefono", "email",
                      "rep_nombre", "rep_dni", "rep_domicilio", "notario",
                      "notario_loc", "fecha_escritura", "protocolo", "registro"]

_SUBC_DATOS_SCHEMA = """{
  "empresa": "razon social completa CON forma juridica (S.L., S.L.U., S.A., etc.)",
  "cif": "CIF o NIF de la empresa",
  "domicilio": "domicilio social completo (calle, numero, CP y localidad)",
  "telefono": "telefono de contacto",
  "email": "correo electronico",
  "rep_nombre": "nombre y apellidos del administrador o representante legal",
  "rep_dni": "DNI/NIE del representante legal",
  "rep_domicilio": "domicilio del representante a efectos de notificaciones",
  "notario": "nombre del notario de la escritura de constitucion",
  "notario_loc": "localidad del notario",
  "fecha_escritura": "fecha de la escritura en formato DD/MM/AAAA",
  "protocolo": "numero de protocolo notarial",
  "registro": "datos del Registro Mercantil (Tomo, Folio, Hoja)"
}"""

_SUBC_DATOS_SYSTEM = (
    "Eres un asistente administrativo de Acometidas Europa Saneamiento Tecnico S.L. "
    "Extraes los DATOS IDENTIFICATIVOS de una empresa subcontratista a partir de documentos "
    "(escrituras, presupuestos, facturas, certificados, tarjeta CIF, etc.). "
    "Devuelves SIEMPRE un unico JSON valido, sin texto adicional ni markdown. "
    "Si un dato no aparece en el documento, deja la cadena vacia. NUNCA inventes datos."
)


def extract_subcontrata_datos(files: list[Path], api_key: str = "") -> dict:
    """Extract the subcontractor identity/legal data from one or more documents."""
    key = os.environ.get("ANTHROPIC_API_KEY", "").strip() or api_key.strip()
    if not key:
        raise ValueError("API Key de Anthropic no configurada.")
    client = anthropic.Anthropic(api_key=key)

    content: list[dict] = [{"type": "text", "text": (
        "Extrae los datos identificativos de la empresa subcontratista de estos documentos.\n"
        f"Devuelve UNICAMENTE un JSON con este esquema:\n{_SUBC_DATOS_SCHEMA}\n\nDOCUMENTOS:\n"
    )}]

    img_ext = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
    for fp in files:
        suf = fp.suffix.lower()
        if suf in img_ext:
            data, mt = _encode_image(fp)
            content.append({"type": "image", "source": {"type": "base64", "media_type": mt, "data": data}})
        elif suf == ".pdf":
            content.append({"type": "document",
                            "source": {"type": "base64", "media_type": "application/pdf", "data": _encode_pdf(fp)}})
        elif suf in (".docx", ".doc"):
            try:
                from docx import Document
                doc = Document(str(fp))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                for tbl in doc.tables:
                    for row in tbl.rows:
                        text += "\n" + "\t".join(c.text for c in row.cells)
            except Exception:
                text = fp.read_text(errors="replace")
            content.append({"type": "text", "text": f"\n--- DOCUMENTO: {fp.name} ---\n{text[:12000]}\n"})
        else:
            try:
                text = fp.read_text(encoding="utf-8", errors="replace")
                content.append({"type": "text", "text": f"\n--- DOCUMENTO: {fp.name} ---\n{text[:12000]}\n"})
            except Exception:
                pass

    response = client.messages.create(
        model=MODEL,
        max_tokens=2048,
        system=_SUBC_DATOS_SYSTEM,
        messages=[{"role": "user", "content": content}],
    )

    raw = response.content[0].text.strip()
    raw = re.sub(r"^```[a-z]*\n?", "", raw, flags=re.MULTILINE)
    raw = re.sub(r"```$", "", raw, flags=re.MULTILINE)
    m = re.search(r"\{[\s\S]*\}", raw)
    data = json.loads(m.group(0) if m else raw.strip())
    return {k: str(data.get(k, "") or "").strip() for k in _SUBC_DATOS_FIELDS}
